# -*- coding: utf-8 -*-
"""산출물 10종 전수 검사기.

눈으로 훑으면 매번 뚫린다. 기계가 잡을 수 있는 것은 기계가 잡는다.
새 결함이 나오면 «먼저 여기에 검사 항목을 추가»하고 그다음 고쳐라.

    python audit_docs.py
"""
import datetime
import glob
import os
import pathlib
import re
import sys
from collections import Counter, defaultdict

# 한글 윈도(cp949)에서 그대로 돌리면 「—」 같은 글자에서 출력이 죽는다(2026-08-29 실측).
# `npm run check:docs` 가 그렇게 죽으면 «검사가 통과한 것처럼» 보이지 않고 그냥 오류만 남는다.
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

try:
    from docx import Document
except ModuleNotFoundError:
    # 🛑 트레이스백으로 죽지 마라 — 「고장」인지 「환경에 없음」인지 구별이 안 된다.
    #    2026-08-31 정리 실측: 이 검사는 `python-docx` 가 없어 매번 트레이스백으로 죽었고,
    #    어떤 워크플로·훅도 안 부르고 있어 «있는 줄 알지만 한 번도 안 도는 검사»였다.
    print("⏭️  정부과제 문서 감사 — **못 쟀다**: python-docx 가 이 환경에 없다.")
    print("   재려면: pip install python-docx  (그다음 npm run check:docs)")
    print("   ⚠️ 이 줄이 보이면 「통과」라고 적지 마라 — 검사한 적이 없다.")
    sys.exit(0)

HERE = os.path.dirname(os.path.abspath(__file__))
# 「미래 날짜」 판정 기준은 «돌리는 날»이다. 코드에 박아 두면 날이 갈수록 과거를 미래로 잡는다
# (2026-08-29 실측: 8/21 로 박혀 있어 이미 지난 8/25 가 4건 오탐으로 잡혔다).
_오늘 = datetime.date.today()
기준일 = (_오늘.year, _오늘.month, _오늘.day)
문제 = []


def 알림(문서, 항목, 내용):
    문제.append((문서, 항목, 내용))


def 셀들(doc):
    for ti, t in enumerate(doc.tables):
        for ri, r in enumerate(t.rows):
            for ci, c in enumerate(r.cells):
                yield ti, ri, ci, c.text.strip()


def 문단들(doc):
    for p in doc.paragraphs:
        s = p.text.strip()
        if s:
            yield s


def 전체글(doc):
    return " ".join(list(문단들(doc)) + [s for *_, s in 셀들(doc)])


# ── 1. 약점 자진신고 ─────────────────────────────────────────────
# 이 산출물은 「제출 서식」이 아니라 우리 성과를 보이는 자료다(2026-08-21 PO 지시).
# 묻지도 않은 약점을 우리 손으로 제목·판정 칸에 세우지 않는다.
자진신고 = [
    (r"미달", "「미달」 판정을 우리 문서에 직접 씀"),
    (r"개선\s*과제", "「개선 과제」 = 아직 못 했다는 자백"),
    (r"공백이\s*남아|남아\s*있음", "남은 구멍을 자진 신고"),
    (r"미흡|부실|불충분|부족함", "자기비하 표현"),
    (r"오진|틀렸|정정\s*내역|반성문", "내부 반성 문구가 대외 문서에 샘"),
    (r"미검증|측정\s*못|확인\s*필요|미정|TBD|추후\s*결정|예정임", "빈칸·미완성 표시"),
]
# 오탐 제외 — 우리 약점이 아니라 시장·환자 상황이거나 판정 기준 정의인 것
자진신고_예외 = [
    "장비기준 미달", "과업 목적 미달", "언어 장벽", "인프라 부족",
    "정보 부족", "미달 시", "노후",
    # 판정 기준 정의문 — 자기 약점 서술이 아니라 «채점 잣대»다. 오히려 신뢰도의 근거.
    "주된 요구는 반영되었으나", "이를 이행으로 간주하지 아니함",
    "실제 수행 사례가 없는 항목은",
    # 「등록」과 「평가인증」은 «다른 제도»다. 가산점 3점의 기준은 평가인증이므로
    # 등록 완료만으로 인증을 주장하면 허위가 된다. 이 구분은 일부러 남긴다.
    "인증은 별도 확인 필요",
]

# 이름을 대는 것 자체가 문제가 아니라 «굳이 알릴 필요가 없는 것»이 문제다.
# 아래는 실재하는 파일명이라 바꿀 수 없다(파일명을 문서에서만 다르게 적으면 그게 거짓).
# 설명 칸에서 도구 이름을 뺐고 파일명만 남겼다 — 이 상태를 정상으로 둔다.
내부용어_예외 = ["CLAUDE.md"]


def 검사_자진신고(f, doc):
    for s in list(문단들(doc)) + [s for *_, s in 셀들(doc)]:
        if not s or any(e in s for e in 자진신고_예외):
            continue
        for pat, 사유 in 자진신고:
            if re.search(pat, s):
                알림(f, "약점 자진신고", f"{사유} ▸ {s[:85]}")
                break


# ── 2. 표 구조 ───────────────────────────────────────────────────
def 검사_표구조(f, doc):
    for ti, t in enumerate(doc.tables):
        if not t.rows:
            알림(f, "표 구조", f"표{ti}: 행이 없다")
            continue
        폭 = len(t.rows[0].cells)
        for ri, r in enumerate(t.rows):
            if len(r.cells) != 폭:
                알림(f, "표 구조", f"표{ti} {ri}행: 칸 수 {len(r.cells)} != 머리글 {폭}")
        if any(not c.text.strip() for c in t.rows[0].cells):
            빈 = [i for i, c in enumerate(t.rows[0].cells) if not c.text.strip()]
            알림(f, "표 구조", f"표{ti}: 머리글 빈 칸 {빈}")


# ── 3. 표 안 계산 검산 ───────────────────────────────────────────
합계어 = ("합계", "계", "소계", "총계")


def 검사_합계(f, doc):
    for ti, t in enumerate(doc.tables):
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        for ri, row in enumerate(rows):
            if not row or row[0] not in 합계어:
                continue
            윗 = [r for r in rows[1:ri] if r and r[0] not in 합계어]
            if not 윗:
                continue
            for ci in range(1, len(row)):
                if not re.fullmatch(r"[\d,]+", row[ci] or ""):
                    continue
                뽑 = [r[ci] for r in 윗 if ci < len(r) and re.fullmatch(r"[\d,]+", r[ci] or "")]
                if len(뽑) < 2:
                    continue
                합 = sum(int(v.replace(",", "")) for v in 뽑)
                적힌 = int(row[ci].replace(",", ""))
                if 합 != 적힌:
                    알림(f, "표 계산", f"표{ti} 합계 {ci}열: 적힌 {적힌} != 실제 합 {합}")


# ── 4. 판정 개수 vs 요약 숫자 ────────────────────────────────────
판정어 = ("충족", "부분충족", "대안반영", "미반영")


def 검사_판정정합(f, doc):
    # ⚠️ 판정이 적힌 표가 여럿이다(주 대비표 + 잔여사항표). 잔여사항표는 주 대비표의
    #    일부 항목을 «다시» 싣는 것이라 합치면 이중 계산이 된다. 가장 큰 표 하나만 센다.
    후보 = []
    for t in doc.tables:
        머리 = [c.text.strip() for c in t.rows[0].cells] if t.rows else []
        if "정의" in 머리 or not any(h in ("ID", "요구 ID") for h in 머리):
            continue
        c = Counter()
        for r in t.rows[1:]:
            for v in [x.text.strip() for x in r.cells]:
                if v in 판정어:
                    c[v] += 1
                    break
        if c:
            후보.append(c)
    개별 = max(후보, key=lambda c: sum(c.values())) if 후보 else Counter()
    if not 개별:
        return
    for t in doc.tables:
        for r in t.rows:
            v = [c.text.strip() for c in r.cells]
            if v and v[0] in ("합계", "계") and len(v) >= 6:
                try:
                    적힌 = {"충족": int(v[2]), "부분충족": int(v[3]),
                            "대안반영": int(v[4]), "미반영": int(v[5])}
                except ValueError:
                    continue
                for k, n in 적힌.items():
                    if 개별.get(k, 0) != n:
                        알림(f, "판정 불일치",
                             f"요약표 「{k} {n}건」 != 개별 항목 실제 {개별.get(k, 0)}건")


# ── 5. 같은 꼬리표에 다른 숫자 (문서 간) ─────────────────────────
꼬리표 = [
    ("제휴 의료기관 수", r"제휴\s*의료기관[^\d]{0,6}(\d+)\s*(?:개|곳)"),
    ("지원 언어 수", r"(\d+)\s*개\s*국?어|(\d+)\s*개\s*언어"),
    ("실제 문의 건수", r"문의\s*(\d+)\s*건"),
    ("접근성 점수", r"접근성[^\d]{0,6}(\d+)\s*점"),
    ("단위 시험 건수", r"단위\s*(?:테스트|시험)[^\d]{0,20}([\d,]{3,})\s*건"),
    # 「유치 목표」는 뺐다 — 11번 표에 「공고 최소 10건 / 수행계획서 목표 12건」이 나란히
    # 놓여 있어 한 값으로 읽으면 반드시 오탐이 난다(2026-08-21 확인).
]


def 수집_꼬리표(f, doc, 통):
    본문 = 전체글(doc)
    for 이름, pat in 꼬리표:
        for m in re.finditer(pat, 본문):
            값 = next((g for g in m.groups() if g), None)
            if 값:
                통[이름].add((값.replace(",", ""), f))


# ── 6. 미래 날짜 ─────────────────────────────────────────────────
def 검사_날짜(f, doc):
    본문 = 전체글(doc)
    for m in re.finditer(r"20(2\d)\s*[.\-년]\s*(\d{1,2})\s*[.\-월]\s*(\d{1,2})", 본문):
        y, mo, d = 2000 + int(m.group(1)), int(m.group(2)), int(m.group(3))
        if not (1 <= mo <= 12 and 1 <= d <= 31):
            continue
        if (y, mo, d) <= 기준일:
            continue
        앞 = 본문[max(0, m.start() - 30):m.start()]
        # 사업 기간·목표·예정은 미래라서 정상이다. 「실측·완료·기준」에 붙은 미래 날짜만 결함.
        if re.search(r"기간|목표|예정|까지|계획|수검|평가", 앞 + m.group(0)):
            continue
        알림(f, "미래 날짜", f"{y}.{mo}.{d} ▸ …{앞}{m.group(0)}")


# ── 7. 내부 용어 새어나감 ────────────────────────────────────────
내부용어 = [
    (r"\bTODO\b|\bFIXME\b|\bXXX\b", "코드 주석 표시"),
    (r"scratchpad|node_modules|\.venv|worktree", "내부 경로"),
    (r"어시스턴트|클로드|Claude|GPT", "작성 도구 이름"),
    (r"이번 세션|다음 세션|핸드오프", "세션 작업 용어"),
]


def 검사_내부용어(f, doc):
    for s in list(문단들(doc)) + [s for *_, s in 셀들(doc)]:
        if any(e in s for e in 내부용어_예외):
            continue
        for pat, 사유 in 내부용어:
            if re.search(pat, s, re.I):
                알림(f, "내부 용어", f"{사유} ▸ {s[:80]}")
                break


# ── 8. 빈 칸 ─────────────────────────────────────────────────────
def 검사_빈칸(f, doc):
    for ti, t in enumerate(doc.tables):
        if len(t.rows) < 2:
            continue
        for ri, r in enumerate(t.rows[1:], 1):
            vals = [c.text.strip() for c in r.cells]
            if len(vals) < 2:
                continue
            if vals[0].startswith("—") or vals[0].startswith("["):
                continue          # 표 안 구분줄(머리글 역할) — 비어 있는 게 정상
            if vals[0] and all(not v for v in vals[1:]):
                알림(f, "빈 칸", f"표{ti} {ri}행 「{vals[0][:40]}」 나머지 전부 비어 있음")


# ── 9. 같은 말이 한 칸 안에서 되풀이 ─────────────────────────────
# 제자리 교체(update_reports.py)는 «돌 때마다 덧붙는» 구조라 같은 문장이 2~4번
# 쌓일 수 있다. 2026-08-21 실제로 09 산출물목록에서 한 문장이 두 번 찍혀 있었다.
def 검사_중복문구(f, doc):
    """같은 «줄»이나 같은 «문장»이 한 칸 안에서 두 번 나오면 잡는다.

    앞부분만 같은 것은 세지 않는다 — 표 이름을 「칸 목록」줄과 「파일 위치」줄에
    각각 적는 것은 정상이고, 「면력한방병원」과 「신촌면력한방병원」도 다른 기관이다.
    """
    for s in list(문단들(doc)) + [x for *_, x in 셀들(doc)]:
        if len(s) < 24:
            continue
        줄 = [x.strip() for x in s.splitlines() if len(x.strip()) >= 15]
        겹줄 = [x for x, n in Counter(줄).items() if n > 1]
        if 겹줄:
            알림(f, "문구 중복", f"같은 줄이 두 번 ▸ {겹줄[0][:60]}")
            continue
        문장 = [x.strip() for x in re.split(r"[.。]\s*", s) if len(x.strip()) >= 15]
        겹문 = [x for x, n in Counter(문장).items() if n > 1]
        if 겹문:
            알림(f, "문구 중복", f"같은 문장이 두 번 ▸ {겹문[0][:60]}")


# ── 10. 워드판 ↔ 웹판 대조 ───────────────────────────────────────
# 웹판(.html)은 손으로 만든 시각화판이라 «자동 생성이 안 된다». 2026-08-21 실제로
# 워드판은 90.9% 인데 웹판은 86.4% 로 조용히 어긋나 있었다(같은 산출물, 다른 말).
def 검사_웹판대조():
    for hp in sorted(glob.glob(os.path.join(HERE, "*.html"))):
        dp = hp[:-5] + ".docx"
        if not os.path.exists(dp):
            continue
        이름 = os.path.basename(hp)
        html = pathlib.Path(hp).read_text(encoding="utf-8", errors="ignore")
        doc = Document(dp)
        # 요구항목 판정 개수를 양쪽에서 세어 맞대 본다
        웹 = Counter(v for _, v in re.findall(
            r'class="rid mono">([A-Z]+-[A-Z]+-\d+)</span>.*?class="verdict [a-z]+">([^<]+)</span>', html))
        워드 = Counter()
        후보 = []
        for t in doc.tables:
            머리 = [c.text.strip() for c in t.rows[0].cells] if t.rows else []
            if "정의" in 머리 or not any(h in ("ID", "요구 ID") for h in 머리):
                continue
            c = Counter()
            for r in t.rows[1:]:
                for v in [x.text.strip() for x in r.cells]:
                    if v in 판정어:
                        c[v] += 1
                        break
            if c:
                후보.append(c)
        if 후보:
            워드 = max(후보, key=lambda c: sum(c.values()))
        if 웹 and 워드 and 웹 != 워드:
            알림(이름, "웹판 어긋남", f"판정 개수 웹 {dict(웹)} != 워드 {dict(워드)}")
        # 백분율 값도 맞대 본다 — 단, style="width:..%" 같은 «화면 꾸밈값»은 뺀다.
        # 막대 그래프 너비는 워드판에 있을 수가 없어 그대로 두면 영원히 오탐이 난다.
        본문만 = re.sub(r"<[^>]*>", " ", html)
        웹률 = set(re.findall(r"(\d{1,3}\.\d)%", 본문만))
        워드글 = 전체글(doc)
        워드률 = set(re.findall(r"(\d{1,3}\.\d)%", 워드글))
        빠진 = 웹률 - 워드률
        if 빠진:
            알림(이름, "웹판 어긋남", f"웹판에만 있는 비율값 {sorted(빠진)} — 워드판 갱신 누락 의심")


def main():
    통 = defaultdict(set)
    파일 = sorted(glob.glob(os.path.join(HERE, "*.docx")))
    for path in 파일:
        f = os.path.basename(path)
        doc = Document(path)
        for fn in (검사_자진신고, 검사_표구조, 검사_합계, 검사_판정정합,
                   검사_날짜, 검사_내부용어, 검사_빈칸, 검사_중복문구):
            fn(f, doc)
        수집_꼬리표(f, doc, 통)

    검사_웹판대조()

    for 이름, 값들 in 통.items():
        if len({v for v, _ in 값들}) > 1:
            상세 = " / ".join(f"{v}({os.path.splitext(fn)[0][:16]})" for v, fn in sorted(값들))
            알림("(문서 간)", "숫자 불일치", f"{이름}: {상세}")

    print(f"산출물 전수 검사 — 문서 {len(파일)}개\n")
    if not 문제:
        print("  걸린 것 없음")
        return 0
    묶음 = defaultdict(list)
    for 문서, 항목, 내용 in 문제:
        묶음[항목].append((문서, 내용))
    for 항목 in sorted(묶음, key=lambda k: -len(묶음[k])):
        print(f"■ {항목} — {len(묶음[항목])}건")
        for 문서, 내용 in 묶음[항목][:14]:
            print(f"   [{os.path.splitext(문서)[0][:24]}] {내용}")
        if len(묶음[항목]) > 14:
            print(f"   … 외 {len(묶음[항목]) - 14}건")
        print()
    print(f"합계 {len(문제)}건")
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
