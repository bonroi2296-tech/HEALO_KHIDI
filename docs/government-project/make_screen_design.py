# -*- coding: utf-8 -*-
"""01-1_화면설계서.docx 생성.

왜 만드나 : 사업계획서(2026-04-15 제출본) 1단계 산출물에 「요구사항 정의서 · 화면 설계서」가
명시돼 있는데 화면 설계서만 실물이 없었다(2026-08-19 확인). 이 파일이 그 산출물이다.

원칙
  · lo-fi 유지: 색·폰트가 들어간 실제 화면 캡처는 「결과물」이지 「설계」가 아니다.
    그림은 scripts/docs/make_wireframes.py 가 회색 상자로만 그린다.
  · 실선 = 이미 만들어 돌아가는 것 / 점선 = 아직 만들지 않은 것. 숨기지 않는다.
  · 수치는 _facts.py 와 운영DB 실측만 쓴다(추정 금지).

쓰는 법:  python docs/government-project/make_screen_design.py
"""
import os
import subprocess
import sys

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import _facts as F

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
WF_DIR = os.path.join(HERE, "wireframes")
WF_MAKER = os.path.join(ROOT, "scripts", "docs", "make_wireframes.py")
OUT = os.path.join(HERE, "01-1_화면설계서.docx")


# ── 그림 먼저 그린다(설계서와 그림이 어긋나지 않게 항상 같이 만든다) ──────────
def build_wireframes():
    os.makedirs(WF_DIR, exist_ok=True)
    subprocess.run([sys.executable, WF_MAKER, WF_DIR], check=True,
                   stdout=subprocess.DEVNULL)


# ── docx 서식 (다른 산출물 생성기와 같은 서식을 쓴다) ────────────────────────
def set_font(run, size=10, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rPr.append(rFonts)


def heading(doc, text, size=14, color=(0, 70, 127)):
    p = doc.add_paragraph()
    set_font(p.add_run(text), size=size, bold=True, color=color)
    return p


def para(doc, text, size=10, indent=0, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    set_font(p.add_run(text), size=size, color=color)
    return p


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        set_font(c.paragraphs[0].add_run(h), size=9, bold=True)
        set_cell_bg(c, "E8F0F2")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            set_font(cells[i].paragraphs[0].add_run(str(v)), size=9)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    return t


# ── 화면 정의 (근거는 전부 코드·운영DB 실측 2026-08-19) ─────────────────────
SCREENS = [
    dict(
        no="P-01", file="P-01_home.png", name="홈(다국어 진입)",
        ict="① 병원 안내 및 매칭", path="/  (언어별 /ru · /kz · /en …)",
        who="누구나(비로그인)",
        purpose="카자흐스탄·러시아권 환자가 자국어로 첫 화면을 만나고, 두 갈래(바로 대화 / 문의 남기기)"
                " 중 하나로 즉시 진입하게 한다.",
        parts=[
            "상단 바: 로고 + 메뉴 6개(원격협진·진료 과목·병원·치료 여정·비자 가이드·보험 가이드) + [언어 변경] + [로그인] [회원가입]. 폭이 좁은 화면에서는 [≡] 안으로 접힌다",
            "첫 화면 큰 배너(히어로): 큰 제목 「한국 최고 종양 전문의의 제2의견」 + 부제 3줄 + 단추 [무료 상담] 하나",
            "본문 구역 10개: 왜 한국인가 / 우리 의료팀 / healwith 가 하는 일 / 어떻게 진행되나 / "
            "주요 암종 / 파트너 병원 / 진단부터 회복까지 동행 / 자주 묻는 질문 / 긴급 도움 / 오늘 시작하기",
            "우하단: 화면에 떠 있는 대화 단추(「도움말」 말풍선). AI 사전상담으로 들어가는 입구다",
            "맨 아래: 유치업 등록번호 · 보증보험 · 법정 고지 · 개인정보 처리방침",
            "첫 방문 시 화면 아래 쿠키 설정 띠",
        ],
        data=[
            "다국어 사전(i18n): 화면이 쓰는 문구 1,941개를 6개 언어에 모두 채웠다. 러시아어·카자흐어는 자동 검사로 매 변경마다 100%를 확인한다(2026. 8. 20. 실측)",
            "병원(hospitals): 발행 8곳(협진 대학병원 4 + 면력한방병원 4지점). 나머지 1곳은 시험용 등록이라 공개하지 않는다",
            "화면 설정(site_settings): 노출 문구·배너를 관리자 화면에서 고친다",
        ],
        nexts="→ P-02 문의 갈림길 · P-03 AI 사전상담(우하단 대화 단추)",
        notes=["언어 전환은 주소 앞자리(/ru, /kz …)를 바꾸며 선택은 다음 방문까지 유지된다.",
               "모든 문구는 사전에서 나오고 화면에 직접 쓰는 것(하드코딩)은 자동 검사로 막는다."],
    ),
    dict(
        no="P-02", file="P-02_inquiry.png", name="문의(연락 방법 고르기 → 신청서)",
        ict="② 사전 진료의뢰 및 상담", path="/inquiry  (한 화면에서 단계 전환)", who="누구나(비로그인)",
        purpose="먼저 「연락하기 편한 방법」을 고르게 해 문턱을 낮추고, 신청서를 고른 사람만 입력 양식으로 보낸다. "
                "대학병원 의뢰에 필요한 자료는 접수 뒤 2단계에서 받는다.",
        parts=[
            "갈림길: ① AI 상담원(러시아어 24시간) / ② 실시간 상담원(메신저) / ③ 신청서 작성",
            "신뢰 표시: 무료 상담 · 병원에 직접 결제(수수료 없음) · AES-256 암호화 · 1영업일 내 회신 · "
            "KHIDI 지원사업 · 등록 유치업자",
            "③을 고르면 1단계 접수: 이름 · 연락 수단 · 국적 · 선호 언어 / 암종 · 병기 / 하고 싶은 말",
            "동의 4종(필수): 개인정보 수집 · 민감정보(건강) · 제3자 제공(병원) · 국외이전",
            "접수 뒤 2단계: 별도 주소 /inquiry/intake (문의번호와 접근 토큰이 있어야 한다). 첨부 최대 10개, 건당 200MB",
        ],
        data=[
            "문의(inquiries): 실제 환자 문의 8건. 자동 검사가 만든 것을 포함한 시험 문의는 표식으로 갈라 실적에서 뺀다",
            "환자 상세(cancer_patient_intakes): 암종·병기·치료 이력을 담는다",
            "첨부: 브라우저에서 저장소로 직행 업로드(건당 최대 200MB)",
        ],
        nexts="→ P-06 코디네이터 인박스(접수 즉시 표시·알림)",
        notes=["/inquiry 는 폼이 아니라 「갈림길」이 먼저 나온다.",
               "2단계는 접수한 뒤에만 들어갈 수 있다. 문의번호나 접근 토큰 없이 열면 값이 없다는 안내가 뜬다.",
               "1단계만 채워도 접수된다. 자료가 없어 못 보내는 일을 없앤다.",
               "시험 문의는 표식이 붙어 목록·알림·성과지표에서 모두 빠진다."],
    ),
    dict(
        no="P-03", file="P-03_ai_chat.png", name="AI 사전상담 (근거 검색형)",
        ict="① · ②", path="홈 위젯 · 메신저(텔레그램)", who="누구나(비로그인)",
        purpose="업무시간 밖 문의를 24시간 받되, 지어낸 답이 나가지 않도록 자체 DB에서 찾은 근거로만 답한다.",
        parts=[
            "대화 영역: 환자/AI 말풍선",
            "근거 패널: DB에서 찾은 문서 3~5건(병원명·시술·출처)",
            "안전장치: DB에 없는 병원은 추천 불가, 연락처 없이 「접수됨」이라 말하지 않음",
            "입력창: 6개 언어 자동 감지",
        ],
        data=[
            "AI 상담 대화(chat_threads) 326건, 주고받은 말(chat_messages) 1,068건",
            "근거 문서(rag_documents) 21건. 병원·진료 정보를 담으며, AI 는 답할 때 이 안에서만 근거를 찾는다",
            "품질 자동 시험 기록(ai_response_evaluations, ai_regression_runs)",
        ],
        nexts="→ 3턴을 넘기면 실제 문의로 자동 승격 → P-06 코디네이터 인박스",
        notes=["동작: 질문 → 의미 변환 → 자체 DB 검색 → 근거 읽기 → 출처와 함께 답변(2~3초).",
               "품질은 매일 자동 시험한다. 의료 레드라인(완치 단정·지어낸 병원) 위반을 검출해 경보한다."],
    ),
    dict(
        no="P-04", file="P-04_consult.png", name="원격협진 상담방",
        ict="② · ④", path="/c/<초대토큰>", who="환자(비로그인 초대) · 코디네이터 · 의료진",
        purpose="설치 없이 브라우저만으로 한국 의료진과 환자가 만나고, 언어 장벽은 실시간 자막으로 없앤다.",
        parts=[
            "영상 영역: 상대 화면(크게) / 내 화면(작게)",
            "참가자·대기실: 입장 요청 승인·거절",
            "실시간 자막: 상대 발화를 내 언어로(한 줄 스택, 가운데 정렬)",
            "자막 켜기·끄기: 누를 때만 켜진다(자동 켜짐 없음). 잡음 제거는 단추 없이 자동으로 걸린다",
            "도구 막대: 마이크 · 카메라 · 화면 공유 · 채팅 · 언어 · 자막 크기 · 음성 출력 · [종료]",
        ],
        data=[
            "상담 세션(consultation_sessions) 133건, 입퇴장 기록 325건",
            "통역 자막(consultation_translations) 3,277건(시험 방 제외). 러→한 2,487 · 한→러 393 · 카자흐→한 117 등",
            "상담 기록·자막은 암호화 저장(AES-256)",
        ],
        nexts="→ P-05 소견 확인 · P-07 케이스 상세(기록 축적)",
        notes=["성과지표는 코디네이터가 「상담 일정」 화면에서 [상담 완료]를 눌러야 잡힌다. 상담방에서 [종료]만 누르면 세지 않는다.",
               "초대는 57자짜리 주소로 나가고 계정 없이 들어올 수 있다. 방은 참가자 2명부터 시작으로 본다."],
    ),
    dict(
        no="P-05", file="P-05_opinion.png", name="소견 확인(환자용)",
        ict="② · ③", path="/opinion/<토큰>", who="환자(로그인 없이 링크로)",
        purpose="병원이 발급한 공식 소견 문서를 환자 언어로 안전하게 전달한다.",
        parts=[
            "환자 안내 머리말: 담당 의료진·작성일·소속 병원",
            "소견 문서 미리보기: 병원이 발급한 공식 문서 그대로",
            "언어: 환자 언어 자동 선택(문의서 언어 → 브라우저 → 영어)",
            "첨부: 검사지·영상 판독본",
            "다음 단계 안내: 추가 자료 요청·원격협진 예약·방한 일정 문의",
        ],
        data=[
            "제2의료소견(case_opinions) 6건. 전건 환자에게 전달 완료. 소견 요청(opinion_requests) 2건",
            "서류 번역(attachment_translations) 12건. 외국 검사지와 진료기록을 AI 가 읽어 표의 행과 칸까지 구조를 살려 옮긴다",
        ],
        nexts="→ P-04 원격협진 예약 · P-08 사후관리",
        notes=["환자에게 전달한 시점이 성과지표 판정 기준이다. 작성만 하면 세지 않는다."],
    ),
    dict(
        no="P-06", file="P-06_inbox.png", name="코디네이터 인박스(문의함)",
        ict="운영 화면(① ~ ⑥ 공통)", path="/coordinator/inbox  (언어 앞자리 없음)", who="코디네이터",
        purpose="들어온 문의를 한 곳에서 받고, 기준일을 넘긴 케이스를 눈에 띄게 해 다음에 손댈 것을 놓치지 않게 한다.",
        parts=["왼쪽 메뉴 15개: 대시보드 · 문의함 · AI 상담 리드 · 의뢰·케이스/병원배정 · 상담 일정 · "
               "유치 전환 · 사후관리·만족도 · 파트너 발굴 · 메시지 · 비자 트래킹 · 견적 · 증상 알림 · "
               "콘텐츠 편집 · 개선 요청함 · 설정",
               "거르기 세 갈래: 전체 / 추가 정보 필요 / 매칭 준비 완료 (각 건수 표시)",
               "목록 표: 이름 | 국적 | 암종 | 연락방법 | Step 완료(1만 / 1+2) | 접수일 | 상태",
               "상태 칸에 빨간 배지 「N일째 정체」: 단계 기준일을 넘긴 케이스를 표시"],
        data=["문의(inquiries): 목록에 뜨는 건수 8건. 실환자 수와 같다(시험 문의는 자동으로 빠진다)",
              "알림(notifications) 1,043건. 새 문의가 들어오면 즉시 발송한다"],
        nexts="→ P-07 케이스 상세",
        notes=["시험 문의는 목록·알림·성과지표에서 모두 빠진다."],
    ),
    dict(
        no="P-07", file="P-07_case.png", name="케이스 상세(문의 하나를 끝까지)",
        ict="② · ③", path="/coordinator/inbox/[번호]", who="코디네이터",
        purpose="환자 자료를 한국 의료진이 읽을 수 있는 형태로 만들어 협진 병원에 의뢰하고, 받은 소견을 환자 언어로 되돌린다.",
        parts=["머리말: 환자 이름 · 회원/비회원 · 담당자 · 문의 번호 · 접수 시각 · Step 완료 · 상태",
               "케이스 브리프(AI 초안): 환자 요약 · 원하는 것 · 「코디가 볼 포인트」 · 「주의」 + "
               "「진단이 아니며 검수 필요」 고지",
               "연락 정보 / 의료·여정 정보 / 문의 메시지 원문 / 추가 정보(인테이크)",
               "동의 항목 5종: 개인정보 · 민감정보(건강) · 국경 간 이전 · 제3자(병원) 제공 · 마케팅(선택) "
               "+ 동의 시각(KST) · 동의서 판번호",
               "첨부 서류 + [전부 받기] · 파일마다 [한] [EN] [RU] 변환: AI 문서 번역 진입점",
               "환자 대신 서류 올리기: PDF·이미지·Word 200MB / ZIP·RAR·DICOM 200MB",
               "전문의 소견: [소견 요청 링크 만들기] · [이미 받은 소견 직접 입력] → 도착 소견 목록 → "
               "환자 언어 자동 번역 초안 → 코디 교정 후 [에이전시에 공개]"],
        data=["협진 의뢰(cotreatment_referrals) 4건, 공유 서류(case_shared_documents) 8건, 제2의료소견(case_opinions) 6건 전부 전달 완료",
              "서류 번역(attachment_translations) 12건. 표의 행과 칸까지 구조를 살려 옮기되 숫자·단위·정상범위는 원문 그대로 둔다"],
        nexts="→ P-05 소견 확인(환자) · P-04 원격협진",
        notes=["「접수 후 추가 정보」는 저장할 때 암호화된다.",
               "번역은 요약이 아니라 1:1 옮김이다. 한국 의료진이 원본과 대조할 수 있어야 한다."],
    ),
    dict(
        no="P-08", file="P-08_patient.png", name="사후관리(환자 화면)",
        ict="④ · ⑤ · ⑥", path="/patient  (로그인)", who="환자",
        purpose="귀국 뒤에도 한국 의료진과 끊기지 않게 경과를 주고받고, 재방문 시점을 놓치지 않게 한다.",
        parts=[
            "인사말 + 알림 띠 「화상 상담 예약이 있습니다 · 날짜」 [대기실 입장]",
            "빠른 메뉴 6종: ①내 상담 ②의료 문서(/patient/documents) ③암 치료 가이드(/education) "
            "④증상 기록(/patient/symptoms) ⑤재진 예약(/patient/rebooking) ⑥비자 가이드(/visa)",
            "내 문의 · 내 상담 목록: 접수한 문의와 잡힌 상담이 이어서 표시",
            "휴대전화에서는 아래 고정 이동 막대(홈 · 의료 문서 · 더보기)",
            "화면 뒤에서 자동으로 도는 것: AI 가 증상 기록을 읽어 위험도를 매기고 담당자에게 알린다",
            "[남은 것] 단계별 교육 자동 발송: 귀국 후 시기에 맞춰 보내는 부분은 아직 만들지 않았다",
        ],
        data=[
            "교육 콘텐츠(education_contents) 18건. 위·유방·간·폐·갑상선 5종이며 러시아어 본문까지 전건 채워 발행했다",
            "증상 기록(symptom_reports) 1건 · 사후관리 일정(followup_schedules) 0건. 기능은 완성했으나 유치 확정 환자가 아직 없어 실사용 기록이 없다",
            "알림 기기 등록(device_tokens) 11건",
        ],
        nexts="→ P-04 원격 경과상담 · P-09 성과 대시보드",
        notes=["교육 콘텐츠는 환자 전용 화면이 아니라 공개 화면 /education 을 그대로 쓴다. 6개 언어로 실제 서비스하고 있다.",
               "「이상 징후 알림」은 환자 화면에 있는 칸이 아니라, 화면 뒤에서 자동으로 도는 판정이다."],
    ),
    dict(
        no="P-09", file="P-09_dashboard.png", name="유치 전환 상세 (성과 대시보드)",
        ict="성과지표 자동집계", path="/admin/khidi/conversion", who="관리자",
        purpose="성과지표를 사람이 세지 않고 화면이 세게 한다. 보고서 숫자와 화면 숫자가 갈리지 않도록 한 곳에서 집계한다.",
        parts=["기간 고르기: 최근 30일 / 90일 / 1년 + 「실적만」 스위치(켜면 시험 데이터 제외)",
               "유치 깔때기: 문의 접수 → 사전상담 완료 → 견적·비자 진행 → 유치 확정 → 사후관리 완료",
               "전환율 카드 3종: 문의→사전상담 · 사전상담→유치 · 전체 유치율",
               "국가별 표 / 채널별(유입경로) 표 / 기관별 상담·사후관리 표",
               "유치 확정 대기: [유치 확정] [이탈] [테스트로 표시] 로 결과를 코디가 입력"],
        data=[f"{F.AS_OF} 기준: 문의 8건, 사전상담 완료 1건(12.5%), 유치 0건, 사후관리 0건",
              "국가별 KZ 5 · ET 1 · KG 1 · RU 1 / 채널별 웹 문의폼 6 · 소급등록 2",
              "성과 기록(kpi_snapshots) 94건. 목표값은 파일 한 곳에만 두어 화면과 문서가 어긋나지 않게 한다"],
        nexts="— (보고서 산출)",
        notes=["깔때기의 사전상담(1건)은 문의에 연결된 건만 세고, 기관별 표의 사전상담(3건)은 전체 세션을 센다. 세는 기준이 다르다.",
               "시험 데이터는 「실적만」 스위치로 갈라 본다."],
    ),
    dict(
        no="P-10", file="P-10_flow.png", name="화면 흐름도 (전체)",
        ict="전체", path="—", who="—",
        purpose="환자·코디네이터·의료진이 어떤 순서로 화면을 지나는지 한 장에 담았다.",
        parts=["환자 동선: 홈 → AI 상담 → 문의 접수 → (코디) → 원격협진 · 소견 확인 → 사후관리",
               "운영 동선: 인박스 → 케이스 상세 → 병원 의뢰 → 회신 등록 → 환자 전달",
               "권한 5종: 환자·코디네이터·에이전시·해외 의료기관·병원"],
        data=["모든 화면의 실적 이벤트가 성과 대시보드(P-09)로 자동 모인다"],
        nexts="—",
        notes=["환자는 로그인 없이도 초대 링크로 상담방(P-04)·소견(P-05)에 들어올 수 있다."],
    ),
]


def main():
    build_wireframes()
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.27), Inches(11.69)
    s.left_margin = s.right_margin = Cm(2.0)
    s.top_margin = s.bottom_margin = Cm(2.0)

    # ── 표지 ────────────────────────────────────────────────────────────
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(F.PROJECT["사업명"]), size=12, bold=True, color=(90, 90, 90))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("화 면 설 계 서"), size=26, bold=True, color=(0, 70, 127))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("(저해상도 뼈대 그림 기반)"), size=12, color=(110, 110, 110))
    for _ in range(4):
        doc.add_paragraph()
    table(doc, ["구분", "내용"], [
        ("문서명", "화면 설계서 (산출물 A-05)"),
        ("플랫폼", F.PROJECT["플랫폼"]),
        ("수행기관", F.PROJECT["수행기관"]),
        ("작성 기준일", F.AS_OF),
        ("근거", "배포 코드·운영데이터베이스 실측"),
    ], widths=[4.0, 12.0])
    doc.add_page_break()

    # ── 1. 문서 개요 ────────────────────────────────────────────────────
    heading(doc, "1. 문서 개요")
    heading(doc, "1.1 목적", size=12)
    para(doc, "본 문서는 HEALO 플랫폼의 화면 구성과 화면 사이 이동을 설계 수준에서 규정한다. "
              "사업계획서 1단계 산출물로 명시된 「요구사항 정의서·화면 설계서」 중 화면 설계서에 해당하며, "
              "요구사항 정의서(01)의 기능 요구를 화면 단위로 배치한 것이다.")
    heading(doc, "1.2 표기 규칙", size=12)
    table(doc, ["표기", "뜻"], [
        ("회색 상자(실선)", "이미 만들어져 실서비스에서 돌아가는 화면 요소"),
        ("점선 상자", "설계는 확정했으나 아직 만들지 않은 화면 요소. 숨기지 않고 그대로 표시한다"),
        ("저해상도 뼈대 그림(lo-fi)", "색과 글꼴을 쓰지 않는다. 무엇이 어디에 있고 무엇을 누르면 어디로 가는지만 그린다. "
                  "색이 들어간 실제 화면 캡처는 「결과물」이지 「설계」가 아니므로 설계서에 넣지 않는다."),
    ], widths=[3.5, 12.5])
    heading(doc, "1.3 화면 목록", size=12)
    table(doc, ["번호", "화면", "경로", "6대 ICT", "이용 주체"],
          [(x["no"], x["name"], x["path"], x["ict"], x["who"]) for x in SCREENS],
          widths=[1.6, 4.4, 4.0, 3.2, 3.0])
    doc.add_page_break()

    # ── 2. 시스템 구성도 ────────────────────────────────────────────────
    # 사업계획서(2026-05-14 제출본) 마일스톤 M1 이 「요구사항 정의서 · UI/UX 설계서 ·
    # 시스템 구성도」를 결과물로 약속했다. 앞 둘은 실물이 있었는데 구성도만 없어서
    # 이 문서 안에 넣는다(2026-08-20). 별도 문서를 새로 만들지 않는 이유는 M1 결과물
    # 둘이 한 문서에 담기는 편이 대조하는 사람에게 낫기 때문이다.
    # ⚠️ 이 한 쪽만 «가로»다. 구성도는 16:9 라 세로 쪽에 넣으면 글자가 읽을 수 없게
    #    작아진다(6.3인치 폭 → 본문 글씨 6pt 수준). 쪽을 눕혀 9.5인치로 넣는다.
    land = doc.add_section(WD_SECTION.NEW_PAGE)
    land.orientation = WD_ORIENT.LANDSCAPE
    land.page_width, land.page_height = Inches(11.69), Inches(8.27)
    land.left_margin = land.right_margin = Cm(1.5)
    land.top_margin = land.bottom_margin = Cm(1.5)

    heading(doc, "2. 시스템 구성도")
    para(doc, "이용자부터 데이터 보관까지 플랫폼 전체가 어떻게 이어지는지 한 장에 담았다. "
              "점선 상자는 바깥 서비스이고 나머지는 직접 만들어 운영한다.")
    arch = os.path.join(WF_DIR, "S-01_architecture.png")
    if os.path.exists(arch):
        doc.add_picture(arch, width=Inches(10.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    back = doc.add_section(WD_SECTION.NEW_PAGE)
    back.orientation = WD_ORIENT.PORTRAIT
    back.page_width, back.page_height = Inches(8.27), Inches(11.69)
    back.left_margin = back.right_margin = Cm(2.0)
    back.top_margin = back.bottom_margin = Cm(2.0)

    # ── 3. 화면별 상세 ──────────────────────────────────────────────────
    heading(doc, "3. 화면별 상세")
    for i, x in enumerate(SCREENS):
        heading(doc, "%s. %s" % (x["no"], x["name"]), size=12)
        table(doc, ["구분", "내용"], [
            ("6대 ICT", x["ict"]),
            ("경로", x["path"]),
            ("이용 주체", x["who"]),
            ("목적", x["purpose"]),
        ], widths=[3.0, 13.0])
        doc.add_paragraph()
        img = os.path.join(WF_DIR, x["file"])
        if os.path.exists(img):
            doc.add_picture(img, width=Inches(6.3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para(doc, "주요 구성요소", size=10)
        for it in x["parts"]:
            para(doc, "· " + it, size=9, indent=0.5)
        para(doc, f"다루는 데이터 ({F.AS_OF} 실측)", size=10)
        for it in x["data"]:
            para(doc, "· " + it, size=9, indent=0.5)
        # 흐름도(P-10)처럼 이어질 화면이 없는 칸은 「—」만 두 줄 먹으므로 통째로 뺀다.
        if x["nexts"].strip() != "—":
            para(doc, "다음 화면 연결", size=10)
            para(doc, x["nexts"], size=9, indent=0.5)
        if x.get("notes"):
            para(doc, "설계 메모", size=10)
            for it in x["notes"]:
                para(doc, "※ " + it, size=9, indent=0.5, color=(110, 110, 110))
        # 마지막 화면 뒤에 넣으면 빈 쪽 한 장이 남는다.
        if i < len(SCREENS) - 1:
            doc.add_page_break()

    # ── 3. 권한별 접근 범위 ────────────────────────────────────────────
    heading(doc, "4. 권한별 화면 접근 범위")
    para(doc, "화면 접근 범위는 역할에 따라 나뉜다. 권한은 서버에서 판정하며, 화면에서 감추는 것만으로는 막지 않는다.", size=9)
    table(doc, ["계층", "권한 저장", "전용 화면", "설명"],
          [(t[1], t[2], t[3], t[4]) for t in F.TIERS], widths=[3.0, 3.0, 4.5, 5.5])

    doc.add_paragraph()
    heading(doc, "5. 후속 고도화 항목", size=14)
    table(doc, ["항목", "화면", "상태", "계획"], [
        ("단계별 교육 콘텐츠 자동 발송", "P-08", "판정 로직 완료 · 화면 연결 남음",
         "2026년 9월 중간평가 이전 완료 목표"),
        ("예약 리마인더: 의료진 일정 연동", "P-08 · P-09", "부분 구현",
         "실환자 유치 확정 시점에 맞춰 연동"),
    ], widths=[5.0, 3.0, 4.5, 3.5])

    doc.save(OUT)
    print("저장 완료:", OUT)
    print("와이어프레임:", WF_DIR)


if __name__ == "__main__":
    main()
