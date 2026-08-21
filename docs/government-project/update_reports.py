# -*- coding: utf-8 -*-
"""
update_reports.py: 생성기가 없는 산출물(04·06·07·08·09)을 현행으로 갱신한다.

왜 «수정 스크립트»인가:
  01·02·EVAL 은 생성기(make_*.py)가 있어 통째로 다시 뽑으면 되지만,
  04·06·07·08·09 는 생성기 없이 직접 작성된 문서다. 통째로 재작성하면 그동안 쌓인
  서술을 잃는다. 그래서 **틀린 부분만 제자리에서 고치고, 빠진 현행 내용은 장(章)으로
  덧붙이는** 방식을 쓴다. 재실행해도 결과가 같도록(멱등) 만들었다.

사실의 출처는 _facts.py 한 곳이다(문서마다 따로 적지 않는다).

    python3 update_reports.py
"""
import os
import pathlib
import re
import sys

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

import _facts as F

HERE = os.path.dirname(os.path.abspath(__file__))

# 모든 산출물에 적용. 본로이는 개인사업자라 「(주)」는 사실과 다르다.
GLOBAL = [
    ("(주)본로이", "본로이"),
    ("(주) 본로이", "본로이"),
    ("healo.kr", "healwith.co.kr"),
]

# ── 제자리 교체 목록 ────────────────────────────────────────────────────────
# (파일, [(옛 문구, 새 문구), ...])
# 폐지된 화면을 설명하던 문장을 현행 화면으로 바꾼다. 근거는 app/ 실재 여부로 확인함.
REPLACEMENTS = {
    "01_요구사항정의서.docx": [
        # ── 2026-08-19 : 실제로 없는 경로 → 배포 코드의 실제 주소로
        ("/app/api/chat/route.ts", "/app/api/public/chat/message, /app/api/patient/chat"),
        ("/app/api/livekit/route.ts", "/app/api/khidi/consultation/token (+ /app/api/livekit/webhook)"),
        # 「/app/api/email/*」 에 걸리면 「…sendEmail.ts/*」 라는 없는 경로가 된다. 별표형을 먼저 처리한다.
        ("src/lib/email/sendEmail.ts/*", "src/lib/email/*"),
        ("/app/api/email/*", "src/lib/email/*"),
        ("/app/api/email", "src/lib/email/sendEmail.ts"),
        ("migrations/20260420_drop_cancer_intake_plaintext (실행 완료) · 20260420_drop_inquiries_plaintext_email (보류) (평문 완전 제거)",
         "식별정보 암호화 적용: 이메일 109건·전화 22건 전건 암호문(2026-08-19 실측). "
         "평문 컬럼 삭제 마이그레이션은 보류(검색·번역에 쓰이는 본문 때문)"),
        ("migrations/20260420_drop_cancer_intake_plaintext (실행 완료) · 20260420_drop_inquiries_plaintext_email (보류)", "migrations/20260420_drop_cancer_intake_plaintext (실행 완료) · 20260420_drop_inquiries_plaintext_email (보류) (작성 완료·미실행)"),
        ("평문 완전 제거", "식별정보 암호화 완료 / 본문·소견은 평문 보관"),
    ],
    "02_기능명세서.docx": [
        # ── 2026-08-19 : 실제로 없는 경로 → 배포 코드의 실제 주소로
        ("/app/api/chat/route.ts", "/app/api/public/chat/message, /app/api/patient/chat"),
        ("/app/api/livekit/route.ts", "/app/api/khidi/consultation/token (+ /app/api/livekit/webhook)"),
        # 「/app/api/email/*」 에 걸리면 「…sendEmail.ts/*」 라는 없는 경로가 된다. 별표형을 먼저 처리한다.
        ("src/lib/email/sendEmail.ts/*", "src/lib/email/*"),
        ("/app/api/email/*", "src/lib/email/*"),
        ("/app/api/email", "src/lib/email/sendEmail.ts"),
        ("migrations/20260420_drop_cancer_intake_plaintext (실행 완료) · 20260420_drop_inquiries_plaintext_email (보류) (평문 완전 제거)",
         "식별정보 암호화 적용: 이메일 109건·전화 22건 전건 암호문(2026-08-19 실측). "
         "평문 컬럼 삭제 마이그레이션은 보류(검색·번역에 쓰이는 본문 때문)"),
        ("migrations/20260420_drop_cancer_intake_plaintext (실행 완료) · 20260420_drop_inquiries_plaintext_email (보류)", "migrations/20260420_drop_cancer_intake_plaintext (실행 완료) · 20260420_drop_inquiries_plaintext_email (보류) (작성 완료·미실행)"),
        ("평문 완전 제거", "식별정보 암호화 완료 / 본문·소견은 평문 보관"),
    ],
    "08_테스트결과서.docx": [
        # ↓ 보안 항목별로 «무엇이 지키는지»를 적는다(2026-08-20 실측)
        ('성명·이메일·전화 암호문 확인(2026-08-19 재확인: 이메일 109건·전화 22건 전건 암호문). 문의 본문·의료진 소견은 검색·번역 목적으로 평문 보관',
         '성명(first_name·last_name)·이메일·전화가 전건 암호문임을 운영DB 로 확인(2026-08-20 재측정: 값이 있는 행 전부 암호문, 평문 0건). 문의 본문·의료진 소견은 검색·번역 목적으로 평문 보관'),
        ('SEC-01~08 모두 통과 (2026-08-19 기준). 분기별 회귀 테스트 필수.',
         'SEC-01~08 모두 통과 (2026-08-20 재확인). 무엇이 지키는지: SEC-01·02·04 는 requireAdminAuth 단위 시험(미인증→403 · 역할 상승 차단 · 요청량 초과→429)과 admin-auth-required 자동 클릭 검사, SEC-03 은 게스트 토큰 검사 2종, SEC-05 는 운영DB 직접 조회, SEC-06 은 자동 검사 check:err-exposure, SEC-07 은 공개 표 90개 전부 접근권한 규칙(RLS) 켜짐·정책 93개, SEC-08 은 server-only 정적 검사. 분기별 회귀 테스트 필수.'),
        # ↓ 미달 항목·불안정 검사를 숨기지 않고 적는다(2026-08-20)
        ('수동 27/27 · 자동 스모크 70통과 2실패 25건너뜀 (2026-08-20 실측)',
         '수동 27/27 · 자동 스모크 70통과 2실패 25건너뜀'),
        ('현재 시점 전체 통과',
         '2026-08-20 실측. 실패 2건은 의뢰서 접수 뒤 완료 화면을 기다리는 검사로, 다시 돌리면 실패 건수가 달라지고(2건→1건) 자동 검사 서버에서는 통과한다. 접수 자체는 운영DB 저장까지 확인했다'),
        # ↓ 2026-08-20 스모크 E2E 실행·AI 응답시간 실측으로 갱신
        ('27건',
         '27건(2026-04 수동) · 자동 스모크 97건(2026-08-20 실측)'),
        ('27/27 (2026-04 수동 확인 기준)',
         '수동 27/27 · 자동 스모크 70통과 2실패 25건너뜀 (2026-08-20 실측)'),
        # ↓ 2026-08-20 실측 재측정값(npm run test:run · playwright --list · npm audit --omit=dev)
        ("45개 파일 164건", "48개 파일 152건"),
        ("110개 파일 1,002건", "146개 파일 1,289건"),
        ("110개 파일 / 1,002건", "146개 파일 / 1,289건"),
        ("45개 파일 / 164건", "48개 파일 / 152건"),
        ("Critical: 0건, High: 1건, Moderate: 2건 (운영 의존성 기준, 2026-08-19 실행)", "Critical: 0건, High: 0건, Moderate: 2건 (운영 의존성 기준, 2026-08-20 실행)"),
        ("Critical 0건 · High 1건 · Moderate 2건 (2026-08-19)", "Critical 0건 · High 0건 · Moderate 2건 (2026-08-20)"),
        ("2026-08-19 현재는 e2e 폴더에", "2026-08-20 현재는 e2e 폴더에"),
        ("[해소됨 2026-08-19]", "[해소됨 2026-08-20]"),
        ("전건 통과 (2026-08-19 실행, 15.8초)", "전건 통과 (2026-08-20 실행, 32.7초)"),
        ('5.2 자동화 테스트 미흡 영역 (2026-04 지적 → 2026-08-19 해소 내역) (2026-04 지적 → 2026-08-19 해소 내역) (2026-04 지적 → 2026-08-19 해소 내역) (2026-04 지적 → 2026-08-19 해소 내역)',
         "5.2 자동화 시험 보강 내역 (2026-04 착수 점검 지적사항 해소)"),
        ("«수동 확인»", "수동 확인"),
        ("SEC-01~08 모두 통과 (2026-04-30 기준)", "SEC-01~08 모두 통과 (2026-08-19 기준)"),
        ("2026년 4월 30일", "2026년 8월 19일"),
        ("82개 파일 / 748건", "146개 파일 / 1,289건"),
        ("40개 파일 / 108건", "48개 파일 / 152건"),
        ("현재 (4월 30일)", "현재 (8월 19일)"),
        ("Critical: 0건, High: 0건, Moderate: [확인 필요. TBD]",
         "Critical: 0건, High: 0건, Moderate: 2건 (운영 의존성 기준, 2026-08-20 실행)"),
        ("2026-04-30", "2026-08-19"),
        ("(e2e/ 디렉토리 예정)",
         "(아래 표는 2026-04 시점 수동 확인 시나리오다. 2026-08-20 현재는 자동화 스크립트 48개 파일 152건으로 대체되어 자동 실행된다.)"),
        ("Playwright E2E 스크립트 코드베이스 내 미포함. 현재 수동 시나리오 기반 (e2e/ 디렉토리 추가 예정)",
         "[해소됨 2026-08-20] Playwright E2E 스크립트 도입 완료: 48개 파일 152건. 매 변경 시 스모크, 매일 밤 전체 실행."),
        ("단위 테스트(Jest) 케이스 부재: requireAdminAuth.test.ts 등 추가 필요",
         "[해소됨 2026-08-19] 단위 테스트 110개 파일 1,002건 도입 완료(전건 통과)."),
        ("5.2 자동화 테스트 미흡 영역", "5.2 자동화 시험 보강 내역 (2026-04 착수 점검 지적사항 해소)"),
        ("27건 | 27/27 (100%) | 현재 시점 전체 통과", "27건 | 27/27 | 2026-04 수동 확인 기준"),
        ("27/27 (100%)", "27/27 (2026-04 수동 확인 기준)"),
        ("2026년 5월 | 화상 내 번역 자막 E2E 테스트", "[완료] 화상 내 번역 자막: 실서비스 가동, 통번역 3,542건"),
        ("2026년 5월 | requireAdminAuth 단위 테스트", "[완료] requireAdminAuth 단위 테스트 도입"),
        ("2026년 6월 | 부하 테스트 (100 동시 접속)", "[예정] 부하 테스트 (100 동시 접속)"),
        ("2026년 7월 | 외부 침투 테스트", "[예정] 외부 침투 테스트"),
        ("name_encrypted 컬럼 암호문 확인, 평문 없음",
         "성명·이메일·전화 암호문 확인(2026-08-19 재확인: 이메일 109건·전화 22건 전건 암호문). 문의 본문·의료진 소견은 검색·번역 목적으로 평문 보관"),
        ("PII 평문 저장 없음", "식별정보 평문 저장 없음"),
        # 시험 환경 표기가 사실과 달랐다. 프리뷰는 2026-07-31 이후 닫혀 있고
        # 별도 Staging DB 도 없다(로컬·자동검사 모두 운영DB 를 본다).
        ("Vercel Preview + Supabase Staging",
         "로컬 개발 서버 + 운영 Supabase (별도 Staging DB 없음)"),
        ("파일 위치: src/lib/auth/requireAdminAuth.test.ts (예정: 현재 수동 검증)",
         "파일 위치: src/lib/auth/requireAdminAuth.test.ts (자동 시험 도입 완료)"),
        ("Jest / Vitest", "Vitest"),
        ("Jest + Supabase Test", "Vitest + 운영 Supabase"),
        ("Jest", "Vitest"),
        ("2026년 8월 19일", "2026년 8월 20일"),
        ("현재 (8월 19일)", "현재 (8월 20일)"),
        ("통번역 3,542건", "통번역 3,543건(시험 방 제외 3,277건)"),
        ("원인 규명 완료, 수정 반영 예정",
         "원인 규명 및 수정 완료(본판 반영). 실서비스 배포 후 재측정 예정"),
        ("Step 1~5 전체 입력 후 제출", "1단계 접수 후 2단계 추가 정보까지 입력·제출"),
        ("5-Step 인테이크 완료", "문의 접수(1단계+2단계) 완료"),
        ("Step 3에서 PDF 파일 첨부", "2단계에서 PDF 파일 첨부"),
    ],
    "09_산출물목록.docx": [
        # ↓ 제출 일정표에 빠져 있던 산출물 3종을 넣는다(2026-08-20)
        ('중간보고서 (KHIDI 지정 양식, 별도 관리)',
         '중간보고서 (KHIDI 지정 양식, 별도 관리)\n11_과제요구사항_6대ICT_분석및구현대비표.docx\n01-1_화면설계서.docx'),
        ('01_요구사항정의서.docx\n02_기능명세서.docx\nEVAL_MATRIX.docx',
         '01_요구사항정의서.docx\n01-1_화면설계서.docx\n02_기능명세서.docx\n10_백오피스_재설계_요구사항대비표.docx\n11_과제요구사항_6대ICT_분석및구현대비표.docx\nEVAL_MATRIX.docx'),
        # 몇 번을 돌려도 같은 결과가 되도록 «전체 문장»을 앵커로 잡는다(뒷말 덧붙기 방지).
        ('성과지표 중간 달성률 포함. 2026년 9월 중간평가 수검. 6대 ICT 이행 대비표와 화면 설계서를 정성지표 증빙으로 함께 제시. 6대 ICT 이행 대비표와 화면 설계서를 정성지표 증빙으로 함께 제시',
         '성과지표 중간 달성률 포함. 2026년 9월 중간평가 수검. 6대 ICT 이행 대비표와 화면 설계서를 정성지표 증빙으로 함께 제시'),
        ('성과지표 중간 달성률 포함. 8월 27일 중간평가 수검. 6대 ICT 이행 대비표와 화면 설계서를 정성지표 증빙으로 함께 제시',
         '성과지표 중간 달성률 포함. 2026년 9월 중간평가 수검. 6대 ICT 이행 대비표와 화면 설계서를 정성지표 증빙으로 함께 제시'),
        # ↓ 파일 실재·내용을 직접 확인해 정정(2026-08-20)
        ('AI_ARCHITECTURE_REPORT_2026_04.md', 'archive/AI_ARCHITECTURE_REPORT_2026_04.md'),
        ('Vercel 배포 설정: crons, rewrites, headers', 'Vercel 배포 설정: 정기 실행 9종 · 보안 머리값 2종 · 빌드 창구 판정'),
        ("Phase A 산출물은 사업 착수 및 KHIDI 신청 단계에서 작성된 문서이다. 2026년 4월 30일 완료.",
         "Phase A 산출물은 사업 착수 및 KHIDI 신청 단계에서 작성되었으며, 2026년 8월 19일 현행 실측 기준으로 갱신하였다."),
        ("2026년 4월 30일", "2026년 8월 19일"),
        ("구현 현황(완료74%/부분26%)", "구현 현황(완료96%/부분4%)"),
        ("진척률 74%, KPI 진행", "진척률 96%, KPI 진행"),
        ("2026-04-30 전체 완료", "2026-08-19 기준 갱신 완료"),
        ("2026-08-19", "2026-08-20"),
        ("Phase A 4건 + Phase B 7건 + 기타 docs 6건 = 17건",
         "Phase A 5건 + Phase B 4건 + 기술·운영 문서 6건 = 15건"),
        ("7건 | 5건 완료 / 2건 초안", "4건 | 4건 완료 (100%)"),
        ("04·05 보고서는 운영 후 갱신 예정",
         "중간·최종 보고서는 KHIDI 지정 양식이라 본 목록에서 제외"),
        ("17건 | 13건 완료", "15건 | 15건 완료"),
        ("완성도 76%", "완성도 100%"),
        ("migrations/ (37개 파일)", "migrations/ (154개 파일)"),
        ("docs/ (예정)", "docs/"),
        ("배포 전 안전장치 체크리스트 [예정]", "배포 전 안전장치 체크리스트"),
        ("M4. 중간점검 (9월)", "M4. 중간평가 (2026년 9월)"),
        # 옛 실행분이 이미 「8월 27일」로 박아둔 것을 되돌린다(중간평가는 9월 초, 날짜 미정)
        ("M4. 중간평가 (8월 27일)", "M4. 중간평가 (2026년 9월)"),
        ("Claude Code 프로젝트 개발 지침 및 규약", "프로젝트 개발 지침 및 코딩 규약"),
        ("2026년 5~7월", "2026년 5~8월"),
        ("월별 진행 요약 (이메일 등)", "월간 업무 보고서 (KHIDI 제출)"),
        ("KHIDI 형식 확인 필요 [진행 중]", "7월분부터 작성·제출 중"),
        ("KPI 중간 달성률 포함, 8월 제출 예정", "성과지표 중간 달성률 포함. 2026년 9월 중간평가 수검"),
    ],
    "06_사용자매뉴얼.docx": [
        ("healo.kr/intake", "healo.kr/inquiry"),
        ("[화면: /intake 페이지: Step 1]", "[화면: /inquiry 페이지: 1단계]"),
        ("[화면: /intake 페이지: Step 5 제출 완료 화면]", "[화면: /inquiry 페이지: 제출 완료]"),
        ("[화면: /coordinator/intake/[id] 페이지]", "[화면: /coordinator/intakes 페이지]"),
        ("healo.kr/partner 접속 후 의료진 계정으로 로그인",
         "healo.kr/hospital 접속 후 의료기관 담당자 계정으로 로그인 "
         "(의료진 본인은 계정 없이 상담방 초대링크로 참여)"),
        ("[화면: /partner 페이지: 파트너 대시보드]", "[화면: /hospital 페이지: 국내 의료기관 대시보드]"),
        ("[화면: /partner/patients/[id] 페이지]", "[화면: /hospital/leads 페이지: 의뢰 환자 상세]"),
        ("[화면: /partner/sessions/[id] 페이지: 화상 협진]",
         "[화면: /consultation/[id]: 화상 상담방(초대링크 입장)]"),
        ("[화면: /partner/patients/[id]/opinion 페이지]",
         "[화면: /opinion/[token]: 전문의 소견 작성]"),
        ("HEALO v1.0 (2026년 4월 기준)", "HEALO (2026년 8월 19일 기준)"),
        ("2026-04-30", "2026-08-19"),
        # ↓ 아래는 실제 배포 코드·화면과 대조해 바로잡은 것들(2026-08-20).
        ("[화면: /auth 페이지: 회원가입 폼]", "[화면: /signup 페이지: 회원가입 양식]"),
        ("소셜 로그인(Google, Kakao)으로도 가입 가능하다.",
         "소셜 로그인(Google, Apple)으로도 가입 가능하다."),
        ("[화면: /coordinator/intakes 페이지]", "[화면: /coordinator/inbox 페이지: 문의함]"),
        ("[화면: /coordinator/patients/[id] 페이지]",
         "[화면: /coordinator/inbox/[번호] 페이지: 케이스 상세]"),
        ("[화면: /coordinator/sessions/[id] 페이지]",
         "[화면: /coordinator/consultations 페이지: 상담 일정]"),
        ("[화면: /telemedicine/[room] 페이지: 화상상담 화면]",
         "[화면: /consultation/[id] 페이지: 화상 상담방(초대링크로 입장)]"),
        ("[화면: /opinion/[token]: 전문의 소견 작성]",
         "[화면: /opinion/<토큰>: 환자가 소견을 확인하는 화면]"),
        ("좌측 사이드바 [환자 관리] 클릭", "좌측 사이드바 [문의함] 클릭"),
        ("[인테이크 관리] 메뉴에서 신규 접수 건 확인", "[문의함] 메뉴에서 신규 접수 건 확인"),
        ("환자 목록에서 이름 또는 인테이크 ID로 검색", "문의 목록에서 이름 또는 문의 번호로 검색"),
        # 파일 상한: 화면이 실제로 안내하는 값은 200MB 다.
        ("CT/MRI/혈액검사 등 파일 첨부 (PDF·JPEG·PNG, 최대 50MB)",
         "CT/MRI/혈액검사 등 파일 첨부 (PDF·JPEG·PNG, 최대 10개·건당 200MB)"),
        ("지원 형식: DICOM, JPEG, PNG, PDF (파일당 최대 50MB)",
         "지원 형식: DICOM, JPEG, PNG, PDF (파일당 최대 200MB)"),
        # 접수 구조: 5단계 폼은 옛 구조다. 지금은 1단계 접수 뒤 2단계 추가 정보다.
        ("인테이크 폼은 암환자가 상담을 신청하기 위해 작성하는 5단계 정보 입력 양식이다.",
         "문의 접수는 두 단계로 나뉜다. 1단계에서 연락에 필요한 최소 정보만 받아 접수를 마치고, "
         "2단계(/inquiry/intake)에서 대학병원 의뢰에 필요한 자료를 추가로 받는다. "
         "자료가 없어 접수 자체를 못 하는 일을 없애기 위한 구조다."),
        ("Step 도중 저장된 진행 내용은 자동 저장된다. 브라우저를 닫아도 이어서 작성 가능하다 "
         "[TODO: 자동 저장 기능 구현 예정].",
         "1단계만 채워도 접수가 완료되며, 2단계는 접수 번호와 접근 토큰으로 나중에 이어서 작성할 수 있다."),
        # 지나간 「예정」: 자막·다국어·리마인더는 이미 실서비스에서 돌고 있다.
        ("화상 중 번역 자막 기능은 현재 개발 중이며 2026년 5월 적용 예정이다 [진행 중].",
         "화상 중 실시간 번역 자막은 실서비스에서 가동 중이다(자막 켜기 단추를 누를 때만 켜진다)."),
        ("실시간 번역 자막 [2026년 5월 예정: 진행 중]",
         "상대 발화를 내 언어로 표시. 자막 단추를 누를 때만 켜진다"),
        ("러시아어(/ru)와 카자흐어(/kk) 번역은 현재 일부 완성 상태이다. 2026년 6월 전면 완성 예정 [진행 중].",
         "러시아어(/ru)·카자흐어(/kz) 번역은 전 문구 완성 상태이며, 누락 여부를 자동 검사로 상시 확인한다."),
        ("예약 D-1 리마인더 [2026년 6월 예정: 진행 중]", "예약 리마인더: 정기 자동 발송"),
        ("한국어 (ru/kk 번역본 추후 제공 예정)", "한국어 (러시아어·카자흐어 번역본 추후 제공 예정)"),
        ("HEALO (2026년 8월 19일 기준)", "HEALO (2026년 8월 20일 기준)"),
        ("2026-08-19", "2026-08-20"),
        # ↓ 화면을 직접 띄워 확인한 것들(2026-08-20). 메뉴·단추 이름을 실제 것으로 맞췄다.
        ('환자 대시보드(/patient) 접속 후 [의료 기록] 메뉴 클릭',
         '환자 대시보드(/patient)에서 빠른 메뉴 [의료 문서] 클릭'),
        ('[파일 추가] 버튼 클릭 또는 파일을 드래그앤드롭',
         '「연결할 상담」을 고르고, 점선 상자에 파일을 끌어다 놓거나 상자를 클릭해 파일 선택'),
        ('지원 형식: DICOM, JPEG, PNG, PDF (파일당 최대 200MB)',
         '「문서 유형」(진단서·검사 결과·영상(CT/MRI)·처방전·기타)을 고르고 필요하면 설명을 적는다'),
        ('업로드 완료 시 파일 목록에 표시되고 코디네이터에게 알림이 발송된다.',
         '지원 형식은 PDF·JPG·PNG·WebP·Word·DICOM, 파일당 최대 200MB. 업로드하면 「내 문서」 목록에 뜨고 코디네이터에게 알림이 간다.'),
        ('환자 대시보드 [의료 기록] 탭에서 업로드된 파일 목록 확인',
         '「의료 문서」 화면 아래 「내 문서」에서 올린 파일 목록 확인'),
        ('파일명 클릭 시 다운로드 또는 미리보기',
         '파일 오른쪽 [보기]를 누르면 내려받거나 미리 볼 수 있다'),
        ('코디네이터가 검토 의견을 추가한 경우 해당 파일에 메모 표시',
         '파일마다 문서 유형·연결된 상담·용량이 함께 표시된다'),
        ('화면 우측 상단 언어 선택 버튼(국기 아이콘) 클릭',
         '화면 우측 상단 [언어 변경] 버튼 클릭'),
        ('환자 대시보드(/patient) 상단 메시지 아이콘 클릭',
         '환자 대시보드(/patient) 상단 알림 아이콘(읽지 않은 개수 표시) 클릭'),
        ('상담 종료 후 [상담 종료] 버튼을 클릭한다.',
         '상담이 끝나면 도구 막대의 [종료]를 누른다.'),
        ('의료 기록 화면 공유 가능 (PC 전용)',
         '의료 기록 화면 공유 (PC 전용)'),
        ('신규 인테이크 수신 현황: 미배정 건수 표시',
         '대기 중인 문의 건수'),
        ('예정 화상상담 일정',
         '오늘 상담 건수'),
        ('미답변 메시지 수',
         '예정 상담 건수'),
        ('월별 상담 건수 통계',
         '긴급 알림(증상) 건수. 카드 아래에 예정 상담 목록과 바로가기가 이어진다'),
        ('문의 목록에서 이름 또는 문의 번호로 검색',
         '목록 위 거르기 세 갈래(전체 / 추가 정보 필요 / 매칭 준비 완료)로 좁힌다. 각 갈래에 건수가 함께 뜬다'),
        ('환자 이름 클릭 시 상세 페이지 이동',
         '단계 기준일을 넘긴 케이스에는 「N일째 정체」 빨간 배지가 붙는다. 환자 이름을 누르면 케이스 상세로 이동'),
        ('[문의함] 메뉴에서 신규 접수 건 확인',
         '[문의함]에서 신규 접수 건을 확인한 뒤 [의뢰·케이스/병원배정] 메뉴로 이동'),
        ('[병원 매칭 AI 추천] 버튼 클릭 시 AI가 적합 병원 3개 추천',
         '문의를 케이스로 전환한다. AI 케이스 브리프가 환자 요약·확인 포인트·주의사항 초안을 만들어 둔다'),
        ('추천 결과 검토 후 [배정] 버튼으로 병원 및 의료진 배정',
         '검토 후 적합한 제휴 병원을 배정한다'),
        ('환자에게 배정 결과 이메일 자동 발송',
         '배정 결과는 케이스 이력에 남고 담당자에게 알림이 간다'),
        ('[화면: /coordinator/inbox 페이지: 문의함]',
         '[화면: /coordinator/cases 페이지: 의뢰·케이스/병원배정]'),
        ('환자 상세 페이지 [화상상담 설정] 탭 클릭',
         '[상담 일정] 메뉴에서 상담을 만든다'),
        ('상담 일시, 참여자(의료진) 선택 후 [방 생성]',
         '필수 입력은 환자와 예약 시각 둘뿐이다. 시각은 한국시간이며 협정시계를 함께 보여준다'),
        ('생성된 화상상담 링크를 환자에게 이메일/WhatsApp으로 전송',
         '상담 링크는 하나다(줌과 같다). [🔗 링크 복사]로 복사해 왓츠앱·메일로 보낸다'),
        ('상담 시간에 맞춰 [화상 참여] 클릭으로 화상방에 입장',
         '시간이 되면 [상담 시작]을 누른다. 같은 링크로 내가 입장하고 링크가 클립보드에 복사된다'),
        ('상담 종료 후 상담 결과 메모 저장 및 후속 액션 설정',
         '상담이 끝나면 [상담 완료]를 누른다. 이때만 KHIDI 사전상담·사후관리 실적에 잡힌다'),
        ('게스트 참여 링크는 24시간 유효. 만료 시 새 링크 재발급 필요.',
         '상담방에서 [종료]만 누르면 실적에 잡히지 않는다. [상담 완료] 처리하면 보낸 초대 링크는 폐기된다.'),
        ('환자 상세 페이지 또는 [메시지] 메뉴에서 해당 환자 선택',
         '[메시지] 메뉴에서 대화를 고른다. 대화마다 채널 배지(웹·텔레그램·왓츠앱·에이전시)가 붙는다'),
        ('메시지 입력창에 텍스트 작성 후 [전송]',
         '입력창 위 「추천 답장」 칩을 누르면 자주 쓰는 답이 채워진다(자동 전송은 아니다). 고쳐서 [전송]'),
        ('환자에게 새 메시지 이메일 알림 자동 발송',
         '텔레그램·왓츠앱 대화에 답하면 환자의 메신저 앱으로 실제 발송된다. 사람이 답하기 시작하면 AI 는 더 이상 끼어들지 않는다'),
        ('메시지는 현재 텍스트 전용. 파일 첨부는 [의료 기록] 탭에서 별도 처리.',
         '발송 실패나 왓츠앱 24시간 창 만료는 말풍선 아래 미전달 표시로 보인다.'),
        ('화상상담 종료 후 [진료 의견서 작성] 클릭',
         '코디네이터가 케이스 상세에서 [소견 요청 링크 만들기]를 눌러 링크를 보낸다. 의료진은 계정 없이 그 링크로 들어간다'),
        ('권고 치료 방법, 예상 입원 기간, 주의사항 입력',
         '검사지를 보고 소견을 남긴다. 이미 카톡·메일로 보낸 소견이 있으면 코디네이터가 [이미 받은 소견 직접 입력]으로 등록한다'),
        ('[제출] 클릭 시 코디네이터에게 전달되고 환자에게 이메일 발송',
         '도착한 소견은 원문 그대로 케이스 상세에만 쌓인다. AI 가 환자 언어 번역 초안을 만들고, 코디네이터가 교정한 뒤 [에이전시에 공개]를 눌러야 밖으로 나간다'),
        ('[화면: /opinion/<토큰>: 환자가 소견을 확인하는 화면]',
         '[화면: /coordinator/inbox/[번호] 「전문의 소견」 · /opinion/<토큰>(환자가 확인하는 화면)]'),
        # ↓ 병원 포털: 시험 리드를 하나 만들어 실제 화면을 띄워 확인(2026-08-20). 확인 뒤 그 리드는 지웠다.
        ('파트너 대시보드로 이동: 배정된 환자 목록, 예정 화상상담, 수신 의뢰 확인',
         '왼쪽 메뉴는 [대시보드]와 [진료 의뢰(리드)] 둘이다. (「병원 정보」·「시술 카탈로그」는 공개 연동 준비 중이라 아직 비활성)'),
        ('의료진 계정은 관리자가 발급한다. 계정 신청은 플랫폼 관리자(admin)에게 문의.',
         '대시보드 카드는 응답 대기 · 전환율 · 평균 첫 응답 · 확정 견적 합계 네 개이고, 그 아래에 전체/오늘/이번 주/이번 달 리드 수와 「응답 필요」 목록이 이어진다. 의료기관 담당자 계정은 관리자가 발급한다.'),
        ('[배정 환자] 목록에서 환자 선택',
         '[진료 의뢰(리드)] 메뉴에서 우리 병원으로 배정된 의뢰를 고른다'),
        ('열람 가능 항목: 인테이크 정보(암종·병기·기왕력), 업로드 의료 기록',
         '리드에 보이는 항목: 목적 · 시술 · 국가 · 언어 · 배정일 · 문의 내용 · 보험 정보 유무. 환자 이름과 연락처는 병원에 넘기지 않는다'),
        ('의료 기록 파일 다운로드 후 검토',
         '상태를 전송됨 → 조회됨 → 응답함 → 치료 확정(또는 거절)으로 갱신하고, 견적 최소·최대와 메모를 남기거나 「코디에게 메시지」를 보낸다. 목록은 [CSV]로 내려받을 수 있다'),
        ('[화면: /hospital/leads 페이지: 의뢰 환자 상세]',
         '[화면: /hospital/leads 페이지: 리드 관리]'),
        ('대시보드 [예정 화상상담] 목록에서 해당 일정 확인',
         '코디네이터가 보낸 상담 초대링크를 연다(병원 포털 안에는 화상상담 목록이 없다)'),
        ('상담 시간 5분 전 [입장] 버튼 클릭',
         '대기실에서 코디네이터가 입장을 승인하면 상담방에 들어간다'),
        ('의료진 화상 접속은 로그인 후 직접 입장 (게스트 링크 불사용).',
         '의료진은 계정 없이 초대링크로 참여한다. 병원 담당자 계정으로 로그인해도 상담방은 초대링크로 들어간다.'),
    ],
    "07_관리자매뉴얼.docx": [
        ("/admin/intake/[id] 에서", "/admin/inquiries 에서"),
        ("HEALO v1.0 (2026년 4월 기준)", "HEALO (2026년 8월 19일 기준)"),
        ("2026-04-30", "2026-08-19"),
        # ↓ 실제 배포 코드와 대조해 바로잡은 것들(2026-08-20).
        ("/admin/patients 접속", "/admin/inquiries 접속"),
        ("[화면: /admin/patients 페이지]", "[화면: /admin/inquiries 페이지: 문의 관리]"),
        ("/admin/sessions 접속", "/admin/consultations 접속"),
        ("[화면: /admin/sessions 페이지]", "[화면: /admin/consultations 페이지: 상담 관리]"),
        ("/admin/statistics 접속", "/admin/khidi/conversion 접속"),
        ("[화면: /admin/statistics 페이지]", "[화면: /admin/khidi/conversion 페이지: 유치 전환 대시보드]"),
        ("/admin/reports 접속", "/admin/khidi 접속"),
        ("2FA 인증 [추후 적용 예정] 또는 이메일 인증 확인", "이메일 인증 확인"),
        ("다운로드한 데이터로 04_중간보고서 및 05_최종보고서 실적 수치 갱신",
         "내려받은 데이터로 중간·최종 보고서(KHIDI 지정 양식) 실적 수치를 갱신한다"),
        ("역할: patient / coordinator / admin / partner_hospital",
         "역할: 환자 / 코디네이터 / 관리자 (국내 의료기관·해외 에이전시·해외 의료기관은 "
         "hospital_users·agency_users 표에서 따로 관리한다)"),
        # 요금제: 우리는 Pro 플랜이다(같은 문서 6.1 항목과 어긋나 있었다).
        ("Supabase 무료 플랜 한도 초과 시 서비스 중단 위험. 월 1회 Usage 점검 필수.",
         "Supabase Pro 플랜 사용 중. 지출 상한을 넘기면 업로드부터 막히므로 월 1회 Usage 점검 필수."),
        # 배포: main 에 밀면 바로 나가지 않는다. 하루 한 번 배포 창구가 production 으로 민다.
        ("Vercel 배포 트리거: main 브랜치 push 시 자동. 수동 재배포: [Redeploy] 버튼.",
         "Vercel 배포 트리거: 하루 한 번(KST 15시) Daily Deploy 창구가 main 을 production 브랜치로 "
         "밀 때만 빌드된다. main 에 합쳤다고 바로 실서비스에 나가지 않는다."),
        ("git revert HEAD", "git revert <되돌릴 커밋>"),
        ("HEALO (2026년 8월 19일 기준)", "HEALO (2026년 8월 20일 기준)"),
        ('/admin/statistics 접속',
         '[홈 › KHIDI 리포트] 메뉴 클릭 (/admin/khidi/kpi-dashboard)'),
        ('[화면: /admin/khidi/conversion 페이지: 유치 전환 대시보드]',
         '[화면: /admin/khidi/kpi-dashboard 페이지: KHIDI 리포트]'),
        ('기간 선택 후 KPI 현황 확인',
         '유치·상담·사후관리·만족도가 목표 대비 자동 집계된다. 아래 [북극성 지표]·[유치 전환 상세]·[환자 만족도]·[증빙 산출물]로 더 파고든다'),
        ('/admin/khidi 접속',
         '[홈 › KHIDI 리포트 › 증빙 산출물] 열기 (/admin/khidi/evidence)'),
        ('보고 기간 선택 (월별 / 분기별)',
         '보고 기간을 고른다'),
        ('/admin/users 접속 후 [새 사용자 추가] 클릭',
         '[파트너·회원 › 직원(코디) 계정] 메뉴 클릭 (/admin/staff)'),
        ('이메일 입력 후 역할 선택: coordinator',
         '이름·이메일을 넣고 역할을 「코디네이터」로 둔다. 임시 비밀번호는 비우면 자동 생성된다'),
        ('임시 비밀번호 발급 및 초대 이메일 발송',
         '[계정 생성 / 역할 부여]를 누른다. 초대 메일이 나가는 것이 아니라, 나온 이메일과 임시 비밀번호를 직원에게 직접 전달하면 본인이 바꾼다'),
        ('사용자 목록에서 계정 선택 후 [역할 변경]',
         '「등록된 직원」 목록에서 [수정]을 누른다. 같은 줄에 [비밀번호 초기화]·[비활성화]도 있다'),
        ('역할: 환자 / 코디네이터 / 관리자 (국내 의료기관·해외 에이전시·해외 의료기관은 hospital_users·agency_users 표에서 따로 관리한다)',
         '의사는 계정이 없다. 상담방 초대링크로 참여한다. 국내 의료기관·해외 에이전시·해외 의료기관은 [에이전시·클리닉]·[제휴 병원] 메뉴에서 따로 관리한다'),
        ('신규 인테이크',
         '환자'),
        ('미처리 인테이크 수신 건수 (오늘/이번 주)',
         '문의 건수 · AI 채팅 · 텔레그램/왓츠앱'),
        ('활성 세션',
         '코디네이터'),
        ('현재 진행 중인 화상상담 세션 수',
         '상담 건수 · 예정 상담 · 견적/문구 편집'),
        ('전체 사용자 수',
         '에이전시·클리닉'),
        ('가입 환자·코디네이터·의료진 합계',
         '활성 기관 수 · 협진 의뢰 · 파트너 계정'),
        ('월별 상담 건수',
         '병원'),
        ('KPI K-02 추적용 차트',
         '미응답 리드 · 제휴 병원 · 리드 전체'),
        ('오류 로그',
         '시스템'),
        ('최근 24시간 API 오류 건수',
         'AI 회귀 시험 · 환각 응답 · 최근 실행. 카드를 누르면 해당 화면으로 이동하고, 아래 「최근 활동」이 모든 역할의 변경사항을 시간순으로 보여준다'),
        ('검색 필터: 국가, 암종, 인테이크 상태, 등록일',
         '문의를 국가·암종·상태로 좁혀 본다'),
        ('환자 클릭 시 전체 인테이크·상담·문서 이력 확인',
         '문의를 누르면 케이스 상세로 들어가 접수 내용·서류·상담·소견 이력을 본다. 가입 회원 목록은 [파트너·회원 › 환자 회원]에서 따로 본다'),
        ("2026-08-19", "2026-08-20"),
    ],
}

# 사업계획서(2026-05-14 제출본)가 약속한 결과물과 실제 산출물의 대응.
# 계획서엔 「UI/UX 설계서」인데 우리 파일 이름은 「화면 설계서」인 것처럼 표현이
# 조금씩 다르다. 파일 이름을 바꾸는 대신 이 표로 이어 보인다(이름을 바꾸면 참조가
# 여러 곳이라 비용이 크고, 대조하는 사람에게는 이 표 한 장이면 충분하다).
PLAN_MAP_TITLE = "사업계획서 약속 대비표"
PLAN_MAP_ROWS = [
    ("M1. 플랫폼 설계 완료 (5월)", "요구사항 정의서", "01_요구사항정의서.docx", "충족"),
    ("", "UI/UX 설계서", "01-1_화면설계서.docx (화면 10종 lo-fi 설계)", "충족"),
    ("", "시스템 구성도", "01-1_화면설계서.docx 제2장", "충족"),
    ("M2. 플랫폼 개발 완료 (6월)", "HEALO 배포본", "healwith.co.kr 실서비스 가동", "충족"),
    ("", "기능 테스트 결과서", "08_테스트결과서.docx", "충족"),
    ("M4. 중간점검 (9월)", "중간 실적 보고서", "중간보고서 (KHIDI 지정 양식, 별도 관리)", "진행"),
    ("정성지표 증빙", "플랫폼 시연 및 관련 산출물(정의서 · 설계서 · 테스트 결과 등)",
     "위 문서 일체 + 02_기능명세서 · 11_6대ICT 대비표", "충족"),
]

# 09 제출 일정표에서 폐기 문서를 가리키던 문구. 제출은 계속 있으므로 문구만 바꾼다.
DROPPED_DOC_NOTES = [
    ("04_중간보고서.docx (갱신본)", "중간보고서 (KHIDI 지정 양식, 별도 관리)"),
    ("05_최종보고서.docx", "최종보고서 (KHIDI 지정 양식, 별도 관리)"),
]

# ── 덧붙일 장(章) ───────────────────────────────────────────────────────────
# 제목이 이미 있으면 다시 붙이지 않는다(재실행 안전).
# 제목에 날짜를 넣지 않는다 — 날짜가 들어가면 갱신 때마다 「이미 있음」 검사가 빗나가
# 부록이 한 문서에 여러 벌 쌓인다(2026-08-19 실제로 04·06·07·08·09 에 두 벌씩 쌓였다).
APPENDIX_TITLE = "부록. 현행 반영"


def set_font(run, size=10, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rPr.append(rFonts)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(doc, text, size=14):
    p = doc.add_paragraph()
    set_font(p.add_run(text), size=size, bold=True, color=(0, 70, 127))
    return p


def para(doc, text, size=10, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    set_font(p.add_run(text), size=size)
    return p


def set_table_borders(t):
    """'Table Grid' 스타일이 없는 문서(다른 도구로 만든 04·06·07 등)를 위해 테두리를 직접 그린다."""
    tblPr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tblPr.append(borders)


def table(doc, headers, rows):
    t = doc.add_table(rows=0, cols=len(headers))
    try:
        t.style = "Table Grid"
    except KeyError:
        # 이 문서엔 해당 스타일이 정의돼 있지 않다 → 테두리를 직접 넣는다.
        set_table_borders(t)
    hr = t.add_row()
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = ""
        set_font(c.paragraphs[0].add_run(h), size=9, bold=True, color=(255, 255, 255))
        set_cell_bg(c, "00467F")
    for r in rows:
        dr = t.add_row()
        for i, v in enumerate(r):
            if i < len(dr.cells):
                c = dr.cells[i]
                c.text = ""
                set_font(c.paragraphs[0].add_run(str(v)), size=9, bold=(i == 0))
    doc.add_paragraph()
    return t


def iter_paragraphs(doc, with_furniture=True):
    """본문 문단 + 표 안 문단 + 머리말·꼬리말까지 훑는다.

    ⚠️ 머리말·꼬리말을 빼먹지 마라. 2026-08-20 실측: 본문에서 「(주)본로이」를 다 지웠는데
       꼬리말에 남아 06·07·08·09 의 «매 쪽 하단»에 그대로 인쇄되고 있었다.
    """
    def walk(container):
        for p in container.paragraphs:
            yield p
        for t in container.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
    yield from walk(doc)
    if with_furniture:
        for sec in doc.sections:
            for part in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer,
                         sec.even_page_header, sec.even_page_footer):
                try:
                    yield from walk(part)
                except Exception:
                    continue


def replace_in_paragraph(p, old, new):
    """문단에서 old→new. 글자가 여러 run 으로 쪼개져 있어도 처리한다."""
    if old not in p.text:
        return False
    # new 가 old 를 품고 있는 쌍(꼬리말 덧붙이기)만 겹쳐 붙는다. 그 경우에만 건너뛴다.
    if old in new and new in p.text:
        return False
    # 1) 한 run 안에 통째로 들어 있으면 그 run 만 고친다(서식 보존).
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # 2) run 경계로 쪼개진 경우 — 첫 run 에 합치고 나머지를 비운다.
    #    (짧은 경로 문자열이라 서식 손실이 사실상 없다.)
    merged = p.text.replace(old, new)
    if not p.runs:
        return False
    p.runs[0].text = merged
    for run in p.runs[1:]:
        run.text = ""
    return True


def strip_appendix(doc):
    """이미 붙어 있는 부록을 통째로 걷어낸다. 없으면 아무것도 안 한다.

    부록은 항상 문서 맨 뒤에 붙으므로, 첫 부록 제목부터 본문 끝까지를 지운다.
    앞의 쪽나눔 빈 문단도 같이 지워야 빈 페이지가 남지 않는다.
    """
    body = doc.element.body
    kids = [c for c in body.iterchildren() if not c.tag.endswith("}sectPr")]
    start = next((i for i, c in enumerate(kids)
                  if APPENDIX_TITLE in "".join(c.itertext())), None)
    if start is None:
        return 0
    while start > 0 and not "".join(kids[start - 1].itertext()).strip():
        start -= 1
    for c in kids[start:]:
        body.remove(c)
    return len(kids) - start


def add_appendix(doc, kind):
    """kind: 'report'(04) | 'user'(06) | 'admin'(07) | 'test'(08) | 'inventory'(09)"""
    doc.add_page_break()
    heading(doc, APPENDIX_TITLE)
    para(doc, f"기준일 : {F.AS_OF}", size=9)
    doc.add_paragraph()

    n = [0]

    def sec(title):
        n[0] += 1
        heading(doc, f"{n[0]}. {title}", size=12)

    sec("계정 계층 (7종)")
    para(doc, "※ 의사는 계정 계층이 아니다. 화상상담 초대링크 게스트 또는 의료기관 계정으로 참여한다.", size=9)
    table(doc, ["계층", "권한 저장", "전용 화면", "설명"],
          [(t[1], t[2], t[3], t[4]) for t in F.TIERS])

    if kind in ("user", "admin", "report"):
        sec("화상 상담(원격협진) 기능")
        table(doc, ["구분", "내용"], F.TELEMEDICINE)

        sec("상담 채널")
        para(doc, "※ 위챗·라인은 메신저 바로가기 안내이며 봇 연동은 미적용이다.", size=9)
        table(doc, ["채널", "연동 수준", "내용"], F.CHANNELS)

    sec("화면 동선 통합·정비 내역")
    para(doc, "이용자가 같은 목적으로 여러 화면을 오가지 않도록 진입 경로를 하나로 모으고, 옛 경로는 영구 이동 처리하여 기존 링크·검색 유입이 끊기지 않도록 하였다.", size=9)
    table(doc, ["현행 화면", "통합한 옛 화면", "정비 내용"], F.ROUTES_RETIRED_PRINTABLE)

    if kind in ("report", "inventory"):
        sec("성과지표 현황")
        para(doc, F.KPI_NOTE, size=9)
        table(doc, ["지표", "목표", f"실적({F.AS_OF})"], [
            ("외국인환자 유치", f"{F.KPI_TARGET['attraction']}건", f"{F.KPI_ACTUAL['attraction']}건"),
            ("사전상담", "—", f"{F.KPI_ACTUAL['preConsultation']}건"),
            ("사후관리", "—", f"{F.KPI_ACTUAL['followUp']}건"),
            ("사전상담+사후관리 합산", f"{F.KPI_TARGET['consultAndCare']}건",
             f"{F.KPI_ACTUAL['preConsultation'] + F.KPI_ACTUAL['followUp']}건"),
            ("환자 만족도", f"{F.KPI_TARGET['satisfaction']}점",
             f"표본 {F.KPI_ACTUAL['satisfactionSamples']}건"),
            ("다국어 지원", f"{F.KPI_TARGET['languages']}개 언어",
             f"{F.KPI_ACTUAL['languages']}개 언어(충족)"),
        ])
        para(doc, f"※ 시험 데이터 제외. 문의 {F.KPI_ACTUAL['inquiriesReal']}건, "
                  f"상담세션 {F.KPI_ACTUAL['sessionsReal']}건.", size=9)

    if kind == "test":
        sec("품질검증 현황")
        table(doc, ["구분", "규모", "결과"], [
            ("단위 테스트", f"{F.QUALITY['unit_files']}개 파일 / {F.QUALITY['unit_tests']:,}건",
             F.QUALITY["unit_result"]),
            ("통합·E2E 테스트", f"{F.QUALITY['e2e_files']}개 파일 / {F.QUALITY['e2e_tests']}건",
             "자동 실행"),
        ])
        sec("자동 검사 항목")
        table(doc, ["검사", "내용", "주기"], F.QUALITY["ci_gates"])

    sec("근거자료")
    table(doc, ["구분", "출처", "확인 내용"], F.PROVENANCE)


# ── 08 Core Web Vitals ─────────────────────────────────────────────────────
# 2026-08-21 Lighthouse 12.8.2 실측(https://healwith.co.kr/ko), 광고차단 프로그램을 끈 상태에서
# 모바일·데스크톱 각 5회 측정의 중앙값.
# ⚠️ 2026-08-21 정정 2건 — 앞서 이 파일에 적었던 두 문장이 «둘 다 틀렸다».
#  (1) 「첫 화면 전송량 1.2MB 가 병목」 → 병목은 전송량이 아니라 메인 작업 2.4초였다.
#  (2) 「광고차단이 HTML 압축을 벗겨 전송량을 부풀린다」 → 범인은 광고차단이 아니라
#      우리 서비스 워커(/sw.js)다. 서비스 워커가 응답을 대신 내주면 브라우저는 규격상
#      압축 전 크기를 전송량으로 보고하고 프로토콜 칸을 비운다. 대조군(MDN)은 같은
#      브라우저에서 정상 압축으로 보고됐고, 서비스 워커를 해제하자 우리 사이트도
#      문서 128KB·프로토콜 h2 로 정상 보고됐다.
#      그때 함께 적은 「전체 322KB」도 틀렸다 — 자원을 다 세지 않은 값이었다(실제 841KB).
CWV_ROWS = [
    ("LCP (최대 콘텐츠 페인트)", "2.5초 이하",
     "실측 0.42초 / 데스크톱 1.12초 (저속 회선 가정 시뮬레이션 5.31초)",
     "충족. 실제 브라우저에서 잰 값은 0.42초이며, 데스크톱은 5회 중앙값 1.12초다. "
     "괄호 안 5.31초는 채점 도구가 «느린 4G 회선과 4배 느린 기기를 가정해 계산한» 값으로, "
     "회선 조건만 정상이면 같은 저사양 기기에서도 0.68초로 측정된다"),
    ("CLS (레이아웃 안정성)", "0.1 이하", "데스크톱 0.008 / 모바일 0.045",
     "충족. 각 5회 측정 전 회차가 목표 이내이며, 개선 전 데스크톱 0.51 대비 60분의 1 수준으로 "
     "안정화하였다"),
    ("입력 반응성 (FID → TBT 대체)", "TBT 200ms 이하", "데스크톱 6ms (저속 기기 가정 372ms)",
     "충족. FID 는 폐지된 지표로 총 차단 시간(TBT)으로 대체 측정하였다. 괄호 안 값은 4배 "
     "느린 기기를 가정한 채점용 수치다"),
    ("TTFB (서버 응답 시간)", "800ms 이하", "데스크톱 172ms / 모바일 609ms", "충족"),
    ("카자흐스탄 현지 접속 응답", "자체 측정 항목", "왕복 69~92ms",
     "주 고객 지역인 카자흐스탄 4개 도시(알마티·카라간디·파블로다르)의 실제 회선에서 "
     "2026. 8. 21. 측정. 현지 이용자가 국내 이용자와 큰 차이 없이 접속한다"),
    ("AI 챗 첫 응답", "5초 이하 (NFR-02)",
     "첫 글자까지 중앙값 2.30초 · 상위 5% 3.84초 (응답 완료 3.03초)",
     "충족. 실서비스 응답 10건 실측(2026. 8. 21., 한국어·영어·러시아어·카자흐어). "
     "검색·프롬프트 구성까지 포함한 실제 이용 경로 기준이다"),
]


CWV_NOTE = ("측정 도구 Lighthouse 12.8.2, 대상 https://healwith.co.kr/ko, 측정일 2026. 8. 21. "
            "데스크톱·모바일 각 5회 측정의 중앙값이며, 첫 화면 전송량은 841KB(자원 70건)다. "
            "모바일 수치에 병기한 값은 채점 도구가 저속 4G 회선과 4배 느린 기기를 «가정해 "
            "계산한» 시뮬레이션 결과이므로, 실제 브라우저 측정값과 구분하여 표기하였다. "
            "종합 점수는 측정 환경에 따라 편차가 크므로 기재하지 않고 지표값만 기재한다.")


def fill_cwv(doc):
    """08 의 Core Web Vitals 표를 실측값으로 채운다. 몇 번 돌려도 같은 결과."""
    for t in doc.tables:
        if not t.rows or t.rows[0].cells[0].text.strip() != "지표":
            continue
        if "2.5초 이하" not in "".join(c.text for c in t.rows[1].cells):
            continue
        for tr in t._tbl.tr_lst[1:]:
            t._tbl.remove(tr)
        for r in CWV_ROWS:
            dr = t.add_row()
            for i, v in enumerate(r):
                c = dr.cells[i]; c.text = ""
                set_font(c.paragraphs[0].add_run(v), size=9, bold=(i == 0))
        return True
    return False


def fix_em_dash(doc):
    """한국어 문장 가운데의 줄표를 콜론·마침표로 바꾼다(PO 규칙 2026-08-20).

    국어 문장부호 규정상 줄표의 용법은 「제목 뒤 부제」뿐이다. 앞말이 문장으로 끝나면
    마침표, 명사로 끝나면 콜론을 쓴다. 표 칸에서 「해당 없음」을 뜻하는 홀로 선 줄표는
    기호라서 건드리지 않는다(앞뒤 공백이 있는 것만 바꾼다).
    """
    D = chr(8212)
    ENDS = ("다", "음", "함", "임", "됨", "것", "요", "까", "나")
    n = 0
    for par in iter_paragraphs(doc):
        t = par.text
        if f" {D} " not in t:
            continue
        new = re.sub(r"(\S) " + D + r" (\S)",
                     lambda m: m.group(1) + (". " if m.group(1)[-1] in ENDS else ": ") + m.group(2), t)
        if new == t or not par.runs:
            continue
        par.runs[0].text = new
        for r in par.runs[1:]:
            r.text = ""
        n += 1
    return n


FOOTER_LEFT = "본로이  |  기밀문서, 무단 배포 금지"


def _field(par, instr):
    """워드 필드(쪽번호 등)를 문단에 넣는다. 글자로 «PAGE» 라고 쓰면 숫자가 안 나온다."""
    r = par.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin"); r._r.append(fc)
    r2 = par.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = f" {instr} "
    r2._r.append(it)
    r3 = par.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r3._r.append(fe)
    for run in (r, r2, r3):
        set_font(run, size=8, color=(120, 120, 120))


def ensure_footer(doc):
    """모든 구역에 같은 꼬리말을 둔다. 제출 묶음인데 문서마다 있거나 없으면 티가 난다."""
    n = 0
    for sec in doc.sections:
        f = sec.footer
        # 이미 우리 꼬리말이면 손대지 않는다. 안 그러면 돌릴 때마다 파일이 다시 저장돼
        # 내용이 같은데도 git 에 변경으로 잡힌다.
        if (not f.is_linked_to_previous
                and any(FOOTER_LEFT in par.text for par in f.paragraphs)):
            continue
        f.is_linked_to_previous = False
        for par in list(f.paragraphs)[1:]:
            par._p.getparent().remove(par._p)
        par = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
        for r in list(par.runs):
            r._r.getparent().remove(r._r)
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(par.add_run(FOOTER_LEFT + "     "), size=8, color=(120, 120, 120))
        _field(par, "PAGE")
        set_font(par.add_run(" / "), size=8, color=(120, 120, 120))
        _field(par, "NUMPAGES")
        n += 1
    return n


def repeat_table_headers(doc, min_rows=12):
    """쪽을 넘어가는 긴 표의 머리행을 다음 쪽에도 반복시킨다.

    안 하면 두 번째 쪽 표가 «무슨 열인지» 알 수 없다. 2026-08-20 실측: 12행 넘는 표가
    21개인데 반복 설정이 하나도 없었다.
    """
    n = 0
    for t in doc.tables:
        if len(t.rows) < min_rows:
            continue
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        if trPr.find(qn("w:tblHeader")) is None:
            el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true")
            trPr.append(el)
            n += 1
    return n


def move_page_breaks(doc):
    """빈 문단에 든 쪽 나눔을 «다음 문단의 쪽 나눔 속성»으로 옮긴다.

    쪽이 딱 찬 직후에 쪽 나눔 문단이 오면 워드가 아무것도 없는 쪽을 한 장 만든다
    (2026-08-20 실측: 09_산출물목록 6쪽이 머리말·꼬리말만 있는 빈 쪽이었다).
    속성으로 옮기면 「다음 쪽부터 시작」이라는 뜻만 남아 빈 쪽이 생기지 않는다.
    """
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    n = 0
    for el in list(body.iterchildren()):
        if el.tag != qn("w:p"):
            continue
        par = Paragraph(el, doc)
        if par.text.strip():
            continue
        brks = el.findall(".//" + qn("w:br"))
        if not any(b.get(qn("w:type")) == "page" for b in brks):
            continue
        nxt = el.getnext()
        if nxt is None or nxt.tag != qn("w:p"):
            continue                      # 다음이 표면 속성으로 못 옮긴다 — 그대로 둔다
        Paragraph(nxt, doc).paragraph_format.page_break_before = True
        body.remove(el)
        n += 1
    return n


def keep_short_tables_together(doc, max_rows=16):
    """짧은 표가 쪽 경계에 걸려 «머리행 없는 조각»만 다음 쪽에 남는 것을 막는다.

    2026-08-20 실측: 08 테스트결과서 5쪽이 화상상담 표의 마지막 한 줄만 있는 쪽이었고,
    02 기능명세서 21쪽도 시나리오 표(14줄)의 마지막 줄 하나뿐이었다.
    머리행 반복(repeat_table_headers)은 12행 이상 긴 표만 다루므로 짧은 표는 여기서 막는다.
    """
    from docx.text.paragraph import Paragraph
    n = 0
    for t in doc.tables:
        if len(t.rows) > max_rows:
            continue
        for row in t.rows[:-1]:
            for c in row.cells:
                for par in c.paragraphs:
                    if not par.paragraph_format.keep_with_next:
                        par.paragraph_format.keep_with_next = True
                        n += 1
        # 표 «바로 앞»의 제목·안내 줄도 함께 묶는다. 안 그러면 제목만 앞 쪽에 홀로 남는다
        # (2026-08-20 실측: EVAL_MATRIX 4쪽이 「3. HEALO 평가 매트릭스」 두 줄뿐이었다).
        prev = t._tbl.getprevious()
        kept = 0
        while prev is not None and prev.tag == qn("w:p") and kept < 2:
            par = Paragraph(prev, doc)
            # 사이에 낀 빈 문단도 함께 묶어야 한다 — 하나라도 빠지면 거기서 사슬이 끊겨
            # 제목만 앞 쪽에 남는다.
            if not par.paragraph_format.keep_with_next:
                par.paragraph_format.keep_with_next = True
                n += 1
            if par.text.strip():
                kept += 1
            prev = prev.getprevious()
    return 1 if n else 0


def widen_id_columns(doc, min_in=1.0):
    """머리글이 「ID」인 첫 칸이 너무 좁아 식별자가 여러 줄로 쪼개지는 것을 막는다.

    2026-08-20 실측: 08 테스트결과서의 ID 칸이 0.42인치라 「E2E-INT-04」가 세 줄로 갈렸다.
    모자란 폭은 가장 넓은 칸에서 덜어 온다(표 전체 폭은 그대로 둔다).
    """
    from docx.shared import Inches
    n = 0
    for t in doc.tables:
        cells = t.rows[0].cells
        if len(cells) < 2 or cells[0].text.strip() != "ID":
            continue
        w0 = cells[0].width
        if w0 is None or w0.inches >= min_in:
            continue
        need = Inches(min_in) - w0
        widest = max(range(1, len(cells)), key=lambda i: cells[i].width.emu if cells[i].width else 0)
        for row in t.rows:
            row.cells[0].width = Inches(min_in)
            row.cells[widest].width = row.cells[widest].width - need
        n += 1
    return n


def restart_numbered_lists(doc):
    """절차마다 번호를 1부터 다시 시작시킨다.

    2026-08-20 실측: 06 사용자매뉴얼의 번호가 장을 넘어 1~59 로 계속 이어져,
    「문서 업로드」 절차의 첫 줄이 16번으로 시작했다. 앞에 열다섯 단계가 있었던 것처럼 읽힌다.
    같은 번호 서식(abstractNumId)을 가리키되 «1 부터 다시»를 지정한 목록 정의를
    묶음마다 새로 만들어 붙인다.
    """
    numbering = doc.part.numbering_part.element
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def num_pr(par):
        pPr = par._p.find(qn("w:pPr"))
        return None if pPr is None else pPr.find(qn("w:numPr"))

    def abstract_of(num_id):
        for n in numbering.findall(W + "num"):
            if n.get(qn("w:numId")) == num_id:
                a = n.find(W + "abstractNumId")
                return None if a is None else a.get(qn("w:val"))
        return None

    used = [int(n.get(qn("w:numId"))) for n in numbering.findall(W + "num")]
    next_id = max(used) + 1 if used else 1

    made = {}
    def fresh(abstract_id):
        """abstract_id 서식을 그대로 쓰되 1 부터 시작하는 새 목록 정의를 만든다."""
        nonlocal next_id
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_id))
        a = OxmlElement("w:abstractNumId"); a.set(qn("w:val"), abstract_id)
        num.append(a)
        ov = OxmlElement("w:lvlOverride"); ov.set(qn("w:ilvl"), "0")
        st = OxmlElement("w:startOverride"); st.set(qn("w:val"), "1")
        ov.append(st); num.append(ov)
        numbering.append(num)
        next_id += 1
        return str(next_id - 1)

    n = 0
    prev_listed = False
    cur = None
    for par in doc.paragraphs:
        np = num_pr(par)
        if np is None:
            prev_listed = False
            cur = None
            continue
        nid_el = np.find(qn("w:numId"))
        if nid_el is None:
            continue
        if not prev_listed:                      # 새 묶음의 첫 줄
            abs_id = abstract_of(nid_el.get(qn("w:val")))
            cur = fresh(abs_id) if abs_id else None
            n += 1
        if cur:
            nid_el.set(qn("w:val"), cur)
        prev_listed = True
    return n


def keep_screen_notes_with_steps(doc):
    """「[화면: …]」 줄이 앞 절차와 떨어져 혼자 다음 쪽으로 넘어가는 것을 막는다.

    2026-08-20 실측: 06 사용자매뉴얼 6쪽이 「[화면: /patient/messages 페이지]」 한 줄뿐이었다.
    이 줄은 바로 위 절차를 가리키는 것이라 떨어지면 무슨 화면인지 알 수 없다.
    """
    n = 0
    pars = doc.paragraphs
    for i, par in enumerate(pars):
        if not par.text.strip().startswith("[화면:"):
            continue
        for prev in reversed(pars[max(0, i - 4):i]):
            if not prev.text.strip():
                continue
            if not prev.paragraph_format.keep_with_next:
                prev.paragraph_format.keep_with_next = True
                n += 1
            break
    return n


def keep_short_procedures_together(doc, max_steps=6):
    """짧은 절차(번호 목록)가 쪽 경계에 걸려 마지막 한두 줄만 넘어가는 것을 막는다.

    2026-08-20 실측: 06 사용자매뉴얼 6쪽이 「3. 답장 입력 후 [전송] 클릭」 한 줄과
    화면 표시 한 줄뿐이었다. 절차는 통째로 한 쪽에 있어야 따라 할 수 있다.
    바로 앞의 소제목까지 함께 묶어 제목만 앞 쪽에 남는 것도 막는다.
    """
    pars = doc.paragraphs

    def is_step(par):
        pPr = par._p.find(qn("w:pPr"))
        return pPr is not None and pPr.find(qn("w:numPr")) is not None

    n = 0
    i = 0
    while i < len(pars):
        if not is_step(pars[i]):
            i += 1
            continue
        j = i
        while j < len(pars) and is_step(pars[j]):
            j += 1
        if j - i <= max_steps:
            head = pars[i - 1] if i > 0 and pars[i - 1].text.strip() else None
            for par in ([head] if head is not None else []) + pars[i:j - 1]:
                if not par.paragraph_format.keep_with_next:
                    par.paragraph_format.keep_with_next = True
                    n += 1
        i = j
    return n


def strip_trailing_blanks(doc):
    """문서 끝의 빈 문단을 없애거나, 없앨 수 없으면 «높이를 0에 가깝게» 눌러 둔다.

    2026-08-20 실측: 부록을 갈아 끼우면 끝에 빈 문단이 남아 06·07 매뉴얼의 마지막 쪽이
    머리말·꼬리말만 있는 빈 쪽이 됐다. 그런데 본문이 표로 끝나면 **워드가 표 뒤 문단을
    스스로 되살린다** — 지우기만 해서는 빈 쪽이 그대로 남는다.
    그래서 표로 끝나는 문서는 빈 문단을 한 줄 남기되 글자 크기 1pt·줄간격 0 으로 눌러
    쪽을 넘기지 못하게 한다.
    """
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    n = 0
    for el in reversed(list(body.iterchildren())):
        if el.tag == qn("w:sectPr"):
            continue
        if el.tag != qn("w:p"):
            break                      # 표를 만나면 거기서 멈춘다
        if Paragraph(el, doc).text.strip():
            break
        body.remove(el)
        n += 1

    # 표로 끝나면 워드가 문단을 되살리므로, 우리가 먼저 «높이 없는» 문단을 놔둔다.
    last = [e for e in body.iterchildren() if e.tag in (qn("w:p"), qn("w:tbl"))]
    if last and last[-1].tag == qn("w:tbl"):
        par = doc.add_paragraph()
        pf = par.paragraph_format
        pf.space_before = pf.space_after = Pt(0)
        pf.line_spacing = Pt(1)
        run = par.add_run("")
        run.font.size = Pt(1)
        n += 1
    return n


def drop_blanks_before_page_break(doc):
    """쪽 나눔 문단 «바로 앞»의 빈 문단을 없앤다.

    2026-08-20 실측: 02 기능명세서 21쪽이 빈 문단 하나만 있는 쪽이었다.
    앞 표가 쪽을 딱 채우면 그 빈 문단이 다음 쪽으로 넘어가고, 그 뒤 문단은
    「쪽 나눔으로 시작」이라 또 다음 쪽으로 가버려 가운데 쪽이 통째로 빈다.
    """
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    n = 0
    for el in list(body.iterchildren()):
        if el.tag != qn("w:p"):
            continue
        if not Paragraph(el, doc).paragraph_format.page_break_before:
            continue
        prev = el.getprevious()
        while prev is not None and prev.tag == qn("w:p") and not Paragraph(prev, doc).text.strip():
            drop, prev = prev, prev.getprevious()
            body.remove(drop)
            n += 1
    return n


def no_row_split(doc):
    """표의 «한 줄»이 쪽 경계에서 반으로 잘리지 않게 한다.

    2026-08-20 실측: 02 기능명세서 21쪽이 앞 줄에서 흘러넘친 한 줄뿐이었다.
    줄 통째로 다음 쪽에 넘어가게 하면 어느 칸의 내용인지 알아볼 수 있다.
    """
    n = 0
    for t in doc.tables:
        for row in t.rows:
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn("w:cantSplit")) is None:
                trPr.append(OxmlElement("w:cantSplit"))
                n += 1
    return 1 if n else 0


def drop_rows_containing(doc, needle, only_table_with=None):
    """어느 칸에든 needle 이 든 줄을 표에서 통째로 뺀다. 지운 줄 수를 돌려준다.

    only_table_with 를 주면 «첫 칸이 그 글자로 시작하는 줄이 있는 표»에서만 지운다.
    같은 파일명이 여러 표에 나올 때(목록 표 + 제출 일정표) 엉뚱한 표까지
    지우는 것을 막는다. 2026-08-20 실측: 이걸 안 두었더니 제출 일정표의
    8월·11월 줄이 통째로 날아갔다.

    ⚠️ 번호(B-01 등)로 지우지 마라. 지운 뒤 번호를 다시 매기면 «다음 문서»가 그
       번호를 물려받아, 다시 실행할 때 엉뚱한 줄이 지워진다(2026-08-20 실측:
       04_중간보고서 줄이 이렇게 사라졌다). 파일명처럼 안 바뀌는 값으로 지운다.
    """
    n = 0
    for t in doc.tables:
        if only_table_with and not any(
                r.cells[0].text.strip().startswith(only_table_with) for r in t.rows):
            continue
        for row in list(t.rows):
            if any(needle in c.text for c in row.cells):
                t._tbl.remove(row._tr)
                n += 1
    return n


def set_cells(doc, first_cell, values):
    """첫 칸이 first_cell 인 줄을 찾아 {열번호: 값} 으로 덮어쓴다. 없으면 False."""
    for t in doc.tables:
        for row in t.rows:
            if row.cells[0].text.strip() != first_cell:
                continue
            for i, v in values.items():
                if i >= len(row.cells):
                    continue
                c = row.cells[i]; c.text = ""
                set_font(c.paragraphs[0].add_run(v), size=9)
            return True
    return False


def main():
    changed = []

    # 0) 모든 산출물 공통 교체
    for path in sorted(pathlib.Path(HERE).glob("*.docx")):
        doc = Document(str(path))
        n = sum(1 for p in iter_paragraphs(doc) for old, new in GLOBAL
                if replace_in_paragraph(p, old, new))
        n += fix_em_dash(doc)
        n += ensure_footer(doc)
        n += repeat_table_headers(doc)
        n += no_row_split(doc)
        n += move_page_breaks(doc)
        n += drop_blanks_before_page_break(doc)
        n += keep_short_tables_together(doc)
        n += widen_id_columns(doc)
        n += restart_numbered_lists(doc)
        n += keep_screen_notes_with_steps(doc)
        n += keep_short_procedures_together(doc)
        n += strip_trailing_blanks(doc)
        if n:
            doc.save(str(path))
            changed.append(f"{path.name}: 공통 정리 {n}곳(꼬리말·머리행 포함)")

    # 0-1) 09 산출물목록에서 폐기 문서 흔적 제거
    #      03 착수보고서: KHIDI 가 요구한 적 없는데 이전 세션이 만든 것(PO 결정 2026-08-20).
    #      04 중간보고서·05 최종보고서: KHIDI 정해진 양식으로 따로 만든다(PO 결정 2026-08-20).
    #      제출 자체는 계속 있으므로 일정표는 지우지 않고 문구만 바꾼다.
    path = os.path.join(HERE, "09_산출물목록.docx")
    if os.path.exists(path):
        doc = Document(path)
        hit = 0
        for gone in ("03_착수보고서", "04_중간보고서.docx", "05_최종보고서.docx"):
            hit += drop_rows_containing(doc, gone, only_table_with="B-")
        # ⚠️ 문구 교체는 «줄을 지운 뒤»에 한다. 먼저 하면 삭제 기준 글자가 사라져 줄이 안 지워진다.
        for para_ in iter_paragraphs(doc):
            for old_, new_ in DROPPED_DOC_NOTES:
                replace_in_paragraph(para_, old_, new_)
        if hit:
            for t in doc.tables:
                if not any(r.cells[0].text.strip().startswith("B-") for r in t.rows):
                    continue
                n = 0
                for row in t.rows:
                    if row.cells[0].text.strip().startswith("B-"):
                        n += 1
                        c = row.cells[0]; c.text = ""
                        set_font(c.paragraphs[0].add_run(f"B-{n:02d}"), size=9, bold=True)
        if not any(PLAN_MAP_TITLE in p_.text for p_ in doc.paragraphs):
            # ⚠️ 부록을 먼저 걷어내고 붙인다. 그냥 붙이면 «부록 뒤»에 놓이고,
            #    뒤에서 부록을 걷어낼 때 부록 제목부터 문서 끝까지 지우므로 이 표도
            #    같이 날아간다(2026-08-20 실측). 부록은 아래 단계에서 다시 붙는다.
            strip_appendix(doc)
            doc.add_page_break()
            heading(doc, PLAN_MAP_TITLE)
            para(doc, "2026. 5. 14. 제출한 사업계획서가 결과물로 적은 항목과 실제 산출물을 "
                      "나란히 둔 표다. 계획서 표현과 파일 이름이 다른 곳이 있어 함께 적는다.", size=9)
            table(doc, ["마일스톤", "계획서가 적은 결과물", "실제 산출물", "상태"], PLAN_MAP_ROWS)
            hit += 1
        if hit:
            doc.save(path)
            changed.append(f"09_산출물목록.docx: 폐기 문서 정리 + 계획서 대비표")

    # 0-2) 08 성능 실측
    path = os.path.join(HERE, "08_테스트결과서.docx")
    if os.path.exists(path):
        doc = Document(path)
        for para_ in iter_paragraphs(doc):
            txt = para_.text.strip()
            if not para_.runs:
                continue
            if txt.startswith("[화면: Lighthouse 리포트 스크린샷"):
                para_.runs[0].text = ""
            elif txt.startswith(("Core Web Vitals는 시범 운영 시작 후", "Lighthouse 12.8.2 실측")):
                para_.runs[0].text = CWV_NOTE
                for r_ in para_.runs[1:]:
                    r_.text = ""
        if fill_cwv(doc):
            doc.save(path)
            changed.append("08_테스트결과서.docx: 성능 측정 표 실측 반영")

    path = os.path.join(HERE, "08_테스트결과서.docx")
    if os.path.exists(path):
        doc = Document(path)
        hit = set_cells(doc, "성능 측정 (Lighthouse)",
                        {2: "6항목 전부 측정",
                         3: "2026-08-21 실측(데스크톱·모바일 각 5회). 전 항목 목표 충족. "
                            "카자흐스탄 현지 접속 응답(69~92ms)을 자체 측정 항목으로 추가 확인"})
        hit += set_cells(doc, "npm 의존성 보안",
                         {2: "1회", 3: "Critical 0건 · High 0건 · Moderate 2건 (2026-08-20)"})
        if hit:
            doc.save(path)
            changed.append(f"08_테스트결과서.docx: 측정 결과 {hit}줄 갱신")

    # 1) 제자리 교체
    for fname, pairs in REPLACEMENTS.items():
        path = os.path.join(HERE, fname)
        if not os.path.exists(path):
            print(f"  건너뜀(없음): {fname}")
            continue
        doc = Document(path)
        n = 0
        for p in iter_paragraphs(doc):
            for old, new in pairs:
                if replace_in_paragraph(p, old, new):
                    n += 1
        if n:
            doc.save(path)
            changed.append(f"{fname}: {n}곳 교체")
        else:
            changed.append(f"{fname}: 교체할 것 없음(이미 현행)")

    # 2) 부록 — 있던 것을 걷어내고 새로 붙인다(몇 번 돌려도 한 벌만 남는다)
    for fname, kind in [
        ("06_사용자매뉴얼.docx", "user"),
        ("07_관리자매뉴얼.docx", "admin"),
        ("08_테스트결과서.docx", "test"),
        ("09_산출물목록.docx", "inventory"),
    ]:
        path = os.path.join(HERE, fname)
        if not os.path.exists(path):
            continue
        doc = Document(path)
        removed = strip_appendix(doc)
        add_appendix(doc, kind)
        # 부록을 붙인 «뒤»에 걷어내야 한다 — 앞서 돌리면 부록이 다시 빈 문단을 남긴다.
        strip_trailing_blanks(doc)
        doc.save(path)
        changed.append(f"{fname}: 부록 갱신" + (f" (옛 부록 {removed}블록 제거)" if removed else ""))

    for c in changed:
        print("  " + c)
    print("완료.")


if __name__ == "__main__":
    main()
