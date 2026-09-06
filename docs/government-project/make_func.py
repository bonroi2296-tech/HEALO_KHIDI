import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import _facts as F   # 사실의 단일 소스 — 2026-09-06 부터 이 생성기도 읽는다(전엔 안 읽어 드리프트)

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

def set_font(run, size=10, bold=False, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.append(rFonts)

def add_heading(doc, text, level=1, color=(0,70,127)):
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    size = 14 if level == 1 else (12 if level == 2 else 11)
    set_font(run, size=size, bold=True, color=color)
    return p

def add_para(doc, text, indent=0, size=10):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_header_row(table, cols, bg='00467F'):
    row = table.add_row()
    for i, col in enumerate(cols):
        if i < len(row.cells):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(col)
            set_font(run, size=9, bold=True, color=(255,255,255))
            set_cell_bg(cell, bg)

def add_data_row(table, cols, bold_first=False, bg=None):
    row = table.add_row()
    for i, cell in enumerate(row.cells):
        if i < len(cols):
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cols[i]))
            b = bold_first and i == 0
            set_font(run, size=9, bold=b)
            if bg:
                set_cell_bg(cell, bg)

def func_spec_table(doc, spec):
    """기능 명세 표 생성"""
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    add_header_row(tbl, ['항목', '내용'])
    for k, v in spec.items():
        add_data_row(tbl, [k, v], bold_first=True)
    doc.add_paragraph()

# ================================================================
# 표지
# ================================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"{F.PROJECT['플랫폼']} 플랫폼")
set_font(run, 20, True, (0,70,127))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('기능 명세서')
set_font(run, 24, True, (0,70,127))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Functional Specification Document')
set_font(run, 12, False, (80,80,80))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'ICT 기반 외국인환자 사전상담·사후관리 지원 사업  |  {F.DOC_VERSION}  |  {F.DOC_DATE_DOT}')
set_font(run, 10, False, (120,120,120))

doc.add_paragraph()
doc.add_paragraph()

tbl_info = doc.add_table(rows=0, cols=2)
tbl_info.style = 'Table Grid'
for label, val in [
    ('문서 번호', 'healwith-FSD-2026-001'),
    ('작성 기관', '본로이 (Bonroi)'),
    ('작성일', F.DOC_DATE_KO),
    ('버전', F.DOC_VERSION),
    ('관련 문서', 'healwith-REQ-2026-001 요구사항정의서 · 01-1 화면설계서 · 11 6대 ICT 대비표'),
    ('갱신 기준', f'코드(본판 2026-09-06)·운영DB 정확 집계({F.AS_OF}). 사실은 _facts.py 한 곳에서 읽는다'),
]:
    row = tbl_info.add_row()
    row.cells[0].text = ''
    r1 = row.cells[0].paragraphs[0].add_run(label)
    set_font(r1, 10, True, (0,70,127))
    row.cells[1].text = ''
    r2 = row.cells[1].paragraphs[0].add_run(val)
    set_font(r2, 10)

doc.add_page_break()

# ================================================================
# 목차
# ================================================================
add_heading(doc, '목  차', 1)
for num, title in [
    ('1.', '기능 그룹 개요'),
    ('2.', '회원·인증·권한 관리'),
    ('3.', '환자 문의 접수·문서 업로드'),
    ('4.', '병원 매칭·정보 제공'),
    ('5.', '원격 화상상담 (LiveKit)'),
    ('6.', 'AI 챗봇·사람 상담원 상담'),
    ('7.', 'AI 실시간 번역'),
    ('8.', '예약·일정·비자 관리'),
    ('9.', '사후관리·모니터링·교육'),
    ('10.', '다국어 (ko/en/ru/kz/zh/ja)'),
    ('11.', '코디네이터 포털'),
    ('12.', '관리자·국내 의료기관 포털'),
    ('13.', '알림·이메일·푸시'),
    ('14.', 'AI 학습 파이프라인 (RAG)'),
    ('15.', '보안·암호화·감사'),
    ('16.', '해외 파트너 포털 (에이전시·해외 의료기관)'),
    ('17.', '전문의 소견 · 환자 진행상황 · 병원 의뢰서'),
    ('18.', '만족도 설문 · 사후관리 자동 실행'),
    ('19.', '모바일 앱(스토어 판) · 음성 메모 판독'),
    ('20.', '콘텐츠 편집기 · 번역 품질 · 검색 유입'),
    ('21.', 'AI 품질 보증 · 비용 계측'),
    ('22.', '운영 자동화(정기 실행) · 감시'),
    ('부록 A', '기능 구현 상태 요약'),
    ('부록 B', '남은 일 · 기술 부채'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5) if '.' in num and len(num) > 2 else Cm(0)
    run = p.add_run(f'{num}  {title}')
    set_font(run, 10)

doc.add_page_break()

# ================================================================
# 1. 기능 그룹 개요
# ================================================================
add_heading(doc, '1. 기능 그룹 개요', 1)
add_para(doc, f"본 문서는 {F.PROJECT['플랫폼']} 플랫폼의 21개 기능 그룹에 대한 상세 명세를 기술한다. 각 기능마다 시나리오(액터·전제·트리거·흐름·예외), 입출력, 화면 라우트, DB 테이블/컬럼, 현재 구현 상태를 명시한다. 화면 경로는 실제 폴더(app/)와 1:1 이며, 코드에 없는 것은 「없음」이라고 적는다.")
add_para(doc, '※ 현황 표기: 완료 / 부분구현 / 미구현')
doc.add_paragraph()

# 기능 그룹 요약 표
tbl = doc.add_table(rows=0, cols=4)
tbl.style = 'Table Grid'
add_header_row(tbl, ['그룹', '기능 그룹명', '핵심 라우트', '현황'])
groups = [
    ('G-01', '회원·인증·권한', '/signup, /login, proxy.ts', '완료'),
    ('G-02', '환자 문의 접수·문서', '/inquiry, /inquiry/intake, /api/attachments', '완료'),
    ('G-03', '병원 매칭·정보', '/hospitals, /treatments, /api/public/chat/*', '완료'),
    ('G-04', '원격 화상상담', '/consultation/[id], /c/[code], /api/khidi/consultation/*', '완료'),
    ('G-05', 'AI 챗봇·Human 상담', '/api/public/chat/*, /coordinator/messages, /api/webhooks/*', '완료'),
    ('G-06', 'AI 실시간 번역·자막', '/api/translate-text, /api/khidi/consultation/[id]/stt', '완료'),
    ('G-07', '예약·일정·비자', '/coordinator/consultations, /patient/visa, src/lib/calendar', '부분(병원 가용 일정 입력 화면 없음)'),
    ('G-08', '사후관리·모니터링·교육', '/patient/symptoms, /education, /api/cron/dispatch-surveys', '완료'),
    ('G-09', '다국어 UI', '/ru, /kz … proxy.ts, src/lib/i18n', '완료'),
    ('G-10', '코디네이터 포털', '/coordinator/* (메뉴 16개)', '완료'),
    ('G-11', '관리자·국내 의료기관 포털', '/admin/* (7그룹), /hospital/*', '완료 / 병원 포털은 리드 관리만'),
    ('G-12', '알림·이메일·푸시', 'src/lib/email, src/lib/notifications, src/lib/push', '완료(인앱은 30초 폴링)'),
    ('G-13', 'AI RAG 파이프라인', 'src/lib/rag, src/lib/chat/externalSearch.ts', '완료(구글 검색 그라운딩만 미작동)'),
    ('G-14', '보안·암호화·감사', 'src/lib/security, src/lib/auth, /admin/audit', '완료'),
    ('G-15', '해외 파트너 포털', '/agency/*, /clinic/*, /api/agency/*', '완료'),
    ('G-16', '전문의 소견·환자 진행상황·의뢰서', '/opinion/[token], /claim/[token], /api/coordinator/inquiries/[id]/referral-docx', '완료'),
    ('G-17', '만족도 설문·사후관리 자동 실행', '/survey/[token], /api/cron/dispatch-surveys', '완료(실환자 응답 0)'),
    ('G-18', '모바일 앱·음성 메모 판독', 'capacitor.config.ts, /app, /coordinator/voice', '완료(스토어 심사 대기)'),
    ('G-19', '콘텐츠 편집기·번역 품질·검색 유입', '/coordinator/content, scripts/check-i18n-*.mjs, /api/cron/indexnow', '완료'),
    ('G-20', 'AI 품질 보증·비용 계측', 'src/lib/chat/regressionRunner.ts, /admin/khidi/ai-*', '완료'),
    ('G-21', '운영 자동화·감시', 'vercel.json crons(11), sentry.*.config.js, /api/health', '완료'),
]
for g in groups:
    add_data_row(tbl, g, bold_first=True)

doc.add_page_break()

# ================================================================
# 헬퍼: 시나리오 블록 생성
# ================================================================
def add_scenario(doc, scenario_dict):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    add_header_row(tbl, ['시나리오 항목', '내용'], bg='2E6DA4')
    for k, v in scenario_dict.items():
        add_data_row(tbl, [k, v], bold_first=True)
    doc.add_paragraph()

# ================================================================
# 2. 회원·인증·권한 관리
# ================================================================
add_heading(doc, '2. 회원·인증·권한 관리', 1)

add_heading(doc, '2.1 회원가입 및 로그인 (FN-AUTH-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-AUTH-01',
    '기능명': '이메일·소셜 회원가입 및 로그인 (웹 + 스토어 앱)',
    '액터': '환자, 코디네이터, 관리자, 파트너 담당자',
    '전제 조건': '네트워크 연결, 유효한 이메일 주소 보유',
    '트리거': '사용자가 /signup 또는 /login 페이지 접근',
    '기본 흐름': '① 이메일·비밀번호 또는 Google·Apple 로그인 선택(앱 안에서는 네이티브 로그인 — googleNativeSignIn.ts·appleNativeSignIn.ts, nonce 는 «구글엔 해시·Supabase 엔 원본»)\n② Supabase Auth 인증 처리\n③ 세션 쿠키 발급 (@supabase/ssr)\n④ 계층에 따라 적절한 화면으로 이동(환자는 /patient, 코디는 /coordinator …)',
    '예외 흐름': '• 이메일 미인증 → 인증 메일 재발송 안내\n• 잘못된 비밀번호 → 「인증 실패」 메시지 (error.message 직접 노출 금지)\n• 소셜 OAuth 실패 → fallback 이메일 로그인 안내',
    '입력': '이메일, 비밀번호 (또는 OAuth 토큰)',
    '출력': 'Supabase 세션 쿠키, 리다이렉트',
    '화면 경로': '/app/signup, /app/login',
    'DB 테이블': 'auth.users (Supabase 관리)\napp_metadata.role 컬럼',
    '현재 구현 상태': '완료: proxy.ts, src/lib/auth, src/components/auth/AppleSignInButton.jsx\n앱 안 애플 로그인은 아직 끝까지 못 간다(스토어 심사 대응 중, docs/APP_STORE_LISTING.md)',
})

add_heading(doc, '2.2 역할 기반 접근 제어 (FN-AUTH-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-AUTH-02',
    '기능명': 'RBAC: 역할 기반 라우트 보호',
    '액터': '시스템 (Next.js 미들웨어)',
    '전제 조건': '사용자 세션 쿠키 존재',
    '트리거': '보호된 라우트 (/patient/*, /admin/*, /coordinator/*) 접근 시',
    '기본 흐름': '① Next.js proxy.ts(구 middleware)에서 세션 검증\n② /admin 은 proxy 가 app_metadata.role 까지 확인해 불일치 시 /login 으로\n③ /patient·/coordinator·/hospital·/agency·/clinic 은 proxy 가 «세션 존재»만 보고 통과시키고, 세부 권한은 화면·API 가 판정(requireAdminAuth / checkAgencyAuth / hospital_users 조회)\n④ 권한 일치 → 렌더링',
    '예외 흐름': '• 세션 만료 → 재로그인 안내\n• role 미설정 → 기본 환자 권한\n• 화면에서 감추는 것만으로 막지 않는다(서버 판정이 정본)',
    '역할 정의': '비회원(게스트): 상담방 초대링크 토큰으로만 /consultation/[id] 입장\n환자: /patient/* 접근\ncoordinator: /coordinator/* 접근\nadmin: /admin/* 전체 접근\n국내 의료기관: /hospital/* 접근\n해외 에이전시: /agency/* 접근\n해외 의료기관: /clinic/* 접근\n※ 의사는 계정 계층이 아님: 초대링크 게스트 또는 병원 계정으로 참여',
    '화면 경로': 'proxy.ts (루트 레벨)',
    'DB/컬럼': 'auth.users.app_metadata.role (user_metadata 사용 금지)',
    '현재 구현 상태': '완료: proxy.ts, src/lib/auth/requireAdminAuth',
})

add_heading(doc, '2.3 게스트 토큰 (FN-AUTH-03)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-AUTH-03',
    '기능명': '비회원 게스트 공개 토큰 발급',
    '액터': '비회원 환자',
    '전제 조건': '회원가입 없이 상담 시작 요청',
    '트리거': '최초 문의 접수 또는 채팅 시작',
    '기본 흐름': '① UUID 기반 public_token 생성\n② inquiries 테이블에 token 저장\n③ 쿠키/로컬스토리지에 token 보관\n④ 이후 요청에 token으로 상담 이력 연속성 유지',
    '예외 흐름': '• 토큰 만료 → 새 토큰 발급, 이전 이력 연결 불가',
    '화면 경로': '/inquiry (토큰 발급 시작점)',
    'DB 테이블/컬럼': 'inquiries.public_token (migrations/20260125_inquiries_public_token_and_attachments)',
    '현재 구현 상태': '완료: migrations/20260125_inquiries_public_token_and_attachments',
})

doc.add_page_break()

# ================================================================
# 3. 환자 인테이크·문서
# ================================================================
add_heading(doc, '3. 환자 문의 접수·문서 업로드', 1)

add_heading(doc, '3.1 암환자 문의 접수 2단계 (FN-INTAKE-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-INTAKE-01',
    '기능명': '2단계 문의 접수(1단계 접수 → 2단계 임상 자료)',
    '액터': '해외 암환자 (P-01)',
    '전제 조건': '웹/앱 접속, 다국어 UI 선택 (러시아어 등)',
    '트리거': '/inquiry 페이지 접근',
    '기본 흐름': '① 1단계 접수: 이름·연락 수단·국적·선호 언어 / 암종·병기 / 하고 싶은 말\n② 동의 필수 4종(개인정보·민감정보·제3자 제공·국외이전) + 선택 1종(마케팅) 확인 후 제출 — 여기까지만 채워도 접수된다(동의서 판번호·시각 기록)\n③ 접수 뒤 2단계(/inquiry/intake): 문의번호와 접근 토큰이 있어야 열린다\n④ 2단계에서 대학병원 의뢰에 필요한 검사자료·음성 메모를 추가로 받는다(6개 언어, 자동 저장)\n⑤ inquiries·cancer_patient_intakes 저장, 코디네이터 종 알림 + 메일. 시험 문의는 표식으로 갈라 알림·실적에서 뺀다(건너뛴 이유는 admin_notify_skipped 로 남김)\n⑥ 접수한 사람에게 진행상황 링크(/claim/<토큰>)를 그 채널로 돌려준다',
    '예외 흐름': '• 필수 항목 미입력 → 단계 진행 불가\n• 파일 업로드 실패 → 재시도 안내\n• 네트워크 오류 → 임시 저장 (로컬스토리지)',
    '입력': '성명(암호화), 연락처(암호화), 국가, 암 종류, 병기, 파일, 희망 내용',
    '출력': 'inquiry 레코드 생성, 코디네이터 이메일 알림',
    '화면 경로': '/app/inquiry (통합 문의 퍼널)',
    'DB 테이블/컬럼': 'inquiries (id, first_name, last_name, email, phone — 이 넷에 암호문 저장, cancer_type, public_token, status, outcome, is_test, access_log)\ncancer_patient_intakes (암종·병기·치료 이력 — *_encrypted 칸)\nmigrations/20260125_inquiries_intake_progressive',
    '보안': 'AES-256-GCM 암호화 (encryptionV2.ts)\n개인정보보호법 동의 수집',
    '현재 구현 상태': '완료: /app/inquiry (통합 문의 퍼널) · /app/inquiry/intake (2단계) · /app/inquiry/referral (병원 의뢰서 형식)\nmigrations/20260125_inquiries_intake_progressive\n※ 식별정보 값은 전건 암호문이지만 «평문 이름의 칸»에 저장 중이다. 평문 칸 삭제 마이그레이션(20260420_drop_*)은 실행 금지(돌리면 26~29건 소실) — 부록 B',
})

add_heading(doc, '3.2 의료문서 업로드·관리 (FN-INTAKE-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-INTAKE-02',
    '기능명': '의료문서(CT/MRI/검사결과) 업로드 및 관리',
    '액터': '환자, 코디네이터',
    '전제 조건': '문의 접수 완료 또는 상담 진행 중',
    '트리거': '파일 첨부(점선 상자에 끌어놓기 또는 선택)',
    '기본 흐름': '① 환자·코디가 파일 업로드 — 브라우저에서 저장소로 직행(uploadDirect)\n② 형식·용량 검증(src/lib/uploadPolicy.js): PDF·JPG·PNG·GIF·WebP·Word·DICOM·txt·음성 8종(mp3/m4a/wav/ogg/opus/webm/amr), 영상자료는 ZIP/RAR(병원 CD 통째). 건당 200MB. 파일 머리(매직 바이트)로 위장 차단\n③ Supabase Storage 저장 + 표에 메타데이터(러시아어·한글 파일명 그대로 — 2026-09-04 이중 인코딩 수정)\n④ 같은 파일 두 줄 저장 차단, 본인 문서 소프트 삭제\n⑤ 코디네이터·의료진이 열람·[전부 받기]. CT 는 브라우저에서 넘겨 보고, 큰 스캔 서류는 AI 가 읽어 원문 1:1 번역',
    '예외 흐름': '• 허용 외 파일 형식 → 업로드 거부 및 안내\n• 파일 크기 초과 → 압축 요청\n• Storage 오류 → 재시도',
    '입력': '의료 파일 (CT DICOM, MRI, 검사결과 PDF, 진단서)',
    '출력': 'Supabase Storage URL, attachments 레코드',
    '화면 경로': '/app/inquiry (업로드), /app/patient/documents (관리)',
    'DB 테이블/컬럼': 'consultation_documents (id, consultation_id, file_name, file_type, file_size, storage_path, document_type, uploaded_by, created_at)\nconsultation_documents (migrations/20260406_consultation_documents)',
    '현재 구현 상태': '완료: /app/api/attachments, /app/patient/documents, src/lib/uploadPolicy.js, src/lib/imaging/prepareStudy.ts\nmigrations/20260406_consultation_documents',
})

doc.add_page_break()

# ================================================================
# 4. 병원 매칭
# ================================================================
add_heading(doc, '4. 병원 매칭·정보 제공', 1)

add_heading(doc, '4.1 AI 기반 병원·의료진 매칭 (FN-MATCH-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-MATCH-01',
    '기능명': 'AI 기반 병원·의료진 자동 매칭',
    '액터': '환자 (P-01), AI Agent',
    '전제 조건': '접수 정보 또는 채팅 증상 입력 완료',
    '트리거': '상담 시작 또는 「병원 추천」 요청',
    '기본 흐름': '① 환자 증상·희망 치료·예산·국가 입력\n② AI Agent (Gemini Flash 최신 별칭) 쿼리 생성\n③ RAG 1계층: 자체 DB(rag_documents·rag_chunks, pgvector) 벡터 검색 — 답의 근거는 여기서 먼저 찾는다\n④ RAG 2계층: HIRA 공공 API + 네이버 지역검색 실시간 병렬 호출(src/lib/chat/externalSearch.ts)\n⑤ (3계층으로 설계했던 구글 검색 그라운딩은 미작동 — 설치된 SDK 에 없는 옵션이라 조용히 무시된다. 2026-07-31 「웹 검색 결과」 가짜 출처 라벨을 제거했다)\n⑥ 6개 언어로 매칭 결과 반환 (병원명·진료과·특장점·위치). DB 에 없는 병원은 추천하지 않는다',
    '예외 흐름': '• 해당 진료과 병원 미보유 → 유사 병원 추천\n• AI 응답 5초 초과 → 스트리밍 응답으로 부분 표시',
    '입력': '증상 텍스트, 암 종류, 선호 병원 유형, 예산',
    '출력': '병원 목록 (이름·진료과·의료진·특장점·예상비용·위치)',
    '화면 경로': '/app/hospitals, /app/api/chat',
    'DB 테이블/컬럼': 'hospitals (id, name, specialties, location_kr·location_en 등 언어별 칸, i18n JSONB)\nrag_documents (content, trust_tier, source_type, source_url) + rag_chunks (embedding vector)\nmigrations/20260225_rag_vector_v1, 20260223_i18n_jsonb',
    '참조 특허': '특허 10-2745881 (EMR 연동 플랫폼), 특허 10-2868334 (AI 기반 중개)',
    '현재 구현 상태': '완료: src/lib/chat/generateReply.ts, src/lib/chat/externalSearch.ts\nsrc/lib/rag/*, /app/api/public/chat/{start,message,stream,resume}\n※ 구글 검색 그라운딩만 미작동(부록 B)',
})

add_heading(doc, '4.2 병원 목록·상세 조회 (FN-MATCH-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-MATCH-02',
    '기능명': '병원 목록·상세 페이지 (다국어)',
    '액터': '환자, 방문자',
    '트리거': '/hospitals 또는 /hospitals/[id] 접근',
    '기본 흐름': '① 병원 목록 조회 (필터: 진료과·지역·인증 여부)\n② 선택 병원 상세 정보 표시\n③ 언어 설정에 따라 i18n JSONB 컬럼에서 현지어 텍스트 반환\n④ Google Maps로 위치 표시',
    '예외 흐름': '• 선택 언어 번역 미존재 → 한국어 fallback',
    '입력': '필터 조건 (진료과, 지역, 언어)',
    '출력': '병원 카드 목록, 상세 정보, 지도',
    '화면 경로': '/app/hospitals, /app/treatments',
    'DB 테이블/컬럼': 'hospitals (id, name, specialties, address, i18n JSONB, is_published)\ntreatments, specialties\nmigrations/20260223_i18n_jsonb, 20260125_add_is_published',
    '현재 구현 상태': '완료: /app/hospitals, /app/treatments\nmigrations/20260223_i18n_jsonb',
})

doc.add_page_break()

# ================================================================
# 5. 원격 화상상담
# ================================================================
add_heading(doc, '5. 원격 화상상담 (LiveKit WebRTC)', 1)

add_heading(doc, '5.1 화상상담 세션 생성 및 참여 (FN-VIDEO-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-VIDEO-01',
    '기능명': 'LiveKit 기반 WebRTC 화상상담',
    '액터': '환자 (P-01), 코디네이터/의료진 (P-02, P-03)',
    '전제 조건': '상담 예약 확정 또는 즉시 상담 요청',
    '트리거': '「화상상담 시작」 버튼 클릭',
    '기본 흐름': '① 코디가 /coordinator/consultations 에서 상담 생성(필수 입력은 환자·예약 시각 둘, KST + 협정시계 병기) → 구글 캘린더 자동 등록 + 초대 메일(.ics 첨부)\n② /api/khidi/consultation/token 에서 방 이름 기반 접근 토큰 발급\n③ 환자·의료진: 초대 주소(/c/<코드>)로 계정 없이 입장, 대기실에서 코디가 승인\n④ @livekit/components-react 기반 UI — 마이크·카메라·화면 공유·채팅·언어·자막 크기·음성 출력·[종료]\n⑤ 잡음 제거(Krisp NC)는 자동, 하울링 2단 감지, 좀비 방·좀비 탭 정리\n⑥ 코디가 [상담 완료]를 눌러야 consultation_sessions 가 completed 가 되고 성과지표에 잡힌다([종료]만 누르면 세지 않는다)',
    '예외 흐름': '• 카메라/마이크 권한 거부 → 오디오 전용\n• 네트워크 불안정 → LiveKit SDK 기본 재연결\n• 참여자 미접속 → 대기 유지. 19~90시간 열려 있던 좀비 방은 크론이 닫는다\n• 통화시간(livekit_duration_seconds)은 «참가자 2명 이상» 구간만 — 전날 시험 입장·좀비 탭으로 부풀던 것 2026-09-06 수정',
    '입력': '방 이름(room_name), 참여자 ID, 역할',
    '출력': 'LiveKit Access Token, 화상 UI',
    '화면 경로': '/app/consultation/[id]/page.jsx (상담방) · /app/c/[code] (초대 입장) · /app/telemedicine (소개 화면 — 실제 방이 아니다)',
    'DB 테이블/컬럼': 'consultation_sessions (id, inquiry_id, session_type, status, livekit_room_name, started_at, ended_at, livekit_duration_seconds, is_test, patient_timezone)\nconsultation_admissions (입퇴장) · consultation_guest_tokens\nmigrations/20260403_add_consultation_sessions',
    '연동 서비스': 'LiveKit Cloud Ship 플랜 (livekit-server-sdk)\n@livekit/components-react, livekit-client · 녹화(Egress)는 CONSULT_RECORDING_ENABLED 로 켤 때만(기본 꺼짐, 90일 보관 뒤 자동 파기)',
    '현재 구현 상태': '완료: /app/api/khidi/consultation/token, /app/api/livekit/webhook\n/app/consultation/[id]/page.jsx\nmigrations/20260403_add_consultation_sessions',
})

add_heading(doc, '5.2 게스트 참여 토큰 (FN-VIDEO-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-VIDEO-02',
    '기능명': '비회원 환자의 화상상담 참여',
    '액터': '비회원 환자',
    '전제 조건': '공개 상담 링크 보유 (코디네이터가 발송)',
    '트리거': '상담 링크 클릭',
    '기본 흐름': '① 코디네이터가 환자에게 상담 링크 전송\n② 환자가 링크 접속 (로그인 불필요)\n③ 초대 주소(/c/<코드>)가 상담을 찾아 게스트 토큰을 발급\n④ 이름 입력 후 화상상담 참여',
    '예외 흐름': '• 링크 만료 → 코디네이터에게 재발급 요청',
    '화면 경로': '/app/c/[code] (초대 입장 → 상담방)',
    '현재 구현 상태': '완료: /app/c/[code]/route.ts, consultation_guest_tokens(해시 저장)\n초대 주소는 57자, 상담 시각 +12시간(최소 72시간) 뒤 만료. 환자용 링크엔 ?lang= 을 실어 메신저 미리보기도 환자 언어로',
})

doc.add_page_break()

# ================================================================
# 6. AI 챗봇·사람 상담원 상담
# ================================================================
add_heading(doc, '6. AI 챗봇·사람 상담원 상담 시스템', 1)

add_heading(doc, '6.1 AI 챗봇 24시간 상담 (FN-CHAT-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-CHAT-01',
    '기능명': 'Gemini 기반 AI Agent 다국어 상담',
    '액터': '환자 (P-01), AI Agent',
    '전제 조건': '채팅 UI 로드, 사용자 언어 감지',
    '트리거': '환자 메시지 입력 및 전송',
    '기본 흐름': '① 환자가 채팅 입력 (텍스트, 다국어)\n② /api/public/chat/message 호출\n③ src/lib/chat/generateReply.ts 실행\n  - RAG 검색: rag_documents에서 관련 병원·치료 정보 벡터 검색\n  - Gemini Flash 로 응답 생성 (스트리밍)\n④ 응답 반환 (병원 정보, 치료 안내, 비용 안내 등)\n⑤ 복잡 케이스 감지 → Human Agent 이관 플래그 설정\n⑥ 대화 내용 chat_threads 테이블 저장',
    '예외 흐름': '• Gemini API 타임아웃 → 「잠시 후 다시 시도」 메시지\n• 의료 진단 요청 → 「의료진 상담 연결」 안내',
    '입력': '사용자 메시지 텍스트, 언어 코드, 상담 컨텍스트',
    '출력': '스트리밍 텍스트 응답, 병원 카드, Human 이관 여부',
    '화면 경로': '/app/inquiry ①(공개 위젯) · /app/patient/chat · 텔레그램/왓츠앱 봇(/app/api/webhooks/*)',
    'DB 테이블/컬럼': 'chat_threads (id, inquiry_id, status, channel[web|telegram|agency], guest_*, created_at) · chat_messages (본문) · chat_feedback\nrag_documents (content, trust_tier, lang) · rag_chunks (embedding, chunk_index)\nmigrations/20260225_chat_threads, 20260225_rag_vector_v1',
    '현재 구현 상태': '완료: /app/api/public/chat/{start,message,stream,resume,lookup,threads,consent,feedback}\nsrc/lib/chat/generateReply.ts (스트리밍, 첫 글자 중앙값 2~4초)\nsrc/lib/rag/*\n3턴을 넘기면 실제 문의로 승격(연락 수단 없으면 승격 안 함). 안전장치: 완치 단정·지어낸 병원·연락처 없는 「접수됨」 차단(매일 Chat Smoke)',
})

add_heading(doc, '6.2 Human Agent 이관 및 코디네이터 상담 (FN-CHAT-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-CHAT-02',
    '기능명': 'AI→Human 이관 및 코디네이터 직접 상담',
    '액터': '코디네이터 (P-02), 환자 (P-01)',
    '전제 조건': 'AI 이관 플래그 설정 또는 환자 직접 Human 요청',
    '트리거': 'AI 이관 알림 수신 또는 코디네이터 직접 응답',
    '기본 흐름': '① AI가 이관 플래그 설정 → 코디네이터 알림\n② 코디네이터 /coordinator/messages 접속\n③ 환자 대화 이력 전체 확인\n④ 코디네이터가 직접 메시지 작성·전송\n⑤ 상담 완료 후 AI 자동화 재활성화 여부 선택',
    '예외 흐름': '• 자료 업로드·「사람 연결」 요청 시 즉시 사람에게 이관\n• 상담 단계에서 7일 무동작 → 「식은 문의」 크론이 코디에게 알림(시간 기반 슈퍼바이저 승격은 없다)\n• 왓츠앱 24시간 창 만료·발송 실패는 말풍선 아래 미전달 표시',
    '화면 경로': '/app/coordinator/messages (채널 배지 웹·텔레그램·왓츠앱·에이전시) · 코디 텔레그램 양방향 릴레이',
    'DB 테이블/컬럼': 'chat_threads.status (사람 인계 여부 — 사람이 답하기 시작하면 AI 는 침묵)\nnotifications (코디 종 알림)',
    '현재 구현 상태': '완료: /app/coordinator/messages, src/lib/messaging/{telegram,whatsapp,staffRelay}.ts, /app/api/webhooks/{telegram,whatsapp}\n「추천 답장」 칩(자동 전송 아님)',
})

doc.add_page_break()

# ================================================================
# 7. AI 실시간 번역
# ================================================================
add_heading(doc, '7. AI 실시간 번역', 1)

add_heading(doc, '7.1 채팅 메시지 번역 (FN-TRANS-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-TRANS-01',
    '기능명': '채팅 메시지 AI 실시간 번역',
    '액터': '환자, 코디네이터, 의료진',
    '트리거': '메시지 전송 또는 「번역 보기」 버튼 클릭',
    '기본 흐름': '① 사용자 메시지 감지\n② /api/translate-text 호출(허용 도메인 검증 — 2026-08-20 남의 도메인 제거)\n③ Gemini 로 6개 언어(ko·en·ru·kz·zh·ja) 양방향 번역\n④ 상대 발화를 내 언어로 한 줄 자막 표시(두 언어를 나란히 놓지 않는다)',
    '화면 경로': '/app/api/translate-text',
    'DB 테이블/컬럼': '별도 저장 없음 (실시간 처리)',
    '참조 특허': '특허 10-2868334 (실시간 통역 가능 생성형AI 기반 중개 플랫폼)',
    '현재 구현 상태': '완료 (API): /app/api/translate/route.ts\n/app/api/translate-text/route.ts',
})

add_heading(doc, '7.2 화상상담 중 실시간 번역 (FN-TRANS-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-TRANS-02',
    '기능명': '화상상담 중 STT + AI 실시간 통번역',
    '액터': '환자, 의료진',
    '전제 조건': 'LiveKit 화상상담 세션 활성',
    '기본 흐름': '① 음성 → 받아쓰기(브라우저 STT 또는 서버 /api/khidi/consultation/[id]/stt — Gemini 전사+번역 1콜, 6개 언어)\n② 상대 발화를 내 언어로 자막 표시(화자 이름·시각·원어 동봉, 확정 자막과 하단 중간 자막 구분)\n③ 통역 «음성»이 필요하면 LiveKit 통역봇(agents/live-translate)을 온디맨드로 방에 부른다\n④ 자막·번역 기록은 암호화 저장(consultation_translations, is_partial 로 구분)',
    '예외 흐름': '• 자막은 누를 때만 켜진다(자동 켜짐 없음)\n• 지어냄(환각) 방지: 두 번 받아써 대조하는 거르개(2026-08-18) · 「누가 있는지」 설명을 프롬프트에서 제거\n• 번역문이 문장 중간에 잘리던 진범(출력 예산에 «생각» 토큰이 섞임)은 2026-09-01 한도 폐지로 해소\n• 통역봇 세션은 10분마다 약 10초 끊긴다(Gemini 세션 상한)',
    '화면 경로': '/app/consultation/[id]/page.jsx (자막 스택) · /app/api/khidi/consultation/[id]/{stt,interpreter}',
    '현재 구현 상태': '완료: 실서비스 가동 중 — 통역 자막 3,735건(시험 방 제외 3,396) 축적\nsrc/lib/consultation/{sttEngine,transcriptCrypto,…}.ts, agents/live-translate/\nSTT 후보 모델(Gemini 3.5 Transcribe)은 env 스위치로 실험(기본 꺼짐, 실패 시 Flash 폴백)',
})

doc.add_page_break()

# ================================================================
# 8. 예약·일정·비자
# ================================================================
add_heading(doc, '8. 예약·일정·비자 관리', 1)

add_heading(doc, '8.1 진료 예약·일정 관리 (FN-SCHED-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-SCHED-01',
    '기능명': '환자-의료진 일정 조율 및 예약 확정',
    '액터': '환자 (P-01), 코디네이터 (P-02)',
    '트리거': '상담 완료 후 예약 단계 진입',
    '기본 흐름': '① 코디가 환자와 시각을 정해 /coordinator/consultations 에서 상담 생성\n② 초대 메일(.ics 첨부, 환자 시간대 병기) + 구글 캘린더 자동 등록(서비스 계정 캘린더, src/lib/calendar/googleCalendar.ts)\n③ 리마인더 자동 발송(5분·10분 주기 크론, 이름 포함 — 2026-08-31 «이름 없이» 나가던 것 수정)\n④ 환자는 /patient/calendar 와 진행상황 화면에서 일정 확인\n⑤ (없음) 병원 가용 일정을 코디가 «입력할 화면»',
    '예외 흐름': '• 선택 날짜 의료진 부재 → 대안 일정 제시\n• 초대 링크 만료: 상담 시각 +12시간(최소 72시간) — 만료 시 재발급',
    '화면 경로': '/app/patient/calendar',
    'DB 테이블/컬럼': 'consultation_sessions (scheduled_at, status)\nmigrations/20260403_add_consultation_sessions',
    '현재 구현 상태': '부분구현: /app/coordinator/consultations, /app/patient/calendar, src/lib/calendar/googleCalendar.ts, src/lib/email/icsInvite.ts, /app/api/cron/{dispatch-reminders,consultation-reminders}\n남은 것 = 병원 가용 일정 입력 화면(부록 B)',
})

add_heading(doc, '8.2 비자 발급 안내 (FN-SCHED-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-SCHED-02',
    '기능명': '카자흐스탄→한국 의료비자 안내 서비스',
    '액터': '환자 (P-01)',
    '트리거': '예약 확정 후 또는 비자 정보 메뉴 접근',
    '기본 흐름': '① 비자 종류 안내 — 단기(90일 이내) C-3-3 / 장기(91일 이상) G-1-10, K-ETA 규정은 check:visa-freshness 가 기한을 감시\n② 필요 서류 목록(진단서, 예약 확인서, 재정증명 등) 6개 언어\n③ 처리 기간·수수료 안내\n④ 비자 초청장 PDF 발급(src/lib/pdf/VisaInvitationLetter.jsx) · 코디 비자 트래킹(/coordinator/visa)',
    '화면 경로': '/app/patient/visa, /app/visa',
    'DB 테이블/컬럼': 'education_contents (교육 콘텐츠) · visa_applications·visa_status_history (비자 진행)',
    '현재 구현 상태': '완료: /app/patient/visa, /app/visa, /app/coordinator/visa, /app/api/pdf/*\nmigrations/20260406_education_visa_rebooking (visa_applications 실데이터 1건)',
})

doc.add_page_break()

# ================================================================
# 9. 사후관리·모니터링·교육
# ================================================================
add_heading(doc, '9. 사후관리·모니터링·교육', 1)

add_heading(doc, '9.1 경과 모니터링 f/u (FN-POST-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-POST-01',
    '기능명': '귀국 후 환자 경과 추적 및 모니터링',
    '액터': '환자 (P-01), 의료진 (P-03)',
    '전제 조건': '치료 완료 및 귀국',
    '트리거': '정기 체크인 알림 수신 또는 환자 자발적 입력',
    '기본 흐름': '① 환자가 증상 기록 입력 — 로그인 화면(/patient/symptoms) 또는 «계정 없이» 진행상황 링크(/claim/<토큰>, 2026-09-06). 해외 의료기관·환자는 검사결과·영상 업로드\n② 위험도 판정 = 규칙(응급 키워드 + 점수식) → 제미나이 2차 판정(src/lib/followup/aiTriage.ts, 2026-09-06). AI 는 올리기만 하고 확신 50% 미만은 무시, 실패·시간 초과면 규칙 결과 유지. 근거는 환자 원문 인용으로 기록\n③ 위험 시 담당자 경보(symptom_alerts → 코디 「증상 알림」) + 기록은 케이스 상세 「추가 정보」에도 남아 코디가 늘 보던 자리에서 본다\n④ 3일 이상 입력이 없으면 «침묵» 경보(detect-silent-patients 크론)\n⑤ 담당자가 화상 경과상담(follow_up 세션) 또는 메시지로 확인 — 완료 처리한 세션만 K-04 에 잡힌다',
    '예외 흐름': '• 상담 완료 처리 24시간 미실행 → 운영자 경보(unclosedNudge)\n• 심각한 이상 징후 → 응급 연락처 안내(24시간 답변 «약속»은 하지 않는다)',
    '입력': '증상 코드, 통증 점수, 파일, 메모',
    '출력': '경과 기록, 의료진 알림',
    '화면 경로': '/app/patient/symptoms',
    'DB 테이블/컬럼': 'followup_schedules (사후관리 차수·상태를 별도 표로 관리)\nconsultation_sessions (notes, notes_encrypted)',
    '현재 구현 상태': '완료: /app/patient/symptoms, /app/api/inquiries/claim/symptoms, /app/api/khidi/followup, /app/api/cron/detect-silent-patients, src/lib/followup/{symptomAnalyzer,aiTriage}.ts\n증상 보고 22건(시연·점검 포함) · 실환자 사후관리 일정은 아직 0(유치 확정 환자 없음). AI 2차 판정 실측(2026-09-06 시험 문의): 규칙 «중간 46%» → AI «높음 90%, 진통제 무효·38.5도 인용» → 담당자 확인으로 상향, 호출당 316 토큰(US$0.001)',
})

add_heading(doc, '9.2 건강관리 교육 콘텐츠 (FN-POST-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-POST-02',
    '기능명': '암 유형별 맞춤 사후관리 교육 콘텐츠',
    '액터': '환자 (P-01)',
    '트리거': '교육 콘텐츠 메뉴 접근 또는 치료 완료 후 자동 제공',
    '기본 흐름': '① 환자 암 유형에 따른 콘텐츠 필터링\n② 식이요법, 운동가이드, 복약안내, 면역력 관리 콘텐츠\n③ 러시아어·카자흐어 콘텐츠 제공\n④ 영상·카드뉴스·텍스트 형태',
    '화면 경로': '/app/education (공개 화면, 6개 언어) · /app/admin/education (관리) · /app/patient/education 은 /education 으로 보내는 문일 뿐',
    'DB 테이블/컬럼': 'education_contents (암종·단계·범주별 교육 콘텐츠, 다국어)\nmigrations/20260406_education_visa_rebooking',
    '현재 구현 상태': '완료: /app/education, /app/admin/education, src/lib/followup/educationEngine.ts\n콘텐츠 18건(암 5종·6개 언어). 단계별 자동 발송은 /app/api/cron/dispatch-surveys 에 결합돼 D+7·14·30·90·180 에 환자 언어로 나간다(2026-08-25 연결, reminders_scheduled 로 중복 방지). 실환자 발송 이력은 아직 0',
})

add_heading(doc, '9.3 재방문 예약 (Rebooking) (FN-POST-03)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-POST-03',
    '기능명': '경과 기반 재방문 예약 원스톱 서비스',
    '액터': '환자 (P-01), AI Agent',
    '트리거': '경과 모니터링 데이터에서 재방문 필요성 감지',
    '기본 흐름': '① 시스템 제안: 사후관리 차수(화상 경과상담 도래)·증상 경보·의료진 권고에서 재방문 필요성을 판정(rebookingEngine)해 환자 화면에 «제안»으로\n② 환자가 «먼저» 요청(2026-09-06): 로그인 화면(/patient/rebooking) 또는 진행상황 링크(/claim)에서 [재진 상담 요청] 한 번 — 병원 가용 일정은 시스템이 모르므로 시간은 코디가 잡는다\n③ 요청·확정 즉시 코디·관리자 종 + 메일 → 코디가 「상담 일정」에서 환자·시각을 넣어 초대 링크 발송(기존 흐름 재이용)\n④ 재진 예상비용·비자 재발급 안내(C-3-3/G-1-10 체크리스트)는 같은 화면에\n⑤ 6시간 안 중복 요청은 새 행 없이 접수 처리(연타 방지)',
    '화면 경로': '/app/patient/rebooking · /app/claim/[token] · /app/api/portal/followup (GET·POST·PATCH) · /app/api/inquiries/claim/rebooking',
    'DB 테이블/컬럼': 'followup_schedules (status pending|proposed|confirmed|dismissed, schedule.source = followup|symptom|doctor|patient_request) · inquiries.follow_ups(태그 글)',
    '현재 구현 상태': '완료: 2026-09-06 이전엔 [확정]이 DB 상태만 바꾸고 아무에게도 안 알렸고 환자가 먼저 요청할 단추가 없었다(막다른 길). 지금은 요청·확정 모두 코디에게 닿는다. 로컬 실측(시험 문의 #308): 요청 행 생성·추가 정보 태그·2.3초\nmigrations/20260406_education_visa_rebooking',
})

doc.add_page_break()

# ================================================================
# 10. 다국어
# ================================================================
add_heading(doc, '10. 다국어 (ko/en/ru/kz/zh/ja)', 1)

add_heading(doc, '10.1 다국어 UI 라우팅 (FN-I18N-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-I18N-01',
    '기능명': '6개 언어 UI 및 언어별 라우팅',
    '액터': '모든 사용자',
    '트리거': '언어 선택 또는 브라우저 언어 감지',
    '기본 흐름': '① 브라우저 Accept-Language 헤더 감지\n② 또는 /ru, /kz 경로로 직접 접근\n③ src/lib/i18n/index.js에서 언어 메시지 로드\n④ 해당 언어 UI 렌더링',
    '지원 언어': 'ko (한국어), en (영어), ru (러시아어), kz (카자흐어), zh (중국어), ja (일본어)',
    '화면 경로': 'proxy.ts (언어 접두어 처리) · src/lib/i18n/index.js',
    '현재 구현 상태': '완료: 6개 언어 라우트 운영(ko·en·ru·kz·zh·ja)\n러시아어·카자흐어 문구 전건 채움\n키 누락은 자동 검사로 상시 확인',
})

add_heading(doc, '10.2 DB 콘텐츠 자동번역 (FN-I18N-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-I18N-02',
    '기능명': '병원·치료 정보 DB 자동 다국어 번역',
    '액터': '관리자, 시스템 배치',
    '트리거': '신규 병원·치료 정보 등록 시 또는 관리자 수동 실행',
    '기본 흐름': '① 관리자가 병원 정보를 저장하면 triggerMultiLangTranslation(src/lib/translate.ts)이 /api/translate 로 6개 언어 번역\n② hospitals.i18n JSONB 에 저장, 일괄은 scripts/batch-translate-all.mjs\n③ 코디는 콘텐츠 편집기(/coordinator/content)에서 화면이 실제로 그리는 칸을 6개 언어로 직접 고친다(2026-09-06 부터 치료법·5축·암종·FAQ·병원 소개·주소 249칸 편입)\n④ 각 언어 UI 에서 자동 로드. 빈 칸은 영어로 폴백되므로 «공개 병원 표의 번역 빈 칸» 훑기 검사가 매일 본다',
    '화면 경로': '/app/api/translate · /app/api/admin/hospitals · /app/coordinator/content',
    'DB 테이블/컬럼': 'hospitals.i18n JSONB\ntreatments.translations\nmigrations/20260223_auto_translate_fields\nmigrations/20260226_treatment_translations',
    '현재 구현 상태': '완료: migrations/20260223_auto_translate_fields, 20260226_treatment_translations, content_overrides(262행)·content_change_log\n외국어 병원명은 «Immune Hospital» 로 통일(2026-09-06 PO 결정)',
})

doc.add_page_break()

# ================================================================
# 11. 코디네이터 포털
# ================================================================
add_heading(doc, '11. 코디네이터 포털', 1)

add_heading(doc, '11.1 코디네이터 대시보드·환자 관리 (FN-COORD-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-COORD-01',
    '기능명': '코디네이터 포털: 환자 관리 및 상담 처리',
    '액터': '코디네이터 (P-02)',
    '전제 조건': 'coordinator role 로그인',
    '트리거': '/coordinator 접근',
    '기본 흐름': '① 받은함(전체 / 추가 정보 필요 / 매칭 준비 완료 / 종료)에서 문의 확인 — 「N일째 정체」·「환자 새 글」 배지\n② 케이스 상세: AI 케이스 브리프, 첨부 [한][EN][RU] 번역, 병원 의뢰서(원본 양식), 소견 요청 링크, 에이전시 공개\n③ 메시지(웹·텔레그램·왓츠앱·에이전시) 응답, 음성 메모는 읽어서 확인\n④ 상담 일정 생성·[상담 완료] 처리(실적 판정), 견적·비자 트래킹\n⑤ 결과 입력: 유치 확정 / 종료(환자가 안 온다고 함)\n⑥ 유치 전환·사후관리·만족도 현황 확인, 콘텐츠 편집, 개선 요청함',
    '화면 경로': '/app/coordinator/* — 메뉴 16개: 대시보드 · 음성 정리 · 문의함 · AI 상담 리드 · 의뢰·케이스/병원배정 · 상담 일정 · 유치 전환 · 사후관리·만족도 · 파트너 발굴 · 메시지 · 비자 트래킹 · 견적 · 증상 알림 · 콘텐츠 편집 · 개선 요청함 · 설정',
    'DB 테이블/컬럼': 'inquiries (status, outcome, is_test) · case_status_history · case_updates · inquiry_events\nconsultation_sessions (coordinator_id) · cost_estimates · visa_applications · voice_notes',
    '현재 구현 상태': '완료: /app/coordinator/* (layout.jsx 의 NAV 가 정본)',
})

add_heading(doc, '11.2 코디네이터 응답·문서 관리 (FN-COORD-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-COORD-02',
    '기능명': '코디네이터 응답 이력 및 문서 관리',
    '액터': '코디네이터 (P-02)',
    '기본 흐름': '① 케이스 단계 변경·소식은 case_status_history·case_updates 에 남고 환자 진행상황 화면(/claim)에 그대로 비친다\n② 의료문서 분류·병원 전달(case_shared_documents, 「환자에게 보이기」 스위치)\n③ 상담 요약·케이스 브리프 생성(AI 초안, «진단이 아니며 검수 필요» 고지)',
    'DB 테이블/컬럼': 'case_status_history (31행) · case_updates · case_shared_documents (9) · cotreatment_referrals (4)\n※ coordinator_responses 표는 있으나 0행 — 실제 응답은 chat_messages·case_updates 로 남는다',
    '현재 구현 상태': '완료: /app/coordinator/inbox/[id], /app/api/coordinator/*',
})

doc.add_page_break()

# ================================================================
# 12. 관리자·파트너·의사 포털
# ================================================================
add_heading(doc, '12. 관리자·국내 의료기관 포털', 1)

add_heading(doc, '12.1 관리자 포털 (FN-ADMIN-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-ADMIN-01',
    '기능명': '관리자 KPI·사용자·AI 성능 관리',
    '액터': '관리자 (P-04)',
    '전제 조건': 'admin role 로그인',
    '기본 흐름': '① 홈: 대시보드 · KHIDI 리포트(kpi-dashboard·north-star·conversion·satisfaction·evidence) · 문의 통계 · 광고 예산\n② 상담·문의: 문의 받은함 · 케이스 · AI 채팅 · Human Agent 채널 · 화상 상담 · 협진 의뢰 · 후속 리마인더 · 견적 · 비자 · 증상 알림\n③ 파트너·회원: 직원(코디) 계정 · 에이전시 · 클리닉 · 파트너 발굴 · 제휴 병원 · 병원 진료의뢰 · 환자 회원 · 데이터 삭제 요청\n④ 콘텐츠: 문구 편집기 · 환자 교육자료 · AI 지식베이스(RAG)\n⑤ AI 품질: 자가시험 · 품질 판정 · 모델 벤치마크 · 피드백 · 사용량\n⑥ 시스템: 개선 요청함 · 감사로그 · 외부 서비스 사용량 · 알림 관리 · 브랜딩 설정 (+ 메뉴에서 숨긴 화면 9종은 주소로만 열리며 check:dead-screens 가 매달 대조)\n⑦ 증빙·월간 보고 자료는 CSV 내려받기(PDF 출력은 없다)',
    '화면 경로': '/app/admin/* (7그룹, 정본 = app/admin/_components/AdminNav.jsx)',
    'DB 테이블/컬럼': 'admin_audit_logs (id, admin_user_id, admin_email, action, created_at)\nmigrations/20260129_add_admin_audit_logs',
    '현재 구현 상태': '완료: /app/admin/*',
})

add_heading(doc, '12.2 파트너 병원 포털 (FN-PARTNER-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-PARTNER-01',
    '기능명': '국내 의료기관 포털: 진료 의뢰(리드) 관리',
    '액터': '제휴 한국 병원 담당자 (owner / manager / viewer)',
    '전제 조건': 'hospital_users 표에 등록된 계정으로 로그인(별도 role 값은 없다)',
    '기본 흐름': '① [대시보드] 응답 대기 · 전환율 · 평균 첫 응답 · 확정 견적 합계 + 「응답 필요」 목록\n② [진료 의뢰(리드)] 목적·시술·국가·언어·배정일·문의 내용·보험 정보 유무(환자 실명·연락처는 넘기지 않는다)\n③ 상태 갱신: 전송됨 → 조회됨 → 응답함 → 치료 확정(→ inquiries.outcome=admitted 자동 반영) / 거절 / 만료, 견적 최소·최대·메모, 「코디에게 메시지」, CSV 내려받기\n④ (비활성) 「병원 정보」·「시술 카탈로그」는 기능 플래그 HOSPITAL_CONTENT_ENABLED=false 로 꺼져 있다(공개 연동 준비 중)\n⑤ (없음) 진료 결과 입력·가용 일정 화면. 의료진은 계정 없이 상담 초대링크로 참여한다',
    '화면 경로': '/app/hospital, /app/hospital/leads (profile·treatments 는 리디렉트)',
    'DB 테이블/컬럼': 'hospital_users (2행) · hospital_leads · hospitals (발행 8)\n※ partner_doctors·partner_branches 는 0행 — 면력한방병원 의료진 28명은 코드 데이터(src/lib/data/immuneDoctors.js)',
    '현재 구현 상태': '완료(리드 관리) / 비활성(프로필·카탈로그) / 없음(진료 결과·일정): /app/hospital/*, app/hospital/_components/featureFlags.js',
})

doc.add_page_break()

# ================================================================
# 13. 알림·이메일
# ================================================================
add_heading(doc, '13. 알림·이메일·푸시', 1)

add_heading(doc, '13.1 이메일 알림 (FN-NOTIF-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-NOTIF-01',
    '기능명': 'Resend 기반 이메일 알림·리마인더',
    '액터': '시스템, 환자, 코디네이터',
    '트리거': '예약 확정, 상담 완료, 경과 체크인 시점',
    '기본 흐름': '① 트리거 이벤트 발생(접수·상담 생성·리마인더·설문·교육·식은 문의)\n② 발송 지점(크론·API)이 src/lib/email/sendEmail.ts 호출\n③ 템플릿은 src/lib/email/templates/*.ts 의 TS 함수(react-email 아님), 6개 언어\n④ Resend 로 발송, 실패 시 AWS SES 폴백, 둘 다 없으면 ok:false\n⑤ 발송 결과는 호출부 로그에만 남는다 — admin_notification_logs 는 관리자 «테스트 발송»만 기록(부록 B)',
    '이메일 유형': '접수 확인(시험 문의엔 안 보냄), 상담 초대(.ics), 리마인더, 만족도 설문, 교육 콘텐츠, 관리자 신규 문의·환자 새 글·식은 문의 알림',
    '화면 경로': 'src/lib/email/* · 발신 도메인 healwith.co.kr (SPF·DKIM·DMARC p=none 등록)',
    'DB 테이블/컬럼': 'reminders_scheduled (예약 발송 큐) · admin_notification_logs (테스트 발송) · admin_notification_recipients\nmigrations/20260204_add_admin_notification_logs',
    '현재 구현 상태': '완료: src/lib/email/sendEmail.ts, src/lib/email/templates/*, /app/api/cron/dispatch-*',
})

add_heading(doc, '13.2 실시간 In-app 알림 (FN-NOTIF-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-NOTIF-02',
    '기능명': '인앱 알림(벨) + 스토어 앱 푸시(FCM)',
    '액터': '코디네이터, 관리자, 환자',
    '트리거': '신규 문의, 환자 새 글, 소견 도착, 증상 경보, 식은 문의, 상담 임박',
    '기본 흐름': '① 서버가 notifications 표에 행을 넣는다(src/lib/notifications/inApp.ts)\n② 벨(src/components/NotificationBell.jsx)이 30초 폴링으로 읽는다 — Supabase Realtime 이 아니다(Realtime 은 상담방 메시지에만 쓴다)\n③ 앱이 기기를 등록했으면(device_tokens 23건) pushBridge 가 FCM 푸시로도 보낸다(옵트인 배너·정책 pushPolicy.ts)\n④ 알림이 가리키는 화면이 그 값을 실제로 읽는지 check:deeplinks 가 CI 에서 본다',
    '화면 경로': '/app/notifications · 벨은 환자·직원 상단바 공용',
    '현재 구현 상태': '완료: src/lib/notifications/{inApp,pushBridge,pushPolicy}.ts, src/lib/push/{fcm,registerPush,buildPushMessage}.ts, /app/api/push/*\n(옛 문서가 적던 src/hooks/useNotifications.ts 는 어디서도 안 쓰던 죽은 코드라 2026-09-06 삭제)',
})

doc.add_page_break()

# ================================================================
# 14. AI RAG 파이프라인
# ================================================================
add_heading(doc, '14. AI 학습 파이프라인 (RAG 3계층)', 1)

add_heading(doc, '14.1 RAG 3계층 구조 (FN-RAG-01)', 2)
add_para(doc, '사업계획서 p.30~31 AI 학습 기반 상담 자동화 시스템 구현.')
add_scenario(doc, {
    '기능 ID': 'FN-RAG-01',
    '기능명': 'RAG 3계층 기반 AI 응답 생성',
    '개요': '사업계획서(p.31) 「지금은 사람이, 앞으로는 AI 가」: Human Agent 상담 데이터 RAG 학습 파이프라인',
    '1계층 (자체 DB)': 'Supabase pgvector 로 병원·치료·FAQ 벡터 저장 — rag_documents 63건 / rag_chunks 105조각\n/api/rag/ingest 로 코드 수정 없이 올린다\nmigrations/20260225_rag_vector_v1',
    '2계층 (외부 보조 검색)': '답할 때 HIRA 공공 API + 네이버 지역검색을 실시간 병렬 호출(src/lib/chat/externalSearch.ts, 짧은 타임아웃)\n별도로 크롤 파이프라인(hira·google-places·kakao-local·naver-local)이 있으나 crawl_jobs 0행 — 지금은 안 돈다',
    '3계층 (구글 검색 그라운딩)': '설계엔 있으나 «미작동»: 코드가 넘기는 옵션(useSearchGrounding)이 설치된 @ai-sdk/google 에 없어 조용히 무시된다. 2026-07-31 가짜 「웹 검색 결과」 라벨을 제거했다. 살리려면 정식 googleSearch 도구로(부록 B)',
    '학습 파이프라인': '① Human Agent 상담 기록\n② 자동 구조화 (JSON/CSV)\n③ 벡터 임베딩 생성\n④ pgvector 저장\n⑤ RAG 검색으로 AI 응답 품질 향상',
    '화면 경로': 'src/lib/rag/*, src/lib/chat/generateReply.ts',
    'DB 테이블/컬럼': 'rag_documents (id, content, trust_tier, source_type, source_url) · rag_chunks (embedding)\nplaybook_patterns, playbook_usage_events\nmigrations/20260225_rag_*, 20260225_playbook_*',
    '현재 구현 상태': '완료(1·2계층) / 미작동(3계층): src/lib/rag/*, src/lib/chat/externalSearch.ts\nmigrations/20260225_rag_vector_v1, 20260225_playbook_patterns\n자가시험 50건에서 RAG 조각이 실제로 프롬프트에 들어가는지(rag_chunk_count) 매회 기록',
})

doc.add_page_break()

# ================================================================
# 15. 보안·암호화·감사
# ================================================================
add_heading(doc, '15. 보안·암호화·감사', 1)

add_heading(doc, '15.1 환자 PII AES-256-GCM 암호화 (FN-SEC-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-SEC-01',
    '기능명': '환자 PII 암호화 저장',
    '대상 컬럼': 'inquiries.first_name · last_name · email · phone (암호문 저장)\nconsultation_sessions.notes_encrypted · cancer_patient_intakes.*_encrypted 등',
    '암호화 방식': 'AES-256-GCM (src/lib/security/encryptionV2.ts)',
    '암호화 시점': '접수 폼 제출 → 서버 측에서 암호화 후 저장(값이 있는 행 전부 암호문, 평문 0건 — 2026-08-20 운영DB 실측)\n칸 이름은 평문 시절 그대로(first_name 등)이며 «평문 칸 삭제» 마이그레이션은 실행 금지 — 돌리면 26~29건 소실. 정리는 별도 과제',
    '복호화 권한': 'service_role 키 보유 서버 모듈만 복호화 가능\nimport "server-only" 적용',
    'DB 테이블/컬럼': 'inquiries (first_name, last_name, email, phone — 전부 암호문) · cancer_patient_intakes.*_encrypted · consultation_sessions.*_encrypted · consultation_translations.*_encrypted · voice_notes(2026-09-04) — 암호화 칸 20개\n접근 기록: inquiries.access_log(접수·첨부 시점 IP)',
    '현재 구현 상태': '완료: src/lib/security/encryptionV2.ts (AES-256-GCM)\n회전 금지 키: SUPABASE_ENCRYPTION_KEY · ENCRYPTION_KEY_V1',
})

add_heading(doc, '15.2 API 보안 헬퍼 및 Rate Limiting (FN-SEC-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-SEC-02',
    '기능명': 'API 인증·권한 검증 및 Rate Limiting',
    '인증 헬퍼': 'requireAdminAuth(): admin role 검증\nrequireConsultationAccess(): 상담 접근 권한\n모든 새 API 라우트에 의무 적용',
    'Rate Limiting': 'src/lib/rateLimit.ts: IP 기반 요청 수 제한\n공개 POST 엔드포인트 필수 적용',
    '오류 처리': 'API 응답에 error.message 직접 노출 금지\n"internal_error" 등 코드형 오류만 반환',
    'RLS (Row Level Security)': 'Supabase PostgreSQL 행 단위 접근 제어\nmigrations/20260125_security_rls_policies\nmigrations/20260130_enable_rls_inquiries',
    '화면 경로': 'src/lib/auth/*, src/lib/rateLimit.ts',
    '현재 구현 상태': '완료: src/lib/security/*\nsrc/lib/rateLimit.ts\nmigrations/20260125_security_*',
})

add_heading(doc, '15.3 감사 로그 (FN-SEC-03)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-SEC-03',
    '기능명': '관리자 행동 감사 로그 기록',
    '기록 대상': '관리자 로그인, 사용자 역할 변경, 데이터 수정, 병원 정보 변경 등',
    '화면 경로': '/app/admin/audit',
    'DB 테이블/컬럼': 'admin_audit_logs (id, action, inquiry_ids, admin_user_id, ip_address, user_agent, created_at)\nmigrations/20260129_add_admin_audit_logs',
    '현재 구현 상태': '완료: migrations/20260129_add_admin_audit_logs',
})


doc.add_page_break()

# ================================================================
# 16. 해외 파트너 포털
# ================================================================
add_heading(doc, '16. 해외 파트너 포털 (에이전시·해외 의료기관)', 1)
add_para(doc, '해외 의뢰처를 «에이전시»와 «의료기관» 두 계층으로 갈랐다. 공고의 사후관리 3종(경과·검사결과·영상)은 임상 정보라 비의료 에이전시는 올릴 수 없기 때문이다(변경관리 CR-S-06·08).')
add_heading(doc, '16.1 해외 에이전시 포털 (FN-PARTNER-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-PARTNER-02',
    '기능명': '해외 에이전시: 환자 의뢰·진행 확인·케이스 메신저',
    '액터': '해외 유치 파트너 담당자 (agency_users + agencies.partner_type)',
    '전제 조건': '관리자가 [에이전시·클리닉] 메뉴에서 발급한 계정(상대 업무메일이 ID). 자율 가입 없음',
    '기본 흐름': '① /agency 로그인(proxy 가 세션 확인, API 는 checkAgencyAuth 관문)\n② 환자 의뢰(refer) — 환자 대신 문의를 만든다\n③ 케이스별 단계 진행 확인(환자 진행상황과 같은 6단계)\n④ 코디네이터와 케이스 메신저(chat_threads.channel=agency)\n⑤ 코디가 [에이전시에 공개]한 소견 번역본 열람',
    '화면 경로': '/app/agency (PartnerPortal.jsx) · /app/api/agency/{cases,inquiries,refer,translate}',
    'DB 테이블/컬럼': 'agencies (3) · agency_users (3) · inquiries (source) · chat_threads',
    '현재 구현 상태': '완료: 언어 전환·케이스 메신저 포함(E2E agency-portal.spec.ts)',
})
add_heading(doc, '16.2 해외 의료기관 포털 (FN-PARTNER-03)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-PARTNER-03',
    '기능명': '해외 의료기관: 검사결과·영상·소견 업로드 (사후관리 ICT ④)',
    '액터': '환자를 한국으로 의뢰하는 현지 병원 담당자',
    '기본 흐름': '① /clinic 로그인(에이전시와 같은 표, partner_type 으로 구분)\n② 에이전시 기능 전부 + 검사결과·영상·소견을 케이스에 올린다(progress_records)\n③ 업로드 시 케이스 이력에 경과 사건이 남아 코디·의뢰처 진행 화면에 자동 반영\n④ 업로드 채널은 앱·웹·메신저 전부 허용',
    '화면 경로': '/app/clinic · /app/api/agency/* (유형 검증은 공통 관문)',
    'DB 테이블/컬럼': 'progress_records (3, 견본) · case_shared_documents',
    '현재 구현 상태': '완료: 업로드→저장→담당자 열람 전 구간 견본으로 동작 확인(2026-08-25). 실환자 데이터 0(치료 후 단계 미도달)',
})

doc.add_page_break()

# ================================================================
# 17. 전문의 소견 · 환자 진행상황 · 병원 의뢰서
# ================================================================
add_heading(doc, '17. 전문의 소견 · 환자 진행상황 · 병원 의뢰서', 1)
add_heading(doc, '17.1 전문의 소견 요청·작성 (FN-OPINION-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-OPINION-01',
    '기능명': '계정 없는 협진 전문의의 소견 작성 → 환자 언어 번역 → 공개',
    '액터': '코디네이터, 협진 전문의(계정 없음), 환자',
    '기본 흐름': '① 코디가 케이스 상세에서 [소견 요청 링크 만들기] → 전문의에게 링크 전달(opinion_requests)\n② 전문의가 /opinion/<토큰> 에서 환자·임상 정보와 검사지를 보고 소견을 적어 제출(서류 첨부 선택, CT 초견은 AI 초안 — 판독 아님)\n③ 도착한 소견은 원문 그대로 케이스 상세에 쌓인다. 이미 카톡·메일로 받은 소견은 [이미 받은 소견 직접 입력]\n④ AI 가 환자 언어 번역 초안 → 코디 교정 → [에이전시에 공개] / 환자에게는 «공식 문서»로만 전달(화면에 다시 그린 요약은 2026-08-18 폐기)\n⑤ 환자 전달 시점(released_at)이 K-02 「글로 전달한 사전상담」의 판정 기준',
    '화면 경로': '/app/opinion/[token] · /app/coordinator/inbox/[id] 「전문의 소견」 · /app/api/opinions/[token]/{route,upload,imaging,translate,page}',
    'DB 테이블/컬럼': 'opinion_requests (3) · case_opinions (6, 전건 전달) · attachment_translations (23)',
    '현재 구현 상태': '완료: src/lib/opinions/*, src/lib/documents/translateDoc.ts',
})
add_heading(doc, '17.2 환자 진행상황 화면 (FN-CLAIM-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-CLAIM-01',
    '기능명': '가입 없이 링크로 보는 진행상황·소견·서류 (P-11)',
    '액터': '접수한 사람 누구나(환자·가족·에이전시)',
    '전제 조건': '접수 시 그 채널로 돌려준 /claim/<토큰> 링크',
    '기본 흐름': '① 6단계 막대(누를 수 있는 탭) — 아직 안 온 단계는 잠근다\n② 고른 단계에 있었던 일(코디 소식, 새것부터) + 「다음은요」\n③ 우리가 준 것(공식 문서) / 환자가 준 것(추가 자료 200MB·메시지)\n④ 증상 기록 카드(2026-09-06): 심각도 1~10 + 자유 서술 → 규칙+AI 판정 → 코디 종 알림, 「내 암종 관리 가이드」 링크(/education?cancer=…)\n⑤ 환자 글·증상은 코디 목록에 「환자 새 글 · N일째 안 읽음」 배지로 뜬다\n⑥ 아래 띠에서 계정 연결·가입을 «권유»만 한다(어떤 상태여도 위 진행상황은 남는다)',
    '언어': '사람이 고른 언어 → 문의서의 환자 언어 → 브라우저 → 영어. 국적으로 추측하지 않는다',
    '화면 경로': '/app/claim/[token] · /app/api/inquiries/claim/*',
    'DB 테이블/컬럼': 'case_status_history (31) · case_shared_documents (9) · inquiry_events',
    '현재 구현 상태': '완료: 2026-08-04 신설, 2026-09-05 환자 글 알림·배지',
})
add_heading(doc, '17.3 병원 의뢰서·법정 문서 발급 (FN-DOC-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-DOC-01',
    '기능명': '협진 병원 원본 양식 의뢰서(DOCX) + 법정 문서 PDF 3종',
    '액터': '코디네이터',
    '기본 흐름': '① 케이스 상세에서 병원이 준 워드 양식 «원본 그대로»를 XML 로 채워 화면·인쇄·파일이 한 원본에서 나온다(referral-docx). 국적은 여권 외 서류에서도 읽고 영문 양식엔 성별·국적도 영어로\n② 자료 자동채움(referral-fill) — 접수 자료·음성 메모 판독 결과에서\n③ PDF: 동의서 3종(ConsentForms) · 견적서(MedicalQuotation, 유치수수료 상한 15% 를 기계가 막는다) · 비자 초청장(VisaInvitationLetter). 한글·키릴 글꼴은 자체 호스팅, check:pdf-tone 이 옛 톤 복귀를 차단',
    '화면 경로': '/app/api/coordinator/inquiries/[id]/{referral-docx,referral-fill} · /app/api/pdf/{consent/[form],quotation} · /app/inquiry/referral',
    'DB 테이블/컬럼': 'cotreatment_referrals (4) · cost_estimates (7) · cost_estimate_history (12)',
    '현재 구현 상태': '완료: src/lib/referral/*, src/lib/documents/*, src/lib/pdf/*',
})

doc.add_page_break()

# ================================================================
# 18. 만족도 설문 · 사후관리 자동 실행
# ================================================================
add_heading(doc, '18. 만족도 설문 · 사후관리 자동 실행', 1)
add_heading(doc, '18.1 만족도 설문 (FN-SURVEY-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-SURVEY-01',
    '기능명': '5문항 만족도 설문 자동 발송·응답·집계 (K-03 의 유일한 원천)',
    '액터': '시스템, 환자, 코디네이터, 관리자',
    '트리거': '코디가 상담을 [상담 완료] 처리한 뒤 24시간(소급 14일)',
    '기본 흐름': '① dispatch-surveys 크론(매일 09:00 UTC)이 발송 대상 세션을 찾아 surveys 행 생성 + 메일(환자 언어)\n② 환자가 /survey/<토큰> 에서 5문항(1~5점) 응답 → survey_responses\n③ 코디 /coordinator/satisfaction · 관리자 /admin/khidi/satisfaction 에서 평균(100점 환산)·응답률\n④ 표본부족 가드 SATISFACTION_MIN_RESPONSES(기본 0)\n⑤ 시험 문의·시험 방에 딸린 설문은 집계에서 제외(fetchTestSurveyIds)',
    '예외 흐름': '• 발송 조건 시간창이 실행 주기보다 짧아 대부분 안 나가던 결함은 2026-06 수정(CR-O-04)\n• 실환자 응답은 아직 0건 — 발송 조건이 코디의 [상담 완료] 클릭에 걸려 있다',
    '화면 경로': '/app/survey/[token] · /app/api/survey/* · src/lib/surveys/*',
    'DB 테이블/컬럼': 'surveys (4) · survey_responses (2, 전부 시험·시연)',
    '현재 구현 상태': '완료(E2E patient-survey-response.spec.ts) / 실환자 표본 0',
})
add_heading(doc, '18.2 «방문 전» 사후관리 케이던스 (FN-POST-05) — 2026-09-06 신설', 2)
add_scenario(doc, {
    '기능 ID': 'FN-POST-05',
    '기능명': '소견을 받고 아직 오지 않은 환자에게 D+3 · D+14 · D+30 안부·다음 단계 안내 + 무응답 코디 알림',
    '왜': '사후관리 ④⑤⑥은 전부 «치료가 끝난 뒤»에만 시작된다. 실환자 8명 중 소견까지 받은 6명이 한 명도 오지 않았고 그 6명에게 플랫폼은 소견 전달 뒤 아무것도 하지 않았다(첫 실고객이 두 달째 「상담·검토 진행」). PO 지시(2026-09-06)로 «방문 전» 구간에 케이던스를 붙였다',
    '액터': '시스템, 환자, 코디네이터',
    '트리거': '소견 전달일(case_opinions.released_at 최신) 기준 D+3 / D+14 / D+30 — dispatch-surveys 크론(매일 09:00 UTC) 안에서 돈다',
    '기본 흐름': '① 대상 = 소견이 전달됐고 결과(outcome)가 없고 치료 단계 전인 문의(시험 문의 제외)\n② D+3 「소견 잘 받으셨나요」 / D+14 「다음 단계(비용·일정·비자)를 함께 정할까요」 / D+30 「요즘 어떻게 지내시나요」 — 환자 언어 6종, 진행상황 링크로 회신 유도\n③ D+14·D+30 무응답이면 코디·관리자 종 + 메일(「소견 뒤 N일째 무응답」)\n④ 환자가 앵커 이후 글·증상을 남겼으면 독촉(D+14·D+30)은 보내지 않는다 — 코디가 이어간다\n⑤ 한 실행에 케이스당 최대 1통, 도래한 지 21일 넘긴 단계는 보내지 않는다(도입 시점 옛 케이스 소급 발송 방지). 멱등 키 = reminders_scheduled(pre_visit_followup)',
    '예외 흐름': '• 메일 주소가 없는 소급 등록 케이스 → 환자에게는 못 보내고 코디 알림만\n• 발송 실패 → 기록을 남기지 않아 다음 실행이 재시도\n• 끄기 = env PRE_VISIT_FOLLOWUP_ENABLED=0',
    '화면 경로': 'src/lib/followup/preVisitFollowup.ts · src/lib/email/templates/preVisitFollowup.ts · /app/api/cron/dispatch-surveys',
    'DB 테이블/컬럼': 'case_opinions.released_at(앵커) · reminders_scheduled(reminder_type=pre_visit_followup, payload.phase, status sent|skipped) · inquiries.follow_ups(응답 판정)',
    '현재 구현 상태': '완료(단위 시험 8건) · **발송은 PO 보류로 기본 꺼짐**(2026-09-06 밤 「일단 멈춰」 — 실환자에게 자동 메일이 나가는 것을 보고하자 보류). 켜면 첫 실행에 소견 D+14 구간 실환자 2명에게 각 1통. 판정 로직·기록은 그대로 살아 있다',
})

add_heading(doc, '18.3 치료 후 사후관리 자동 케이던스 (FN-POST-04)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-POST-04',
    '기능명': '치료 후 1주·2주·1개월·3개월·6개월·1년 차수 자동 실행',
    '액터': '시스템, 환자, 코디네이터',
    '기본 흐름': '① 케이스가 사후관리 단계에 들어가면 followup_schedules 에 차수가 생긴다\n② 차수 도래 시 경과 설문·복약 확인·화상 경과상담·검사 리뷰 제안을 환자 앱에 제시(암종별 가감 — 위암 식이, 유방암 혈액검사 등)\n③ 같은 시점에 단계별 교육 콘텐츠를 환자 언어로 발송(D+7·14·30·90·180, reminders_scheduled 로 중복 방지)\n④ 증상 무입력 3일 → 침묵 경보. 재방문 필요 판정(rebookingEngine)이 나오면 재진 예약(/patient/rebooking)·견적·비자 절차를 ③의 기능으로 재이용',
    '화면 경로': '/app/api/cron/dispatch-surveys · /app/api/cron/detect-silent-patients · /app/patient/{symptoms,rebooking} · /app/coordinator/satisfaction(사후관리 탭)',
    'DB 테이블/컬럼': 'followup_schedules (2) · symptom_reports (22) · symptom_alerts (6) · progress_records (3)',
    '현재 구현 상태': '완료(견본으로 전 구간 동작 확인) / 실환자 0 — 유치 확정 환자가 생겨야 실적이 생기는 구조',
})

doc.add_page_break()

# ================================================================
# 19. 모바일 앱 · 음성 메모 판독
# ================================================================
add_heading(doc, '19. 모바일 앱(스토어 판) · 음성 메모 판독', 1)
add_heading(doc, '19.1 스토어 앱 (FN-APP-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-APP-01',
    '기능명': 'Capacitor 기반 iOS/Android 앱 — 웹을 그대로 싣고 푸시·딥링크·네이티브 로그인만 덧댄다',
    '액터': '환자, 코디네이터',
    '기본 흐름': '① capacitor.config.ts 가 실서비스 주소를 라이브로 싣는다(정적 파일을 안 싣는다 — webDir 에 사진 74MB 가 실리던 함정 2026-08-31 제거, 82.7MB→8MB)\n② 앱 설치 안내 화면 /app, 구판 감지 「업데이트해 주세요」 띠, 스플래시, 안드로이드 뒤로가기, 딥링크(.well-known/assetlinks.json·apple-app-site-association)\n③ 네이티브 구글·애플 로그인, FCM 푸시 기기 등록, 광고 ID 권한 없음(판 14)\n④ 계정 삭제 화면 /account-deletion(스토어 요구) — 세 곳 동시 삭제\n⑤ AI 로 «무엇이·누구에게» 가는지 대화창에 상시 고지(애플 5.1.1(i)·5.1.2(i) 대응)',
    '빌드': 'Android 는 로컬/리눅스 gradle, iOS 는 Codemagic(맥 빌드). 부품 설정은 capacitor.config.ts 에만(gradle.properties 에 적으면 조용히 안 먹는다)',
    '화면 경로': '/app/app · android/ · ios/ · codemagic.yaml · src/lib/app/*',
    '현재 구현 상태': '완료 / 스토어: 구글 프로덕션 검토 제출(2026-09-02, 판 14) · 애플 4번째 심사 대기(2026-09-03 재제출). 앱 안 애플 로그인은 아직 끝까지 못 간다',
})
add_heading(doc, '19.2 음성 메모 판독 (FN-VOICE-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-VOICE-01',
    '기능명': '왓츠앱·텔레그램 음성 메모를 코디가 듣지 않고 읽는다',
    '액터': '코디네이터',
    '기본 흐름': '① 문의 2단계 또는 케이스 상세에 음성(mp3·m4a·ogg·opus·amr 등)·텍스트 메모 첨부\n② AI 가 받아쓰고 한국어로 옮겨 [음성 정리] 화면에 정리(문의를 만들기 전에도 가능, 2026-09-04)\n③ 여권이 없어도 메모에서 연락처·국적을 채운다(자료 자동채움)',
    '화면 경로': '/app/coordinator/voice · /app/api/coordinator/voice-notes',
    'DB 테이블/컬럼': 'voice_notes (*_encrypted, migrations/20260904_voice_notes.sql)',
    '현재 구현 상태': '완료: 2026-09-02~04 신설',
})

doc.add_page_break()

# ================================================================
# 20. 콘텐츠 편집기 · 번역 품질 · 검색 유입
# ================================================================
add_heading(doc, '20. 콘텐츠 편집기 · 번역 품질 · 검색 유입', 1)
add_heading(doc, '20.1 콘텐츠 편집기 (FN-CONTENT-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-CONTENT-01',
    '기능명': '화면 문구·콘텐츠를 코디가 6개 언어로 직접 고친다(개발 일정에 종속되지 않게)',
    '액터': '코디네이터, 관리자',
    '기본 흐름': '① 공개 화면 문구는 중앙 사전(src/lib/i18n, 2,052키)에서 나오고 하드코딩은 자동 검사가 막는다\n② 코디가 /coordinator/content 에서 문구·치료법 카드·5축(ITCRN)·암종 상세·FAQ·수술 후 관리·제휴 병원 소개·주소를 언어별로 고친다(content_overrides 262행, 이력 content_change_log)\n③ AI 번역으로 채운 칸은 원어민 검수 전 「제안」 상태 표시(docs/rules/I18N_QUALITY.md §3)\n④ CI: ru·kz 키 누락 차단, 번역 품질(사실 유실·언어 섞임·용어 흔들림) 차단, 암종 콘텐츠 6개 언어 완성, 용어집 src/lib/i18n/glossary.js',
    '화면 경로': '/app/coordinator/content · /app/admin/settings/branding · scripts/check-i18n-*.mjs',
    '현재 구현 상태': '완료: 2026-09-06 콘텐츠 칸 249개 편입, 외국어 병원명 «Immune Hospital» 통일',
})
add_heading(doc, '20.2 검색 유입 (FN-SEO-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-SEO-01',
    '기능명': '러시아·카자흐 검색엔진 유입 배관 — 사이트맵·hreflang·구조화 데이터·IndexNow',
    '액터': '시스템',
    '기본 흐름': '① 사이트맵 170쪽(6개 언어 × 28쪽 + 러/카 랜딩), hreflang·canonical·og:locale, JSON-LD(structuredData.js)\n② 공개 화면 내부 링크에 언어 접두어 강제(check:locale-links — 없으면 검색 로봇이 전부 영어로 튕겨 두 달간 미색인이었다)\n③ IndexNow 자동 제출 크론(빙·얀덱스·네이버 공용, 평일 변경분·월요일 전부; 구글은 안 받는다) — 첫 실행 2026-09-06 146건 202\n④ 암종 상세 6쪽이 링크 0개 «고아»였던 것 2026-08-31 해소',
    '화면 경로': 'src/lib/seo/{indexNow,structuredData} · /app/api/cron/indexnow · app/sitemap.* · public/bea9fc…d.txt(IndexNow 키 — 공개가 규약)',
    '현재 구현 상태': '완료(배관) / 효과는 2주 뒤 재측정. 백링크 0·빙 색인 0쪽은 코드로 못 고친다(부록 B)',
})

doc.add_page_break()

# ================================================================
# 21. AI 품질 보증 · 비용 계측
# ================================================================
add_heading(doc, '21. AI 품질 보증 · 비용 계측', 1)
add_heading(doc, '21.1 자가시험·품질 판사 (FN-AIQ-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-AIQ-01',
    '기능명': 'AI 응답 품질을 사람 눈이 아니라 기계가 상시로 잰다',
    '액터': '시스템, 관리자',
    '기본 흐름': '① 자가시험: eval/chat-cases.json 50문항을 실서비스와 «같은» 채팅 경로(streamChatReply)로 주 2회 실행, 통과율·평균 점수·첫 글자 시간·RAG 조각 수 저장(ai_regression_runs 2,512행). 시험 호출은 isRegressionTest 표식으로 판사·경보를 건너뛰고 비용을 분리\n② 판사(Judge): 실서비스 답변마다 환각·안전·관련성 점수(ai_response_evaluations 493행). 2026-08-31 «세션 사실» 칸을 못 봐 정답을 환각으로 찍던 결함 수정 — 8/24 이전 점수는 오염값\n③ 완치 주장 검사기(러·카·중·일·영 거절 문장 오탐 수정) · 지어낸 병원·연락처 없는 「접수됨」 차단 · 매일 Chat Smoke\n④ 경보-검토 폐루프: 검출된 235건이 아무도 안 보던 것 2026-08 알림 경로 연결',
    '실측': f"{F.AI_QUALITY['regression_latest']}\n{F.AI_QUALITY['regression_first_token']}\n{F.AI_QUALITY['judge_since_0824']}",
    '화면 경로': '/app/admin/khidi/{ai-regression,ai-quality,ai-feedback,model-benchmark,agent-analysis} · src/lib/chat/regressionRunner.ts',
    '현재 구현 상태': '완료: src/lib/chat/regressionRunner.realpath.test.ts 가 «실서비스 경로 이탈»을 막는다',
})
add_heading(doc, '21.2 AI 사용량·비용 계측 (FN-AIQ-02)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-AIQ-02',
    '기능명': '어느 기능이 얼마나 쓰는지 표면(surface)별로 기록',
    '기본 흐름': '① logAiUsage 가 호출마다 surface(public_chat·case_brief·doc_translate·consult_translate·consult_stt·judge·regression_*)·토큰·추정 비용·실제 모델판을 기록(ai_usage_events 3,269행)\n② 일일 상한(aiGuard) + 구글 콘솔 spend cap\n③ 화상 통역 라우트는 2026-09-01 까지 기록이 없었다 — 이제 잡힌다. 한 호출 4,000 토큰 초과 시 경고',
    '실측': '최근 30일 추정 US$23 — 자가시험이 약 1/3, 공개 챗 US$2.3, 문서·메모 번역 US$2.7. 비용의 88~99% 가 모델의 «생각» 토큰',
    '화면 경로': '/app/admin/khidi/usage · src/lib/ai/{usageLog,usagePricing,aiGuard}',
    '현재 구현 상태': '완료',
})

doc.add_page_break()

# ================================================================
# 22. 운영 자동화 · 감시
# ================================================================
add_heading(doc, '22. 운영 자동화(정기 실행) · 감시', 1)
add_para(doc, '사후관리·설문·감시·색인 제출은 사람이 잊으면 멈추는 일이라 전부 정기 실행에 걸었다. 정본은 vercel.json 의 crons 이며 Vercel 예약이 깨운다(깃허브 예약은 이 저장소에서 하루 7회만 돌아 시각이 중요한 일에는 못 쓴다).')
tbl = doc.add_table(rows=0, cols=3)
tbl.style = 'Table Grid'
add_header_row(tbl, ['경로', '주기(UTC)', '하는 일'])
for r in F.CRONS:
    add_data_row(tbl, r, bold_first=True)
doc.add_paragraph()
add_heading(doc, '22.1 감시 (FN-OPS-01)', 2)
add_scenario(doc, {
    '기능 ID': 'FN-OPS-01',
    '기능명': '오류 추적·생사 감시·운영 경보',
    '기본 흐름': '① Sentry(client·server·edge) — 개인정보 마스킹, 소스맵 업로드(2026-07-28 확인), 대시보드 /admin/observability\n② /api/health 가 DB 까지 찔러 상태·실서비스 커밋을 낸다 — UptimeRobot(5분, 외부) + 깃허브 uptime.yml(2회 연속 실패만 장애)\n③ 운영 경보: deadman(정기 실행이 안 돈 것)·operationalAlerts, 훑기 대장(npm run sweep) 매일 08:00 KST\n④ 배포 판정은 창구 시각이 아니라 /api/health 의 commit 으로',
    '화면 경로': 'sentry.*.config.js · instrumentation.ts · /app/api/health · src/lib/alerts/*',
    '현재 구현 상태': '완료',
})

# ================================================================
# 구현 상태 요약
# ================================================================
doc.add_page_break()
add_heading(doc, '부록 A. 기능 구현 상태 요약', 1)
add_para(doc, f'코드베이스 검증 기반 현황 ({F.AS_OF} 기준). 본문 기능 ID 42개를 그룹으로 묶은 표다.')
doc.add_paragraph()

sum_tbl = doc.add_table(rows=0, cols=4)
sum_tbl.style = 'Table Grid'
add_header_row(sum_tbl, ['기능 그룹', '구현 상태', '파일 경로', '비고'])
summary_data = [
    ('회원·인증·권한 (FN-AUTH-01~03)', '완료', '/app/signup, /app/login, proxy.ts', '웹+앱 네이티브 로그인 · 초대 토큰. 앱 안 애플 로그인은 심사 대응 중'),
    ('문의 접수·문서 (FN-INTAKE-01~02)', '완료', '/app/inquiry, /app/inquiry/intake, /app/api/attachments', '2단계 접수 · 넓은 첨부 형식 · 값 전건 암호문(칸 이름 정리는 남음)'),
    ('병원 매칭·정보 (FN-MATCH-01~02)', '완료', 'src/lib/chat, /app/hospitals, /app/treatments', 'RAG 1·2계층 가동, 3계층(구글 그라운딩) 미작동'),
    ('화상상담 (FN-VIDEO-01~02)', '완료', '/app/consultation/[id], /app/c/[code]', '초대링크 입장 · 캘린더·ICS · 통화시간 정직화'),
    ('AI 챗·사람 상담 (FN-CHAT-01~02)', '완료', '/app/api/public/chat/*, /app/coordinator/messages', '웹·텔레그램·왓츠앱 봇 · 3턴 승격'),
    ('번역·화상 자막 (FN-TRANS-01~02)', '완료', '/app/api/translate-text, [id]/stt, agents/live-translate', '실서비스 가동 · 자막 3,735건'),
    ('예약·비자 (FN-SCHED-01~02)', '부분', '/app/coordinator/consultations, /app/patient/visa', '병원 가용 일정 입력 화면 없음'),
    ('사후관리·교육 (FN-POST-01~05)', '완료', '/app/patient/symptoms, /app/claim/[token], /app/education, dispatch-surveys', '규칙+AI 2차 위험도 · 교육 자동 발송 · 방문 전 케이던스(9/06) · 계정 없이 증상 기록 · 치료 후 실환자 0'),
    ('다국어 (FN-I18N-01~02)', '완료', 'proxy.ts, src/lib/i18n, /app/coordinator/content', '2,052키 × 6개 언어 · 콘텐츠 칸 249개 편집 가능'),
    ('코디네이터 포털 (FN-COORD-01~02)', '완료', '/app/coordinator/* (메뉴 16)', '음성 정리 · 종료 사유 · 환자 새 글 배지'),
    ('관리자 포털 (FN-ADMIN-01)', '완료', '/app/admin/* (7그룹)', 'KHIDI 리포트 · AI 품질 · 데이터 삭제 요청'),
    ('국내 의료기관 포털 (FN-PARTNER-01)', '부분', '/app/hospital/*', '리드 관리 완료 · 프로필/카탈로그 플래그 비활성 · 진료결과/일정 없음'),
    ('해외 파트너 포털 (FN-PARTNER-02~03)', '완료', '/app/agency, /app/clinic', '에이전시/의료기관 2계층 · 경과 업로드'),
    ('알림·이메일·푸시 (FN-NOTIF-01~02)', '완료', 'src/lib/email, src/lib/notifications, src/lib/push', '인앱은 30초 폴링 · 일반 발송 로그 표 미기록'),
    ('RAG 파이프라인 (FN-RAG-01)', '완료', 'src/lib/rag/*, externalSearch.ts', '63문서/105조각 · HIRA+네이버 실시간'),
    ('보안·암호화·감사 (FN-SEC-01~03)', '완료', 'src/lib/security, src/lib/auth, /app/admin/audit', 'AES-256-GCM · RLS 전 표 · 요청 제한 공용 저장소'),
    ('소견·진행상황·의뢰서 (FN-OPINION-01, FN-CLAIM-01, FN-DOC-01)', '완료', '/app/opinion/[token], /app/claim/[token], referral-docx', '소견 6건 전달 · 원본 양식 의뢰서 · PDF 3종'),
    ('만족도 설문 (FN-SURVEY-01)', '완료', '/app/survey/[token], dispatch-surveys', '실환자 응답 0'),
    ('스토어 앱·음성 판독 (FN-APP-01, FN-VOICE-01)', '완료', 'capacitor.config.ts, /app/coordinator/voice', '구글 검토·애플 심사 대기'),
    ('콘텐츠 편집·검색 유입 (FN-CONTENT-01, FN-SEO-01)', '완료', '/app/coordinator/content, /app/api/cron/indexnow', '효과는 2주 뒤 재측정'),
    ('AI 품질·비용 (FN-AIQ-01~02)', '완료', 'regressionRunner.ts, /app/admin/khidi/ai-*', '자가시험 통과 94%(9/03) · 첫 글자 P95 5.33초(8/21)'),
    ('운영 자동화·감시 (FN-OPS-01)', '완료', 'vercel.json crons(11), sentry.*, /api/health', 'UptimeRobot 5분'),
]
for row in summary_data:
    add_data_row(sum_tbl, row, bold_first=True)

doc.add_paragraph()
add_para(doc, '구현 비율 집계 (기능 ID 42개 기준):')
add_para(doc, '완료: 40개 (95%). 착수 시(2026-04) 부분구현 6건 중 4건이 완료로 전환됐고, 8/20 이후 새 기능 12개(9/06 방문 전 케이던스 포함)가 더해졌다', indent=1)
add_para(doc, '부분구현: 2개 (5%). ①예약 — 병원 가용 일정 입력 화면 ②국내 의료기관 포털 — 프로필·카탈로그(플래그 비활성)·진료 결과·일정', indent=1)
add_para(doc, '미구현: 0개. 다만 «완료» 안에 숨은 남은 일은 부록 B 에 그대로 적었다', indent=1)
add_para(doc, '※ 2026-08-20 판은 「완료 21 / 부분 2」였고 부분 2건(교육 자동 발송·예약)은 8/25 와 9/04 에 각각 연결·보강됐다. 이번 판은 «완료» 판정을 더 엄격히 보아 국내 의료기관 포털을 부분으로 내렸다.')

# 마지막 「※」 한 줄만 다음 쪽으로 떨어져 한 줄짜리 쪽이 생겼다(2026-08-20 실측).
# 집계 네 줄을 한 덩어리로 묶어 같은 쪽에 남게 한다.
for _p in doc.paragraphs[-5:-1]:
    _p.paragraph_format.keep_with_next = True

# ================================================================
# 부록 B. 남은 일 · 기술 부채
# ================================================================
doc.add_page_break()
add_heading(doc, '부록 B. 남은 일 · 기술 부채', 1)
add_para(doc, '「완료」 뒤에 남아 있는 것을 향후 개발자가 바로 찾아가도록 적는다. 사실의 출처는 _facts.py GAPS 이며, 고치면 거기서 지운다.')
doc.add_paragraph()
gap_tbl = doc.add_table(rows=0, cols=3)
gap_tbl.style = 'Table Grid'
add_header_row(gap_tbl, ['영역', '무엇이 남았나', '어디를 보면 되나'])
for g in F.GAPS:
    add_data_row(gap_tbl, g, bold_first=True)
doc.add_paragraph()
add_para(doc, '※ 알려진 결함의 살아있는 목록은 docs/KNOWN_ISSUES.md, 사건 기록은 docs/POSTMORTEMS.md, 변경 이력은 변경관리_대장.md 다.', size=9)

# 저장
import os as _os
# 저장 위치는 «이 스크립트가 있는 폴더» 기준으로 잡는다.
# (전에는 특정 PC 의 절대경로가 박혀 있어 그 PC 밖에서는 재생성이 아예 불가능했다.)
out_path = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), '02_기능명세서.docx')
doc.save(out_path)
print(f'저장 완료: {out_path}')
print(f'총 단락 수: {len(doc.paragraphs)}')
