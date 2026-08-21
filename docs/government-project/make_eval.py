import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

def set_font(run, size=10, bold=False, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.append(rFonts)

def add_heading(doc, text, level=1, color=(0,70,127)):
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    size = 14 if level == 1 else (12 if level == 2 else 11)
    set_font(run, size=size, bold=True, color=color)
    return p

def add_para(doc, text, indent=0, size=10):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_header_row(table, cols, bg='00467F'):
    row = table.add_row()
    for i, col in enumerate(cols):
        if i < len(row.cells):
            cell = row.cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(col)
            set_font(run, size=9, bold=True, color=(255,255,255))
            set_cell_bg(cell, bg)

def add_data_row(table, cols, bold_first=False, bg=None):
    row = table.add_row()
    for i, cell in enumerate(row.cells):
        if i < len(cols):
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(cols[i]))
            b = bold_first and i == 0
            set_font(run, size=9, bold=b)
            if bg:
                set_cell_bg(cell, bg)

# ================================================================
# 표지
# ================================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('HEALO 플랫폼')
set_font(run, 18, True, (0,70,127))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('평가 매트릭스')
set_font(run, 22, True, (0,70,127))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026년 KHIDI ICT 기반 외국인환자 사전상담·사후관리 지원 사업')
set_font(run, 11, False, (80,80,80))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Evaluation Matrix  |  v1.1  |  2026.08.20')
set_font(run, 10, False, (120,120,120))

doc.add_paragraph()
doc.add_paragraph()
tbl_info = doc.add_table(rows=0, cols=2)
tbl_info.style = 'Table Grid'
for label, val in [
    ('문서 번호', 'HEALO-EVAL-2026-001'),
    ('작성 기관', '본로이 (Bonroi)'),
    ('작성일', '2026년 8월 20일'),
    ('평가 근거', '공고문 붙임1 평가기준 (공고문 p.4, p.9)'),
]:
    row = tbl_info.add_row()
    row.cells[0].text = ''
    r1 = row.cells[0].paragraphs[0].add_run(label)
    set_font(r1, 10, True, (0,70,127))
    row.cells[1].text = ''
    r2 = row.cells[1].paragraphs[0].add_run(val)
    set_font(r2, 10)

doc.add_page_break()

# ================================================================
# 1. 문서 목적
# ================================================================
add_heading(doc, '1. 문서 목적 및 평가 기준 출처', 1)
add_para(doc, '본 문서는 2026년 KHIDI 「ICT 기반 외국인환자 사전상담·사후관리 지원 사업」 공고문 붙임1 (p.9)에 명시된 평가기준에 대하여, HEALO 플랫폼의 대응 현황·증빙 산출물·추정 점수를 표 형태로 정리한 것이다. 사업 심사 준비, 자체 점검, Phase B 보강 우선순위 결정에 활용한다.')
doc.add_paragraph()
add_para(doc, '※ 평가 기준 출처: 공고문 붙임1 「평가기준」 (공고문 p.9, 총 100점 + 가산점 5점 = 최대 105점)')
add_para(doc, '※ 추정 점수는 내부 자체 평가 결과이며, 실제 평가위원회 점수와 다를 수 있음.')

doc.add_page_break()

# ================================================================
# 2. 평가항목별 세부 평가기준 (공고문 붙임1 원문 재현)
# ================================================================
add_heading(doc, '2. 평가항목 구조 (공고문 붙임1 원문 기준)', 1)

tbl_orig = doc.add_table(rows=0, cols=4)
tbl_orig.style = 'Table Grid'
add_header_row(tbl_orig, ['평가항목', '평가요소', '배점', '출처'])
orig_data = [
    ('사업계획의 타당성\n(20점)', '대상국가 선정 타당성\n(보건의료 현황, 사업 여건 등)', '5점', '공고문 p.9'),
    ('', '사업계획의 적정성\n(목적·성과목표, 업무범위, 일정 등)', '10점', '공고문 p.9'),
    ('', '사업비 사용계획 적정성\n(정부출연금 사용, 자기부담금 매칭 등)', '5점', '공고문 p.9'),
    ('사업화 성공가능성\n(35점)', '수익실현 가능성\n(시장 자체 수익 창출 여부)', '15점', '공고문 p.9'),
    ('', '사업 지속가능성\n(정부 지원 종료 후 자체 운영 마일스톤)', '10점', '공고문 p.9'),
    ('', '환자유치 기여 가능성\n(대상국 및 주변국 확대 가능성)', '10점', '공고문 p.9'),
    ('사업내용의 현실성\n(30점)', '사업내용의 구체성\n(ICT 서비스 이해도, 추진방안, 목표 달성 계획)', '15점', '공고문 p.9'),
    ('', '사업내용의 실현가능성\n(추진 체계, 위험요소 분석·대응)', '15점', '공고문 p.9'),
    ('수행기관 역량 및 의지\n(15점)', '목표설정의 도전성, 사업수행 전문성\n전략국 네트워크, 사전타당성 조사', '15점', '공고문 p.9'),
    ('가산점\n(최대 5점)', '외국인환자 유치의료기관 평가인증', '3점', '공고문 p.4, p.9'),
    ('', '외국인환자 유치사업자 참여', '2점', '공고문 p.4, p.9'),
    ('', '비수도권 의료기관 참여 (서울/경기/인천 외)', '1점', '공고문 p.4, p.9'),
    ('합계', '', '최대 105점', ''),
]
for r in orig_data:
    add_data_row(tbl_orig, r, bold_first=True)

add_para(doc, '※ 선정 기준: 종합평점 70점 이상 프로젝트 중 고득점 순 5개 이내 우선협상 (공고문 p.4)')

doc.add_page_break()

# ================================================================
# 3. HEALO 평가 매트릭스 (핵심 표)
# ================================================================
add_heading(doc, '3. HEALO 평가 매트릭스', 1)
add_para(doc, '평가항목별 HEALO 대응 현황, 증빙 산출물, 추정 점수 정리')
doc.add_paragraph()

eval_tbl = doc.add_table(rows=0, cols=5)
eval_tbl.style = 'Table Grid'
add_header_row(eval_tbl, ['항목', '가중치', 'HEALO 대응', '증빙 산출물', '추정 점수'])

eval_data = [
    # ---- 사업계획의 타당성 20점 ----
    ('대상국가 선정 타당성\n(카자흐스탄 선택 근거)',
     '5점',
     '• 카자흐스탄 암 발생 연 41,300명 (2024, 카자흐 보건부)\n• 5년 유병자 102,845명 (GLOBOCAN 2022)\n• 2024년 한국 방문 14,475명 (중앙아시아 75.5%, KHIDI)\n• 한국 직항 (인천-알마티, 인천-아스타나), 인터넷 보급률 93%\n• 자국 의료 인프라 부족: 63% 병원 30년 노후, 33% 장비기준 미달\n(사업계획서 p.9~15)',
     '• 사업계획서 Ⅰ장 추진배경\n• GLOBOCAN 2022 인용 데이터\n• KHIDI 2024 외국인환자 유치 통계\n• 본 문서 01_요구사항정의서',
     '4.5/5'),

    ('사업계획의 적정성\n(계획·일정·업무범위)',
     '10점',
     '• 3단계 추진 전략 (플랫폼 고도화→환자유치→사후관리)\n• 5개 마일스톤 (M1~M5) + Gantt 차트 (사업계획서 p.31~32)\n• 컨소시엄 3기관 역할 분담 명확 (본로이 60%/면력 20%/신촌면력 20%)\n• 6대 ICT 서비스 유형 모두 구현 (공고문 p.8 기준)\n• AI Agent + Human Agent 상호학습 구조 (차별화)',
     '• 사업계획서 Ⅱ장 추진일정\n• 본 문서 02_기능명세서\n• HEALO 플랫폼 MVP 스크린샷\n• 시스템 아키텍처 구성도',
     '8.5/10'),

    ('사업비 사용계획 적정성',
     '5점',
     '• 총사업비 1억원 (정부출연금 8천만원 + 자부담 2천만원)\n• 자부담 현금 5% (500만원), 현물 15% (1,500만원) 구성\n• SW 개발·고도화비 편성 (공고문 p.2 요건 충족)\n• 참여기관별 배분: 본로이 60%, 면력 20%, 신촌면력 20%\n(사업계획서 Ⅶ장 사업비)',
     '• 사업계획서 Ⅶ장 사업비 총괄표\n• 별지5호 현금부담 확약서\n• 별지6호 자기부담현금 납입확약서',
     '4.0/5'),

    # ---- 사업화 성공가능성 35점 ----
    ('수익실현 가능성',
     '15점',
     '• 3가지 수익 모델: 환자유치 수수료(15~30%), 플랫폼 이용료(월50~100만원), 의료관광 패키지(건당100~200만원)\n• 2027년 예상 매출 1.8억원 (유치수수료 1.05억 + 이용료 0.45억 + 패키지 0.3억)\n• SaaS 확장: 타 의료기관·유치기관에 플랫폼 제공\n• 2024년 카자흐→한국 의료이용 14,475명, 시장 공백 선점 전략\n(사업계획서 p.11 수익모델)',
     '• 사업계획서 Ⅴ장 수익모델·파급효과\n• 시장 규모 데이터 (KHIDI 2024)\n• 경쟁사 비교 분석 (사업계획서 p.33)',
     '11.5/15'),

    ('사업 지속가능성\n(정부 지원 후 자체 운영)',
     '10점',
     '• 5개년 계획 수립: 2026 서비스 구축→2027 손익분기점→2028 CIS 확장→2030 글로벌 안정화\n• 2027년 자체 수익 기반 전환 마일스톤 제시\n• AI 학습 데이터 축적으로 운영 효율 자동 향상 구조\n• 장기적으로 Telemedicine 협력센터(2028~) 수익화\n(사업계획서 p.12 연차별 계획)',
     '• 사업계획서 연차별 계획 (p.6)\n• 5개년 수익 전망표\n• AI 자동화율 단계별 목표 (사업계획서 p.32)',
     '7.5/10'),

    ('환자유치 기여 가능성\n(국가 확대)',
     '10점',
     '• 카자흐스탄 성공 후 우즈베키스탄·몽골 순차 확대 계획\n• 동남아(베트남 등) 2029년 진출 검토\n• CIS 19,160명 → 향후 전체 타겟\n• 현지 에이전시 8곳 네트워크 구축 중\n• 직항 노선 (인천-알마티, 인천-아스타나) 접근성',
     '• 사업계획서 5개년 계획\n• CIS 시장 확대 전략 (p.12)\n• 현지 에이전시 접촉 현황',
     '7.5/10'),

    # ---- 사업내용의 현실성 30점 ----
    ('사업내용의 구체성\n(ICT 서비스, 추진방안)',
     '15점',
     '• 6대 ICT 서비스 (사전 3대 + 사후 3대) 공고문 기준 전부 구현\n• WebRTC 화상상담 (LiveKit), AI 실시간 번역, 다국어 UI 구체 명시\n• AI Agent + Human Agent 상호학습 아키텍처 상세 설계\n• RAG 3계층 구조 (DB→HIRA→Google) 기술 구체성\n• 특허 2건 (10-2745881, 10-2868334) 기술 차별화 증빙\n(사업계획서 p.27~31, 본 문서 02_기능명세서)',
     '• 02_기능명세서 (FN-AUTH~FN-SEC)\n• HEALO 플랫폼 MVP 시연\n• 특허등록증 2건\n• 시스템 아키텍처 다이어그램',
     '12.5/15'),

    ('사업내용의 실현가능성\n(추진체계, 위험관리)',
     '15점',
     '• 컨소시엄 3기관 명확한 역할 분담 + 주간 협의체 운영\n• 리스크 8개 항목 분석 및 대응방안 (사업계획서 p.44)\n• MVP 이미 보유 → 개발 리스크 최소화\n• 외국인환자 유치사업자 등록 완료 (법적 요건 충족)\n• 면력한방병원 외국인환자 유치의료기관 등록 완료\n• 대학병원 협진 네트워크 기 구축 (이대서울·고대구로·신촌세브란스)',
     '• 외국인환자 유치사업자 등록증\n• 면력한방병원 외국인환자 유치 등록증\n• 대학병원 협진 MOU 서류\n• HEALO MVP 플랫폼 배포 URL',
     '12.5/15'),

    # ---- 수행기관 역량 및 의지 15점 ----
    ('수행기관 역량 및 의지\n(전문성, 도전적 목표, 네트워크)',
     '15점',
     '• 목표 도전성: KHIDI 최소 목표 대비 상향 (유치 120%, 상담 150%, 만족도 112.5%)\n• ICT 전문성: HEALO MVP 개발·운영, AI 학습 시스템 PoC 완료\n• 특허 2건, 연구개발전담부서 인정서\n• 카자흐스탄 현지 파트너 접촉 중 (알마티·아스타나 에이전시)\n• SNS 채널(trek_korea) 운영 → 실제 환자 DB 확보\n• MSO 전문성 기반 의료 현장 이해\n(사업계획서 Ⅳ장 사업추진역량)',
     '• 특허등록증 10-2745881, 10-2868334\n• 연구개발전담부서 인정서\n• HEALO Survey 환자 DB 현황\n• SNS 채널 운영 실적\n• 외국인환자 유치사업자 등록증',
     '12.0/15'),
]

for r in eval_data:
    add_data_row(eval_tbl, r, bold_first=True)

doc.add_paragraph()
add_para(doc, '○ 본 배점 소계 (가산점 제외): 약 80.5점 / 100점 추정')
doc.add_paragraph()

# ================================================================
# 가산점 표
# ================================================================
add_heading(doc, '4. 가산점 분석 (최대 5점)', 2)
bonus_tbl = doc.add_table(rows=0, cols=4)
bonus_tbl.style = 'Table Grid'
add_header_row(bonus_tbl, ['가산점 항목', '가중치', 'HEALO 대응', '충족 여부'])
bonus_data = [
    ('외국인환자 유치의료기관 평가인증',
     '3점',
     '면력한방병원 외국인환자 유치의료기관 등록 완료\n신촌면력한방병원 외국인환자 유치의료기관 등록 완료\n(평가인증과 등록은 구분: 인증은 별도 절차)',
     '등록 완료 (인증은 별도 확인 필요)'),
    ('외국인환자 유치사업자 참여',
     '2점',
     '본로이: 외국인환자 유치사업자 등록 완료 (주관기관)\n(사업계획서 p.4, 사업계획서 붙임 등록증 첨부)',
     '충족: 2점 확보'),
    ('비수도권 의료기관 참여\n(서울/경기/인천 외)',
     '1점',
     '면력한방병원 (서울 소재): 서울은 수도권에 해당\n신촌면력한방병원 (서울 소재): 마찬가지\n비수도권 참여기관 없음 → 해당 없음',
     '미충족: 서울 소재 병원만 참여'),
]
for r in bonus_data:
    add_data_row(bonus_tbl, r, bold_first=True)

doc.add_paragraph()
add_para(doc, '○ 가산점: 유치사업자 등록 2점 확정. 유치의료기관 «등록»은 2개 기관 완료이며, 가산점 3점의 기준인 «평가인증»은 등록과 별개 절차로 취득 시 최대 5점 (현재 2점 확정)')

doc.add_page_break()

# ================================================================
# 5. 종합 점수 추정
# ================================================================
add_heading(doc, '5. 종합 점수 추정', 1)

score_tbl = doc.add_table(rows=0, cols=3)
score_tbl.style = 'Table Grid'
add_header_row(score_tbl, ['구분', '만점', '추정 점수'])
score_data = [
    ('사업계획의 타당성 소계', '20점', '17.0점'),
    ('  • 대상국가 선정 타당성', '5점', '4.5점'),
    ('  • 사업계획의 적정성', '10점', '8.5점'),
    ('  • 사업비 사용계획 적정성', '5점', '4.0점'),
    ('사업화 성공가능성 소계', '35점', '26.5점'),
    ('  • 수익실현 가능성', '15점', '11.5점'),
    ('  • 사업 지속가능성', '10점', '7.5점'),
    ('  • 환자유치 기여 가능성', '10점', '7.5점'),
    ('사업내용의 현실성 소계', '30점', '25.0점'),
    ('  • 사업내용의 구체성', '15점', '12.5점'),
    ('  • 사업내용의 실현가능성', '15점', '12.5점'),
    ('수행기관 역량 및 의지 소계', '15점', '12.0점'),
    ('기본 점수 합계', '100점', '80.5점'),
    ('가산점 (확정)', '2점', '2.0점 (유치사업자 확정)'),
    ('가산점 (조건부)', '3점', '0~3점 (의료기관 인증 여부 확인)'),
    ('최종 추정 합계', '105점', '82.5~85.5점'),
]
for r in score_data:
    bg = 'E8F0FE' if '소계' in r[0] or '합계' in r[0] or '최종' in r[0] else None
    add_data_row(score_tbl, r, bold_first=True, bg=bg)

doc.add_paragraph()
add_para(doc, '○ 선정 기준 70점 이상 충족: 추정 82.5~85.5점으로 우선협상 대상 가능 범위')
add_para(doc, '○ 최대 경쟁 리스크: 유치사업자 수익실현 가능성 점수 편차가 클 수 있음')

doc.add_page_break()

# ================================================================
# 6. Phase B 보강 우선순위
# ================================================================
add_heading(doc, '6. Phase B 보강 우선순위', 1)

add_para(doc, '평가 점수 향상을 위해 Phase B에서 보강이 필요한 항목을 우선순위 순으로 정리한다.')
doc.add_paragraph()

prio_tbl = doc.add_table(rows=0, cols=4)
prio_tbl.style = 'Table Grid'
add_header_row(prio_tbl, ['우선순위', '보강 항목', '기대 점수 향상', '작업 내용'])
prio_data = [
    ('P1 (즉시)',
     '외국인환자 유치의료기관 평가인증 확인\n→ 가산점 3점 확보 여부',
     '+3점 (가산점)',
     '면력한방병원 평가인증 취득 여부 공식 확인\n인증서 취득 추진 (미취득 시)'),
    ('P2 (긴급)',
     '수익 모델 구체화\n→ 수익실현 가능성 점수 향상',
     '+1.5~2.5점',
     '2027년 매출 예측 근거 데이터 보강\n계약 의향서·LOI 확보'),
    ('P3 (중요)',
     'AI 화상 내 실시간 번역 통합\n→ 사업내용 구체성 향상',
     '+1점',
     'FN-TRANS-02 Phase B 완성'),
    ('P4 (중요)',
     '카자흐스탄 현지 에이전시 계약 체결\n→ 수행기관 역량 점수 향상',
     '+1~1.5점',
     'MOU/계약서 체결 및 증빙 확보'),
    ('P5',
     '사후관리 AI 자동감지 완성\n→ 사업내용 현실성 향상',
     '+0.5점',
     'FN-POST-01 AI 이상징후 감지 구현'),
    ('P6',
     '러시아어·카자흐어 UI 100% 완성\n→ 다국어 실현 증빙',
     '+0.5점',
     'FN-I18N-01 전체 번역 완성'),
]
for r in prio_data:
    add_data_row(prio_tbl, r, bold_first=True)

doc.add_page_break()

# ================================================================
# 7. 사람 검토 필요 항목
# ================================================================
add_heading(doc, '7. 사람 검토 필요 항목 (Human Review)', 1)

review_tbl = doc.add_table(rows=0, cols=3)
review_tbl.style = 'Table Grid'
add_header_row(review_tbl, ['항목', '검토 내용', '담당'])
review_data = [
    ('유치의료기관 평가인증',
     '면력한방병원·신촌면력한방병원의 「외국인환자 유치의료기관 평가인증」 취득 여부 실제 확인\n(등록과 인증은 별개: 인증은 KHIDI 평가 필요)',
     '대표/면력한방병원 담당자'),
    ('대학병원 협진 MOU 공식화',
     '이대서울병원·이대목동병원·고대구로병원·신촌세브란스병원과의 협진 MOU 공식 문서 취득\n(사업계획서에 MOU 언급이나 공식 서류 필요)',
     '면력한방병원 담당자'),
    ('카자흐스탄 에이전시 계약',
     '알마티·아스타나 에이전시 「접촉 중」 → 「계약 체결」로 전환\n계약서 또는 LOI 확보',
     '본로이 마케팅팀'),
    ('사업비 예산 검토',
     '국고보조금 구성: SW 개발 인건비 편성 필수 확인\n홍보비/인건비/임차료 상한(각 20백만원) 준수 확인',
     'PM / 회계 담당'),
    ('개인정보 처리방침 카자흐어 버전',
     '카자흐스탄 개인정보보호법(2013) 준거 처리방침 작성 필요\n법률 전문가 검토 권고',
     '법무 담당 / 외부 법률사무소'),
    ('의료법 검토',
     '「사전상담」 서비스가 의료법상 원격진료로 해석될 리스크 재확인\n법무법인 의견서 취득 권고',
     '법무 담당'),
]
for r in review_data:
    add_data_row(review_tbl, r, bold_first=True)

# 저장
import os as _os
# 저장 위치는 «이 스크립트가 있는 폴더» 기준으로 잡는다.
# (전에는 특정 PC 의 절대경로가 박혀 있어 그 PC 밖에서는 재생성이 아예 불가능했다.)
out_path = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), 'EVAL_MATRIX.docx')
doc.save(out_path)
print(f'저장 완료: {out_path}')
print(f'총 단락 수: {len(doc.paragraphs)}')
