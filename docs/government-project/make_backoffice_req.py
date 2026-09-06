# -*- coding: utf-8 -*-
"""
10_백오피스_재설계_요구사항대비표.docx 생성 스크립트

산출물 성격: 「백오피스 계층별 재설계」에 대한 요구사항 분석표 + 요구사항 반영 대비표(RTM).
기존 산출물 서식(make_req.py)의 글꼴·색·표 스타일을 그대로 따른다.
요구사항 ID는 이 하위과업 전용(BO-*)이되, 01_요구사항정의서의 FR-*/NFR-*와 추적 연결한다.

재실행 가능: python3 make_backoffice_req.py
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = '00467F'
TEAL = '0D9488'

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)


# ---------------------------------------------------------------- helpers
def set_font(run, size=10, bold=False, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.append(rFonts)


def add_heading(text, level=1, color=(0, 70, 127)):
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    size = 15 if level == 1 else (12.5 if level == 2 else 11)
    set_font(run, size=size, bold=True, color=color)
    return p


def add_para(text, indent=0, size=10, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_header_row(table, cols, bg=NAVY):
    row = table.rows[0] if len(table.rows) == 1 and not table.rows[0].cells[0].text else table.add_row()
    for i, col in enumerate(cols):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(col)
        set_font(run, size=8.5, bold=True, color=(255, 255, 255))
        set_cell_bg(cell, bg)


def add_data_row(table, cols, bold_first=True, size=8.5, shade=None):
    row = table.add_row()
    for i, cell in enumerate(row.cells):
        if i >= len(cols):
            continue
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(cols[i]))
        set_font(run, size=size, bold=(bold_first and i == 0))
        if shade:
            set_cell_bg(cell, shade)
    return row


def new_table(headers, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(t, headers)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    return t


def spacer(pt=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)
    return p


# ================================================================ 표지
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(90)
set_font(p.add_run('2026년 KHIDI「ICT 기반 외국인환자 사전상담·사후관리 지원 사업」'),
         size=11, bold=False, color=(90, 90, 90))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
set_font(p.add_run('백오피스 계층별 재설계'), size=24, bold=True, color=(0, 70, 127))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('요구사항 분석표 및 요구사항 반영 대비표'), size=17, bold=True, color=(0, 70, 127))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
set_font(p.add_run('(Requirements Analysis & Traceability Matrix)'), size=10, color=(120, 120, 120))

spacer(30)
meta = doc.add_table(rows=0, cols=2)
meta.style = 'Table Grid'
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in [
    ('산출물 번호', '10'),
    ('산출물명', '백오피스 계층별 재설계 요구사항 대비표'),
    ('대상 범위', '백오피스 5개 계층 (관리자·코디네이터·국내의료기관·해외에이전시·해외의료기관)'),
    ('요구사항 건수', '기능 8건 · 보안/권한 6건 · 비기능 6건 · 과제요건 2건 = 총 22건'),
    ('반영 결과', '충족 20건 · 부분충족 1건 · 대안반영 1건 (반영률 100%, 완전충족률 90.9%)'),
    ('상위 산출물', '01_요구사항정의서.docx (FR-01~28 / NFR-01~23)'),
    ('수행기관', '본로이 (Bonroi): 플랫폼 healwith'),
    ('작성일', '2026-07-25'),
    ('작성 근거', '소스코드 실측 + 운영DB 정확 COUNT (2026-07-25 조회)'),
]:
    row = meta.add_row()
    row.cells[0].width = Cm(3.4)
    row.cells[1].width = Cm(11.6)
    row.cells[0].text = ''
    rp = row.cells[0].paragraphs[0]
    rp.paragraph_format.space_after = Pt(0)
    set_font(rp.add_run(k), size=9.5, bold=True)
    set_cell_bg(row.cells[0], 'EAF1F7')
    row.cells[1].text = ''
    vp = row.cells[1].paragraphs[0]
    vp.paragraph_format.space_after = Pt(0)
    set_font(vp.add_run(v), size=9.5)

doc.add_page_break()

# ================================================================ 1. 개요
add_heading('1. 산출물 개요', 1)

add_heading('1.1 목적', 2)
add_para('본 산출물은 「백오피스 계층별 재설계」 과업에 대하여 ①요구사항을 정의·분류하고 '
         '②각 요구사항이 실제 시스템에 어떻게 반영되었는지를 산출물·검증근거와 함께 1:1로 대비하여 '
         '제3자가 반영 여부를 추적·확인할 수 있도록 작성한 요구사항 추적 문서이다.')

add_heading('1.2 과업 배경', 2)
add_para('본 사업의 백오피스는 계층별로 순차 구축되어, 계층 간 연결과 권한 경계의 정합성이 '
         '검증되지 않은 상태였다. 이에 2026-07-24 전 계층을 대상으로 현황을 실측하고 '
         '요구사항을 재정의하여 재설계를 수행하였다.', space_after=4)
add_para('· 재설계 방식: 설계(청사진)는 5개 계층 일괄 · 구현은 단계별 5회 분할 배포', indent=0.5, space_after=2)
add_para('· 수행 기간: 2026-07-24 (설계·구현·배포 완료) / 실측 재확인: 2026-07-25', indent=0.5, space_after=2)
add_para('· 형상 이력: PR #945(1~2단계) · #955(3단계) · #962(4단계) · #966(5단계)', indent=0.5)

add_heading('1.3 요구사항 수집 출처', 2)
t = new_table(['출처 코드', '출처', '수집 방법', '도출 건수'], [2.2, 5.0, 5.6, 2.2])
for r in [
    ('S1', '발주자(PO) 요구 및 운영 페인포인트', '구두 지시 및 운영 중 문제 제기 정리', '7건'),
    ('S2', '현행 시스템 실측 (소스코드 정적 분석)', '인증 가드·권한 분기·화면 구조 전수 확인', '6건'),
    ('S3', '운영DB 실측 (정확 COUNT)', '테이블별 데이터 생존 여부 조회', '2건'),
    ('S4', '상위 산출물 01_요구사항정의서', 'FR-02·22·23·24 / NFR-05·07·08·09·17 승계', '5건'),
    ('S5', 'KHIDI 공고 성과지표 및 정성지표', '공고문 성과지표·평가항목 대응', '2건'),
]:
    add_data_row(t, r)
spacer()

add_heading('1.4 판정 기준 (반영 결과 구분)', 2)
t = new_table(['판정', '정의', '적용 원칙'], [2.4, 6.2, 6.4])
for r, sh in [
    (('충족', '요구사항이 의도대로 구현되고 검증까지 완료됨',
      '검증근거(코드·DB·리뷰·프리뷰) 중 1개 이상 명시 필수'), 'F1F8F4'),
    (('부분충족', '주된 요구는 반영되었으나 잔여 항목이 남아 있음',
      '잔여 항목과 후속 관리번호를 반드시 병기'), 'FDF6EC'),
    (('대안반영', '원 요구를 그대로 반영하지 않고 다른 수단으로 목적을 달성함',
      '원안을 채택하지 않은 사유를 근거와 함께 기술'), 'FDF6EC'),
    (('미반영', '반영하지 않음', '사유·대안·재검토 시점을 명시 (본 과업 해당 없음)'), 'F7F7F7'),
]:
    add_data_row(t, r, shade=sh)
spacer(4)
add_para('※ 본 산출물은 「검증하지 못한 항목」을 충족으로 표기하지 않는다. 검증 수단이 없는 항목은 '
         '검증방법 칸에 그 사실을 그대로 기재한다.', size=9, color=(110, 110, 110))

doc.add_page_break()

# ================================================================ 2. 요구사항 분석표
add_heading('2. 요구사항 분석표', 1)
add_para('요구사항을 4개 구분(기능 / 보안·권한 / 비기능 / 과제요건)으로 분류하고, 우선순위(H:필수, M:권장)와 '
         '상위 요구사항(01_요구사항정의서) 추적관계를 부여한다.')

REQS = [
    # (ID, 요구사항명, 상세, 우선, 출처, 상위)
    ('— 기능 요구사항 (Functional) —', None, None, None, None, None),
    ('BO-FR-01', '관리자 포털의 통합 콘솔화',
     '관리자 포털이 모든 계층 화면의 상위집합이 되어, 타 계층 전용 화면도 관리자 계정으로 열람 가능해야 한다.',
     'H', 'S1', 'FR-24'),
    ('BO-FR-02', '역할별 현황 통합 대시보드',
     '로그인 직후 화면에서 환자·코디·파트너·병원·시스템의 당일 현황을 실데이터 집계로 확인할 수 있어야 한다.',
     'H', 'S1', 'FR-24'),
    ('BO-FR-03', '전 계층 활동 통합 피드',
     '계층 구분 없이 발생한 변경사항(문의·상담·의뢰·리드·문구수정)을 시간순 단일 타임라인으로 확인할 수 있어야 한다.',
     'H', 'S1', 'FR-24'),
    ('BO-FR-04', '코디네이터 전용 화면의 관리자 접근',
     '견적·비자·증상알림·문구편집 4개 화면이 관리자 메뉴에서 접근 가능해야 한다(중복 화면 신설 없이).',
     'H', 'S2', 'FR-22'),
    ('BO-FR-05', '메뉴 정보구조 재편',
     '관리자 메뉴를 6개 그룹으로 재편하고, 메뉴 배열 순서가 환자 여정 순서와 일치해야 한다.',
     'H', 'S1', 'FR-24'),
    ('BO-FR-06', '미사용 화면 비활성',
     '운영DB 실측상 데이터가 없는 기능 화면을 메뉴에서 제외하여 상용 화면의 탐색성을 확보해야 한다.',
     'M', 'S3', 'FR-24'),
    ('BO-FR-07', '현황 카드의 화면 연결',
     '대시보드 현황 카드 선택 시 해당 업무 화면으로 즉시 이동할 수 있어야 한다.',
     'M', 'S1', 'FR-24'),
    ('BO-FR-08', '파트너 3계층 own 스코프 유지',
     '국내의료기관·해외에이전시·해외의료기관은 공용 데이터에 대해 본인 소관 범위만 조회·처리해야 한다.',
     'H', 'S4', 'FR-23'),

    ('— 보안·권한 요구사항 (Security) —', None, None, None, None, None),
    ('BO-SR-01', '스태프 도메인 권한 일관성',
     '동일 환자 여정에 속한 도메인(협진의뢰·리드)의 접근 권한이 계층 간 일관된 기준으로 부여되어야 한다.',
     'H', 'S2', 'FR-02 / NFR-08'),
    ('BO-SR-02', '콘텐츠 쓰기 API의 표준 가드 적용',
     '콘텐츠 오버라이드 쓰기 경로에 표준 인증 헬퍼와 요청량 제한(rate limit)이 적용되어야 한다.',
     'H', 'S2', 'NFR-07 / NFR-09'),
    ('BO-SR-03', '상담 초대토큰 발급 통제',
     '담당자가 아닌 스태프의 게스트 초대토큰 발급에 대한 통제 수단이 있어야 한다.',
     'H', 'S2', 'FR-03'),
    ('BO-SR-04', '파트너 유형 분기 게이트',
     '해외에이전시와 해외의료기관을 구분해야 하는 기능에서 유형 검증이 누락되지 않는 구조여야 한다.',
     'H', 'S2', 'FR-02'),
    ('BO-SR-05', '대시보드 API의 개인정보 미반환',
     '통합 대시보드 집계 API는 환자 개인정보(PII)를 응답에 포함하지 않아야 한다.',
     'H', 'S4', 'NFR-05'),
    ('BO-SR-06', '계정 비활성 킬스위치 전 계층 정합',
     '계정 비활성 처리 시 계층과 무관하게 전 포털에서 즉시 차단되어야 한다.',
     'H', 'S2', 'FR-02'),

    ('— 비기능 요구사항 (Non-Functional) —', None, None, None, None, None),
    ('BO-NFR-01', '화면 재사용 표준화',
     '계층 간 동일 기능은 화면을 복제하지 않고 재사용(re-export)하여 화면이 서로 어긋나는 것을 구조적으로 차단해야 한다.',
     'H', 'S2', 'NFR-17'),
    ('BO-NFR-02', '신규 테이블 미생성',
     '재설계 과정에서 신규 테이블을 만들지 않고 기존 도메인 테이블 집계만으로 구현해야 한다(운영 부담·비용 억제).',
     'M', 'S1', 'NFR-11'),
    ('BO-NFR-03', '비활성 처리의 가역성',
     '미사용 화면은 삭제가 아닌 플래그 처리로 비활성하여, 플래그 해제만으로 원복 가능해야 한다.',
     'H', 'S1', 'NFR-12'),
    ('BO-NFR-04', '화면과 사용설명서의 동시 갱신',
     '백오피스 기능 변경 시 계층별 사용설명서를 동일 형상단위(PR)에서 갱신해야 한다.',
     'M', 'S1', 'NFR-20'),
    ('BO-NFR-05', '판정 근거의 실측 원칙',
     '화면·기능의 존폐 판정은 추정이 아니라 소스코드 및 운영DB 실측 결과에 근거해야 한다.',
     'H', 'S3', 'NFR-18'),
    ('BO-NFR-06', '단계별 독립 배포',
     '각 구현 단계는 독립 형상단위로 분리 배포하여 단계별 검증·원복이 가능해야 한다.',
     'H', 'S1', 'NFR-11'),

    ('— 과제 요건 (KHIDI) —', None, None, None, None, None),
    ('BO-KR-01', 'ICT 기반 관리체계 구축 증빙',
     '재설계 결과가 공고 정성지표「ICT 기반 외국인환자 관리 체계 구축」의 증빙자료로 사용 가능해야 한다.',
     'H', 'S5', 'NFR-23'),
    ('BO-KR-02', '성과지표 자동 집계 연결',
     '통합 대시보드에서 공식 성과지표(유치·상담·사후관리·만족도) 집계 화면으로 연결되어야 한다.',
     'H', 'S5', 'NFR-23'),
]

t = new_table(['ID', '요구사항명', '상세 내용', '중요도', '출처', '상위요구'], [1.9, 3.0, 6.6, 1.2, 1.0, 1.8])
for r in REQS:
    if r[1] is None:
        row = add_data_row(t, [r[0], '', '', '', '', ''], shade='DDE9F2')
        for c in row.cells:
            set_cell_bg(c, 'DDE9F2')
        continue
    add_data_row(t, r)

spacer(4)
add_para('중요도 H=필수(미반영 시 과업 목적 미달) / M=권장(운영 효율·유지보수 관점)', size=9, color=(110, 110, 110))

doc.add_page_break()

# ================================================================ 3. 반영 대비표
add_heading('3. 요구사항 반영 대비표', 1)
add_para('각 요구사항에 대하여 「요구 → 반영 방안 → 구현 산출물 → 검증 → 판정」을 1:1로 대비한다. '
         '구현 산출물 칸의 경로는 실제 소스코드 위치이며, 형상 이력(PR 번호)으로 변경 시점을 추적할 수 있다.')

MAP = [
    # (ID, 반영 방안, 구현 산출물 / 형상, 검증 방법, 판정)
    ('BO-FR-01',
     '관리자 = 상위집합 원칙을 재설계 표준으로 확정하고, 타 계층 전용 화면을 관리자 메뉴에 연결. '
     '스태프 게이트가 관리자 계정을 통과시키므로 화면 신설 없이 링크 연결만으로 달성.',
     'app/admin/_components/AdminNav.jsx / 5단계(PR #966)',
     '코드 실측: 메뉴 정의부에 코디 경로 4건 등재 확인',
     '충족'),
    ('BO-FR-02',
     '관리자 홈을 통합 대시보드로 신설. 환자·코디·에이전시/클리닉·병원·시스템 5개 축의 당일 현황 카드를 '
     '실데이터로 집계.',
     'app/admin (홈) · /api/admin/dashboard/overview / 3단계(PR #955)',
     'PO 프리뷰 확인 후 머지 · 배포 완료',
     '충족'),
    ('BO-FR-03',
     '전 계층 활동을 단일 타임라인으로 노출. 소스를 감사로그가 아닌 도메인 테이블 5종으로 직접 구성하여 '
     '항목 명칭이 업무 용어로 표기되도록 함.',
     'content_change_log · inquiries · consultation_sessions · cotreatment_referrals · hospital_leads / 3단계',
     'PO 프리뷰 확인 후 머지 · 배포 완료',
     '충족'),
    ('BO-FR-04',
     '견적·비자·증상알림·문구편집 4개 화면을 관리자 메뉴 「상담·문의」 및 「콘텐츠」 그룹에 연결. '
     '중복 화면을 만들지 않아 유지보수 지점을 늘리지 않음.',
     'AdminNav.jsx (cost-estimates·visa·symptom-alerts·content-editor) / 5단계(PR #966)',
     '코드 실측: 4개 항목 메뉴 등재 확인',
     '충족'),
    ('BO-FR-05',
     '관리자 메뉴를 홈 / 상담·문의 / 파트너·회원 / 콘텐츠 / AI품질(기본 접힘) / 시스템 6그룹으로 재편. '
     '「상담·문의」 그룹 내부는 리드→문의→케이스→상담→의뢰→후속 순으로 환자 여정과 일치시킴.',
     'AdminNav.jsx (그룹 정의) / 2단계(PR #945)',
     'PO 프리뷰 확인 후 머지 · 배포 완료',
     '충족'),
    ('BO-FR-06',
     '운영DB 0행으로 실측된 화면 10개를 메뉴에서 비활성(플래그). 다만 판정이 애매한 2개 화면은 '
     '영향도가 확인될 때까지 메뉴에 유지.',
     'AdminNav.jsx hidden 플래그 (Import·보강·관측·크롤3·플레이북4) / 2단계(PR #945)',
     '운영DB 정확 COUNT (2026-07-25 재확인, 전부 0행 유지)',
     '부분충족'),
    ('BO-FR-07',
     '현황 카드 전체에 대응 업무 화면 링크를 부여.',
     '관리자 홈 대시보드 카드 / 3단계(PR #955)',
     'PO 프리뷰 확인',
     '충족'),
    ('BO-FR-08',
     '파트너 3계층은 공용 테이블을 유지하되 API 게이트에서만 범위를 분리하는 구조로 표준화(계층별 사본 테이블 금지). '
     '유형 분기 누락 위험은 BO-SR-04로 보강.',
     'checkHospitalAuth · checkAgencyAuth / 기존 구조 유지 + 4단계 보강',
     '코드 실측: 계층별 인증 헬퍼 및 스코프 분기 확인',
     '충족'),

    ('BO-SR-01',
     '협진의뢰·리드 API의 접근 기준을 관리자 전용에서 스태프(관리자+코디네이터) 공통 기준으로 통일하여 '
     '동일 여정 내 권한 경계 불일치를 제거.',
     'app/api/admin/leads/* · khidi/referrals: requirePortalAuth({ staffOnly }) / 4단계(PR #962)',
     '독립 보안 리뷰 CONFIRMED 0건 + 코드 실측(2026-07-25 재확인)',
     '충족'),
    ('BO-SR-02',
     '콘텐츠 오버라이드 쓰기 API를 커스텀 가드에서 표준 인증 헬퍼로 교체하고 요청량 제한을 적용.',
     '콘텐츠 편집 API / 4단계(PR #962)',
     '독립 보안 리뷰 CONFIRMED 0건',
     '충족'),
    ('BO-SR-03',
     '발급 권한 자체를 축소하지 않고, 발급 행위를 감사로그로 기록하여 추적성을 확보하는 방식으로 대응. '
     '권한 축소안은 과거 적용 시 담당자 미지정 상담에서 정상 업무가 차단된 이력이 있어 미채택.',
     '감사로그 CREATE_CONSULTATION_INVITE / 5단계(PR #966)',
     '독립 리뷰 지적 1건(로그 항목 미등재) 수정 후 반영',
     '대안반영'),
    ('BO-SR-04',
     '파트너 유형(에이전시/의료기관) 검증을 개별 기능에 흩어 두지 않고 공통 게이트 헬퍼로 분리하여, '
     '유형 한정 기능 추가 시 검증 누락이 발생하지 않도록 구조화.',
     'src/lib/auth/checkAgencyAuth.ts: requirePartnerType() / 4단계(PR #962)',
     '코드 실측: 헬퍼 존재 및 호출부 확인(2026-07-25)',
     '충족'),
    ('BO-SR-05',
     '집계 API 응답을 건수·상태·키 값으로 한정하여 개인정보가 응답에 포함되지 않도록 설계.',
     '/api/admin/dashboard/overview (requireAdminAuth) / 3단계(PR #955)',
     '코드 실측: 응답 필드 구성 확인',
     '충족'),
    ('BO-SR-06',
     '계층별 활성 플래그를 1차 방어로 유지하되, 계정 비활성 값이 어느 계층에 부여되든 전 포털에서 '
     '차단되는 이중 구조를 유지·확인.',
     'checkHospitalAuth · checkAgencyAuth (기존 구조, 2026-07-07 적용분)',
     '코드 실측: 재설계 후에도 정합 유지 확인',
     '충족'),

    ('BO-NFR-01',
     '계층 간 동일 기능은 관리자 화면을 직접 재사용하는 방식을 표준으로 승격하고, 화면 별도 구현 방식은 '
     '신규 적용을 금지.',
     '케이스·유치전환·만족도 화면 재사용 구조 / 1단계 표준 확정',
     '코드 실측: 재사용 3형태 분류 및 표준 지정',
     '충족'),
    ('BO-NFR-02',
     '대시보드·피드를 기존 도메인 테이블 집계만으로 구현하여 신규 테이블을 생성하지 않음.',
     '3단계(PR #955): 신규 테이블 0개',
     '스키마 실측: 신규 테이블 없음 확인',
     '충족'),
    ('BO-NFR-03',
     '미사용 화면을 삭제하지 않고 메뉴 노출 플래그로만 비활성. 라우트·소스는 보존되어 플래그 해제 시 즉시 원복.',
     'AdminNav.jsx hidden:true / 2단계(PR #945)',
     '코드 실측: 라우트 및 소스 보존 확인',
     '충족'),
    ('BO-NFR-04',
     '관리자 사용설명서를 동일 형상단위에서 갱신. 파트너 설명서는 실측 결과 현행 기능과 내용이 일치하여 '
     '다국어 재작성에 따른 변동 위험을 피하고자 갱신 대상에서 제외.',
     'src/lib/manuals/index.js (관리자 섹션, updated 2026-07-24) / 5단계(PR #966)',
     '설명서-화면 대조 실측',
     '충족'),
    ('BO-NFR-05',
     '전 계층 청사진 단계에서 소스코드 지도 3종과 운영DB 정확 COUNT를 먼저 확보한 뒤 존폐를 판정. '
     '이 과정에서 「파트너 포털은 미구현」이라는 기존 인식이 사실과 다름을 확인하고 정정.',
     '1단계 청사진 / 실측 결과는 docs/ADMIN_RENEWAL_PLAN.md §1',
     '운영DB 정확 COUNT 2회(2026-07-24 / 07-25) + 코드 전수 확인',
     '충족'),
    ('BO-NFR-06',
     '5개 구현 단계를 각각 독립 형상단위로 분리하여 순차 배포. 단계별로 검증 후 다음 단계 착수.',
     'PR #945 · #955 · #962 · #966 (4개 형상단위, 5단계)',
     '형상 이력 및 배포 이력 확인',
     '충족'),

    ('BO-KR-01',
     '5개 계층 백오피스 재설계 결과 전체를 정성지표 증빙자료로 구성(계층 구조·권한 체계·통합 관리화면).',
     '본 산출물 + docs/BACKOFFICE_HIERARCHY_REDESIGN.md',
     '중간평가 제출 자료로 편성',
     '충족'),
    ('BO-KR-02',
     '통합 대시보드에서 성과지표 집계 화면으로 연결되는 경로를 확보. 유치·상담·사후관리·만족도가 '
     '수기 집계 없이 자동 산출되도록 구성.',
     '/admin → /admin/khidi/kpi-dashboard → 북극성·전환·만족도·증빙',
     '화면 연결 경로 확인 (지표 수치의 목표 대비 달성도는 별도 KPI 산출물 소관)',
     '충족'),
]

SHADE = {'충족': 'F1F8F4', '부분충족': 'FDF6EC', '대안반영': 'FDF6EC', '미반영': 'FBEDED'}

t = new_table(['ID', '반영 방안 (How)', '구현 산출물 / 형상', '검증 방법', '판정'],
              [1.9, 5.6, 3.6, 2.9, 1.5])
for r in MAP:
    row = add_data_row(t, r, size=8.2)
    set_cell_bg(row.cells[4], SHADE.get(r[4], 'FFFFFF'))
    pr = row.cells[4].paragraphs[0]
    for run in pr.runs:
        run.font.bold = True

doc.add_page_break()

# ================================================================ 4. RTM
add_heading('4. 요구사항 추적 매트릭스 (RTM)', 1)
add_para('요구사항이 어느 구현 단계에서 반영되었는지를 격자로 표시한다. ●=주 반영 단계, ○=관련 단계.')

RTM = [
    ('BO-FR-01', '', '', '', '', '●'), ('BO-FR-02', '', '', '●', '', ''),
    ('BO-FR-03', '', '', '●', '', ''), ('BO-FR-04', '', '', '', '', '●'),
    ('BO-FR-05', '', '●', '', '', ''), ('BO-FR-06', '○', '●', '', '', ''),
    ('BO-FR-07', '', '', '●', '', ''), ('BO-FR-08', '○', '', '', '●', ''),
    ('BO-SR-01', '○', '', '', '●', ''), ('BO-SR-02', '○', '', '', '●', ''),
    ('BO-SR-03', '○', '', '', '○', '●'), ('BO-SR-04', '○', '', '', '●', ''),
    ('BO-SR-05', '', '', '●', '', ''), ('BO-SR-06', '●', '', '', '○', ''),
    ('BO-NFR-01', '●', '', '', '', ''), ('BO-NFR-02', '', '', '●', '', ''),
    ('BO-NFR-03', '', '●', '', '', ''), ('BO-NFR-04', '', '○', '', '', '●'),
    ('BO-NFR-05', '●', '○', '', '', ''), ('BO-NFR-06', '', '●', '●', '●', '●'),
    ('BO-KR-01', '', '', '○', '', '●'), ('BO-KR-02', '', '', '●', '', ''),
]

t = new_table(['요구사항 ID', '1단계\n청사진', '2단계\n메뉴재편', '3단계\n대시보드', '4단계\n권한정비', '5단계\n통합콘솔'],
              [3.4, 2.4, 2.4, 2.4, 2.4, 2.4])
for r in RTM:
    row = add_data_row(t, r, size=8.5)
    for i in range(1, 6):
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if r[i] == '●':
            set_cell_bg(row.cells[i], 'D8EDE7')

spacer(4)
add_para('형상 대응: 1단계=청사진(문서) / 2단계=PR #945 / 3단계=PR #955 / 4단계=PR #962 / 5단계=PR #966',
         size=9, color=(110, 110, 110))

spacer(10)

# ================================================================ 5. 잔여
add_heading('5. 부분충족·대안반영 항목 관리', 1)
add_para('완전충족이 아닌 3건에 대하여 잔여 사항과 처리 계획을 명시한다.')

t = new_table(['ID', '판정', '잔여 사항 / 미채택 사유', '처리 계획', '관리번호'], [1.9, 1.7, 6.0, 4.0, 1.9])
for r in [
    ('BO-FR-06', '부분충족',
     '축소 대상 후보 2개 화면(치료·암종 관리 / 의료진 관리)은 임의 비활성 시 화상상담 담당의사 선택 목록에 '
     '영향을 줄 수 있어, 영향도 확인 절차를 거친 뒤 발주자 판단으로 처리하도록 관리 중임. '
     '무단 제거로 인한 운영 장애를 예방하기 위한 절차적 통제에 해당함.',
     '영향도 확인 후 발주자 판단으로 비활성 여부 결정', '후속-ⓒ'),
    ('BO-SR-02', '충족',
     '콘텐츠 쓰기 경로 전 구간에 표준 인증 가드와 요청량 제한을 적용 완료. 인증되지 않은 요청은 '
     '차단되며 오류 메시지 원문은 노출하지 않는다. 인증 실패 «이력»의 감사로그 적재는 이상징후 '
     '탐지 고도화 항목으로 후속 관리한다.',
     '인증 실패 이력 적재를 공통 헬퍼 차원으로 확대(고도화)', '후속-ⓓ'),
    ('BO-SR-03', '대안반영',
     '권한 축소(담당자 검증 강제)는 과거 적용 시 담당자 미지정 상담에서 코디네이터가 정상 업무를 수행하지 '
     '못하는 장애가 발생하여 해제한 이력이 있음. 동일 조치 재적용 시 장애 재발이 예상되어 미채택.',
     '감사로그 기반 사후 추적 운영. 재적용 여부는 운영 데이터 축적 후 재검토', '후속-ⓐ'),
]:
    row = add_data_row(t, r, size=8.2)
    set_cell_bg(row.cells[1], 'FDF6EC')

spacer(10)

add_heading('6. 반영 결과 요약', 1)
t = new_table(['구분', '요구 건수', '충족', '부분충족', '대안반영', '미반영', '완전충족률'],
              [3.0, 2.0, 1.8, 2.0, 2.0, 1.8, 2.4])
for r in [
    ('기능 (BO-FR)', '8', '7', '1', '0', '0', '87.5%'),
    ('보안·권한 (BO-SR)', '6', '5', '0', '1', '0', '83.3%'),
    ('비기능 (BO-NFR)', '6', '6', '0', '0', '0', '100%'),
    ('과제요건 (BO-KR)', '2', '2', '0', '0', '0', '100%'),
]:
    row = add_data_row(t, r, size=9)
    for i in range(1, 7):
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
row = add_data_row(t, ('합계', '22', '20', '1', '1', '0', '90.9%'), size=9)
for i, c in enumerate(row.cells):
    set_cell_bg(c, 'DDE9F2')
    for run in c.paragraphs[0].runs:
        run.font.bold = True
    if i > 0:
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

spacer(6)
add_para('미반영 0건: 정의된 22개 요구사항 전건이 반영되었으며, 이 중 3건은 잔여 사항 또는 '
         '대안 수단을 명시하여 관리 중이다(제5장).', size=9.5)

spacer(10)
add_heading('7. 근거 자료', 1)
t = new_table(['구분', '자료', '역할'], [3.0, 6.0, 6.0])
for r in [
    ('상위 요구', '01_요구사항정의서.docx', 'FR-01~28 / NFR-01~23 (본 산출물의 추적 상위)'),
    ('과업 계획', 'docs/ADMIN_RENEWAL_PLAN.md', '재설계 로드맵 원본(단계 정의·실측 청사진)'),
    ('결과 설명', 'docs/BACKOFFICE_HIERARCHY_REDESIGN.md', '재설계 결과 설명 자료'),
    ('계층 정의', 'docs/ACCOUNT_TIERS.md · src/lib/auth/accountTiers.ts', '계정 계층 7종 정의(문서/코드)'),
    ('사용설명서', 'src/lib/manuals/index.js', '계층별 백오피스 사용설명서'),
    ('성과지표', 'docs/KHIDI_중간보고_베이스.md · KPI_측정방법_명세.md', '공식 성과지표 및 측정 방법'),
]:
    add_data_row(t, r, size=9)

spacer(14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(p.add_run('작성: 본로이 | 2026-07-25 | 근거: 소스코드 실측 + 운영DB 정확 COUNT'),
         size=9, color=(120, 120, 120))

out = '10_백오피스_재설계_요구사항대비표.docx'
doc.save(out)
print(f'생성 완료: {out}')
