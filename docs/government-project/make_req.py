import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import _facts as F   # 사실의 단일 소스 — 2026-09-06 부터 이 생성기도 읽는다

doc = Document()

# 페이지 여백 설정
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

def set_font(run, size=10, bold=False, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.append(rFonts)

def add_heading(doc, text, level=1, color=(0,70,127)):
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    size = 14 if level == 1 else (12 if level == 2 else 11)
    set_font(run, size=size, bold=True, color=color)
    return p

def add_para(doc, text, indent=0, size=10):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_header_row(table, cols):
    row = table.add_row()
    for i, col in enumerate(cols):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(col)
        set_font(run, size=9, bold=True, color=(255,255,255))
        set_cell_bg(cell, '00467F')

def add_data_row(table, cols, bold_first=False):
    row = table.add_row()
    for i, col in enumerate(row.cells):
        if i < len(cols):
            col.text = ''
            p = col.paragraphs[0]
            run = p.add_run(str(cols[i]))
            b = bold_first and i == 0
            set_font(run, size=9, bold=b)

# ================================================================
# 표지
# ================================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"{F.PROJECT['플랫폼']} 플랫폼")
set_font(run, 20, True, (0,70,127))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('요구사항 정의서')
set_font(run, 24, True, (0,70,127))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ICT 기반 외국인환자 사전상담·사후관리 지원 사업')
set_font(run, 12, False, (80,80,80))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Requirements Definition Document  |  {F.DOC_VERSION}  |  {F.DOC_DATE_DOT}')
set_font(run, 10, False, (120,120,120))

doc.add_paragraph()
doc.add_paragraph()
tbl_info = doc.add_table(rows=0, cols=2)
tbl_info.style = 'Table Grid'
info_rows = [
    ('문서 번호', 'healwith-REQ-2026-001'),
    ('작성 기관', '본로이 (Bonroi)'),
    ('작성일', F.DOC_DATE_KO),
    ('버전', F.DOC_VERSION + ' (v1.2 = 2026-08-20, v1.1 = 2026-08-19, v1.0 = 2026-04-30)'),
    ('보안 등급', '대외비'),
    ('관련 사업', '2026년 KHIDI ICT 기반 외국인환자 사전상담·사후관리 지원 사업'),
]
for label, val in info_rows:
    row = tbl_info.add_row()
    row.cells[0].text = ''
    r1 = row.cells[0].paragraphs[0].add_run(label)
    set_font(r1, 10, True, (0,70,127))
    row.cells[1].text = ''
    r2 = row.cells[1].paragraphs[0].add_run(val)
    set_font(r2, 10)

doc.add_page_break()

# ================================================================
# 목차
# ================================================================
add_heading(doc, '목  차', 1)
toc_items = [
    ('1.', '사업 개요', False),
    ('  1.1', '사업 목적', True),
    ('  1.2', '사업 범위', True),
    ('  1.3', '기대효과', True),
    ('2.', '대상 사용자 페르소나', False),
    ('  2.1', '해외 암환자 (P-01)', True),
    ('  2.2', '코디네이터 (P-02)', True),
    ('  2.3', '의료진 (P-03)', True),
    ('  2.4', '관리자 (P-04)', True),
    ('3.', '기능 요구사항 (FR-01 ~ FR-37)', False),
    ('  3-1', '추적표: 요구 → 화면 → 기능 명세', True),
    ('4.', '비기능 요구사항', False),
    ('5.', '외부 시스템 연계', False),
    ('6.', '제약사항', False),
    ('7.', '용어 정의', False),
]
for num, title, indented in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0) if indented else Cm(0)
    run = p.add_run(f'{num}  {title}')
    set_font(run, 10)

doc.add_page_break()

# ================================================================
# 1. 사업 개요
# ================================================================
add_heading(doc, '1. 사업 개요', 1)

add_heading(doc, '1.1 사업 목적', 2)
add_para(doc, '본 문서는 healwith 플랫폼(AI 기반 외국인환자 사전상담·사후관리 통합 플랫폼)의 시스템 요구사항을 정의한다. 2026년 KHIDI 「ICT 기반 외국인환자 사전상담·사후관리 지원 사업」(공고문 p.1~2) 참여 과제로서, 카자흐스탄 암환자의 한국 의료 접근성 향상을 1차 목표로 한다.')
add_para(doc, '사업계획서(p.2)에 명시된 목적을 기반으로, 본 플랫폼은 다음 세 가지 구조적 문제를 해결한다:')
add_para(doc, '(1) 정보 비대칭: 환자가 병원·의료진·치료비를 사전에 비교할 수 없는 구조적 문제', indent=1)
add_para(doc, '(2) 언어 장벽: 러시아어·카자흐어 대응 상담 인프라 부족 (사업계획서 p.9)', indent=1)
add_para(doc, '(3) 사후관리 단절: 귀국 후 경과 모니터링·재방문 연계 체계 부재', indent=1)

add_heading(doc, '1.2 사업 범위', 2)
scope_tbl = doc.add_table(rows=0, cols=2)
scope_tbl.style = 'Table Grid'
add_header_row(scope_tbl, ['구분', '내용'])
scope_data = [
    ('대상 국가', '카자흐스탄 (1차). 향후 우즈베키스탄·러시아 CIS 확대 예정'),
    ('대상 질환', '주로 암(Cancer) 환자: 방사선·표적치료·면역치료 목적'),
    ('사업 기간', '협약체결일 ~ 2026년 11월 20일 (약 8개월, 단년도 사업)'),
    ('원격협진', '7월~11월 (5개월, 필수 4개월 초과 충족)'),
    ('플랫폼 범위', '웹(PC/모바일, healwith.co.kr) + Capacitor 앱(iOS/Android — 구글 프로덕션 검토·애플 심사 대기 중, 2026-09-06)'),
    ('수행 기관', '본로이(주관) + 면력한방병원 + 신촌면력한방병원(참여)'),
    ('협진 병원', '이대서울병원, 이대목동병원, 고려대 구로병원, 신촌세브란스병원'),
]
for r in scope_data:
    add_data_row(scope_tbl, r, bold_first=True)

doc.add_paragraph()

add_heading(doc, '1.3 기대효과', 2)
eff_tbl = doc.add_table(rows=0, cols=3)
eff_tbl.style = 'Table Grid'
add_header_row(eff_tbl, ['구분', '기대효과', '측정 지표'])
eff_data = [
    ('기술적', 'AI+Human 상호학습 기반 상담 자동화 구현\n다국어 AI 번역 시스템 검증\nWebRTC 화상상담 플랫폼 안정화', 'AI 자동응답 정확도 30%+\n화상상담 연결 성공률 95%+'),
    ('경제적', '카자흐스탄 암환자 고부가가치 유치 확대\n2027년 예상 매출 1.8억원 (유치수수료+이용료)\nSaaS 확장 시 추가 수익', '환자 유치 12건/8개월\nKPMG: 헬스케어 AI 투자 10배 성장 추세 부합'),
    ('사회적', '한-카자흐스탄 의료협력 강화\n언어·정보 장벽 해소로 의료 접근성 형평 제고\n환자 치료 성과 지속 모니터링', '만족도 90점+\nCIS 환자 만족도 평균 92.1점 벤치마크'),
    ('정책적', 'KHIDI ICT 사업 공고 6대 서비스 유형 전부 구현\n디지털 전환 표준 모델 제시\n양국 정부 정책 교차점 부합', 'KPI 달성률\n공고문 p.8 ICT 요소 체크리스트 충족'),
]
for r in eff_data:
    add_data_row(eff_tbl, r, bold_first=True)

doc.add_page_break()

# ================================================================
# 2. 대상 사용자 페르소나
# ================================================================
add_heading(doc, '2. 대상 사용자 페르소나', 1)

personas = [
    {
        'title': '2.1 해외 암환자 (P-01)',
        'items': [
            ('페르소나 ID', 'P-01'),
            ('대표 프로필', '알마티 거주, 45세 여성, 유방암 2기 진단. 러시아어 사용. 한국어·영어 불가.'),
            ('기술 수준', '스마트폰 일상 사용, PC 기본 가능'),
            ('주요 니즈', '한국 상급종합병원 암 치료 정보 비교\n예상 치료비·일정 사전 확인\n러시아어로 의료진 상담\n귀국 후 경과 모니터링 지속'),
            ('고충 (Pain Points)', '한국 의료기관 정보 접근 어려움\n언어 장벽으로 상담 불가\n비용·비자 절차 불투명\n귀국 후 의료진과 연락 단절'),
            ('사용 시나리오', '① healwith 웹 접속(러시아어 UI)\n② 증상·희망 치료 입력, AI 병원 매칭\n③ 화상상담 예약, 의료문서 업로드\n④ 내원·치료 (면역치료+대학병원 협진)\n⑤ 귀국 후 사후관리 앱 활용'),
            ('핵심 라우트', '/ru 또는 /kz (다국어 홈), /inquiry (통합 문의 퍼널), /consultation/[id], /patient/*'),
        ]
    },
    {
        'title': '2.2 코디네이터 (P-02)',
        'items': [
            ('페르소나 ID', 'P-02'),
            ('대표 프로필', '본로이 소속, 28세, 러시아어·한국어 능통. 환자 여정 전담 관리.'),
            ('기술 수준', 'PC/앱 업무 처리, CRM·대시보드 능숙'),
            ('주요 니즈', '담당 환자 전체 여정 통합 관리\nAI 이관 상담 처리 및 보완\n의료문서 번역·라우팅\n성과 실적 기록·보고'),
            ('고충', '카카오톡·이메일·전화 분산 채널 통합 어려움\nAI 오답 수정 업무 과부하\n서류 번역·포매팅 반복 작업'),
            ('사용 시나리오', '① 코디네이터 포털 로그인\n② AI 이관 알림 → 상담 처리\n③ 의료문서 업로드·병원 라우팅\n④ 환자 일정·예약 관리\n⑤ 사후관리 상태 모니터링'),
            ('핵심 라우트', '/coordinator/inbox, /coordinator/messages, /coordinator/consultations'),
        ]
    },
    {
        'title': '2.3 의료진: 파트너 병원 (P-03)',
        'items': [
            ('페르소나 ID', 'P-03'),
            ('대표 프로필', '면력한방병원 종양내과 전문의, 42세. 외국인 환자 상담 경험 보유.'),
            ('기술 수준', 'EMR 시스템 사용 경험, 원격 진료 툴 익숙'),
            ('주요 니즈', '환자 의료정보(CT/MRI/검사결과) 사전 검토\nWebRTC 화상상담 수행\n경과 추적 기록 입력\n협진 의뢰서 전송'),
            ('고충', '여러 플랫폼 로그인 분산\n의료문서 형식 불일치\n언어 장벽으로 직접 소통 한계'),
            ('사용 시나리오', '① 상담방 초대링크 수신(계정 발급 불필요)\n② 환자 문서 열람 (사전 검토)\n③ 화상상담 참여 (실시간 자막·통역 지원)\n④ 처방·메모 기록\n⑤ 전문의 소견 전송'),
            ('핵심 라우트', '/consultation/[id] (초대링크 입장), /opinion/[token] (전문의가 소견을 «작성»하는 화면 — 환자 확인은 /claim/[token]), /hospital/* (병원 담당자 계정)'),
        ]
    },
    {
        'title': '2.4 플랫폼 관리자 (P-04)',
        'items': [
            ('페르소나 ID', 'P-04'),
            ('대표 프로필', '본로이 운영팀 PM. 사업 총괄·KPI 관리·KHIDI 보고 담당.'),
            ('기술 수준', 'IT 툴 능숙, 대시보드·통계 분석 가능'),
            ('주요 니즈', '환자 유치 실적 실시간 모니터링\nAI 성능 지표 추적\n병원 정보·의료진 데이터 관리\nKHIDI 성과보고 자료 산출'),
            ('고충', '분산 데이터 수동 집계\nAI 오류 자동 감지 부재\n정부 보고 형식과 내부 데이터 불일치'),
            ('사용 시나리오', '① Admin 포털 로그인\n② KPI 대시보드 확인\n③ AI 성능·RAG 상태 모니터링\n④ 병원·의료진 데이터 관리\n⑤ 성과보고서 출력'),
            ('핵심 라우트', '/admin/*, /admin/analytics, /admin/consultations'),
        ]
    },
]

for persona in personas:
    add_heading(doc, persona['title'], 2)
    ptbl = doc.add_table(rows=0, cols=2)
    ptbl.style = 'Table Grid'
    add_header_row(ptbl, ['항목', '내용'])
    for item_k, item_v in persona['items']:
        add_data_row(ptbl, [item_k, item_v], bold_first=True)
    doc.add_paragraph()

doc.add_page_break()

# ================================================================
# 3. 기능 요구사항
# ================================================================
add_heading(doc, '3. 기능 요구사항', 1)
add_para(doc, '우선순위: H=High(필수) / M=Medium(권장) / L=Low(선택)')
add_para(doc, '현황 표기: 완료 / 부분구현 / 미구현')
doc.add_paragraph()

fr_tbl = doc.add_table(rows=0, cols=5)
fr_tbl.style = 'Table Grid'
add_header_row(fr_tbl, ['FR-ID', '기능명', '우선순위', '상세 설명', '현황'])

fr_data = [
    ('FR-01', '소셜/이메일 회원가입', 'H',
     '이메일·Google·Apple 로그인(웹 + 앱 네이티브). 계정 계층 7종(별표 FR-02). Supabase Auth 사용.',
     '완료 · /app/signup, /app/login, src/lib/auth/{googleNativeSignIn,appleNativeSignIn}.ts (앱 안 애플 로그인은 심사 대응 중)'),

    ('FR-02', '역할 기반 접근 제어(RBAC)', 'H',
     '계정 계층 7종(비회원·환자·코디네이터·관리자·국내 의료기관·해외 에이전시·해외 의료기관)으로 분기. '
     '권한 저장 위치는 계층별로 다름(app_metadata.role / hospital_users / agency_users / 초대링크 토큰). '
     'proxy.ts 가 서버 단계에서 다섯 포털(/admin·/hospital·/patient·/coordinator·/agency·/clinic) 전부 세션을 확인하되 역할(role)까지 보는 곳은 /admin 뿐이고, 나머지는 화면·API(checkAgencyAuth·hospital_users 조회)가 세부 권한을 판정한다. '
     '※ 의사는 계정 계층이 아니라 상담방 초대링크 게스트 또는 병원 계정으로 참여함.',
     '완료 · proxy.ts (구 middleware — Next.js 16 에서 이름이 바뀌었다), src/lib/auth/accountTiers.ts (계층 단일 표준)'),

    ('FR-03', '게스트 토큰 발급', 'H',
     '비회원 환자에게 public_token을 발급하여 회원가입 없이 초기 상담 접근 가능.',
     '완료 · migrations/20260125_inquiries_public_token_and_attachments'),

    ('FR-04', '암환자 문의 접수(1단계+2단계)', 'H',
     '환자 기본정보, 암 종류·병기, 의료기록 업로드, 치료 희망 사항 수집. Progressive 단계별 폼. AES-256-GCM 암호화 저장. (사업계획서 p.27)',
     '완료 · /app/inquiry (통합 문의 퍼널), migrations/20260125_inquiries_intake_progressive'),

    ('FR-05', '의료문서 업로드·관리', 'H',
     'CT/MRI/검사결과/진단서 업로드. Supabase Storage 저장, 암호화. MIME 타입 검증, 파일 크기 제한.',
     '완료 · /app/api/attachments, migrations/20260406_consultation_documents'),

    ('FR-06', '문서 생성·다운로드', 'M',
     'PDF 3종(동의서·견적서·비자 초청장, @react-pdf/renderer) + 협진 병원 원본 워드 양식 의뢰서(DOCX). 상담 요약 PDF 는 없다.',
     '완료 · /app/api/pdf/{consent,quotation}, src/lib/pdf, /app/api/coordinator/inquiries/[id]/referral-docx'),

    ('FR-07', 'AI 기반 병원·의료진 매칭', 'H',
     '환자 증상·희망 치료·예산 기반 병원 자동 추천. Gemini Flash(최신 별칭) + RAG(자체 DB 벡터 → HIRA·네이버 실시간 보조검색; 구글 검색 그라운딩은 미작동). 6개 언어 결과. DB 에 없는 병원은 추천하지 않는다. (사업계획서 p.27)',
     '완료 · /app/api/public/chat/*, src/lib/chat/{generateReply,externalSearch}.ts, src/lib/rag'),

    ('FR-08', '병원 목록·상세 조회', 'H',
     '병원명·진료과·의료진·시설·가격 정보 다국어(ko/en/ru/kz/zh/ja) 제공. i18n JSONB 컬럼 활용.',
     '완료 · /app/hospitals, migrations/20260223_i18n_jsonb'),

    ('FR-09', '예상 비용 산출', 'M',
     '진료 항목별 예상비용 자동 산출 안내. AI 적정가 판독. 비자·숙박 비용 안내 포함. (사업계획서 p.28)',
     '부분구현 · /app/coordinator/cost-estimates (견적 작성·이력 관리 가동. AI 적정가 자동 산출은 남음)'),

    ('FR-10', 'AI 챗봇 상담 (24시간)', 'H',
     'Gemini 기반 AI Agent. 6개 언어(ko·en·ru·kz·zh·ja). RAG 기반 병원 정보 응답. 3턴 넘기면 문의로 승격, 자료·사람 요청 시 이관. 완치 단정·지어낸 병원·연락처 없는 「접수됨」 차단.',
     '완료 · /app/api/public/chat/*, src/lib/chat, src/lib/rag, /app/api/webhooks/{telegram,whatsapp}'),

    ('FR-11', 'Human Agent 상담 (코디네이터)', 'H',
     '코디네이터가 환자 메시지 수신·응답. AI 이관 케이스 처리. 상담 이력 기록.',
     '완료 · /app/coordinator/messages, migrations/20260225_chat_threads'),

    ('FR-12', 'WebRTC 화상상담 (LiveKit)', 'H',
     '브라우저 기반 화상·음성 상담. LiveKit Cloud 연동. 게스트 토큰으로 비회원 참여. (사업계획서 p.28, 공고문 p.8)',
     '완료 · /app/consultation/[id] (상담방), /app/c/[code] (초대 주소), /app/api/livekit'),

    ('FR-13', 'AI 실시간 번역 (상담 중)', 'H',
     '화상상담 및 채팅 중 6개 언어 실시간 번역·자막(화자·시각·원어 표시), 통역 음성 봇. 특허 10-2868334 기반. (사업계획서 p.28)',
     '완료 · /app/api/translate-text, /app/api/khidi/consultation/[id]/{stt,interpreter}, agents/live-translate (통역 자막 3,735건 축적)'),

    ('FR-14', '진료 예약·일정 관리', 'H',
     '상담 예약(환자·시각), 초대 메일(.ics), 구글 캘린더 자동 등록, 리마인더 자동 발송, 환자 달력.',
     '부분구현 · /app/coordinator/consultations, src/lib/calendar/googleCalendar.ts, /app/api/cron/dispatch-reminders (남은 것 = 병원 가용 일정 입력 화면)'),

    ('FR-15', '비자 발급 안내', 'M',
     '카자흐스탄→한국 의료비자 신청 절차, 필요 서류, 처리기간 안내. 다국어 제공.',
     '완료 · /app/patient/visa, /app/visa, migrations/20260406_education_visa_rebooking'),

    ('FR-16', '경과 모니터링 (f/u)', 'H',
     '귀국 후 환자 건강상태 주기적 체크인. 증상 입력, 검사결과 업로드, 규칙 기반 위험도 판정(응급 키워드+점수식 — AI 호출 없음)·3일 침묵 경보·담당자 알림. (사업계획서 p.29, 공고문 p.8)',
     '완료 · /app/patient/symptoms, /app/api/khidi/followup, /app/api/cron/detect-silent-patients, src/lib/followup/symptomAnalyzer.ts'),

    ('FR-17', '건강관리 교육 콘텐츠', 'M',
     '암 유형별 맞춤 사후관리 가이드, 식이요법, 복약 안내. 러시아어 콘텐츠 제공. (사업계획서 p.30)',
     '완료 · /app/education, /app/admin/education, /app/api/cron/dispatch-surveys (콘텐츠 18건·6개 언어. 단계별 자동 발송 D+7·14·30·90·180 연결 2026-08-25, 실환자 발송 0)'),

    ('FR-18', '재방문 예약 (Rebooking)', 'M',
     '경과관리 기반 재방문 필요성 자동 알림. 재진 예약·비자 재발급 안내 원스톱. (사업계획서 p.30)',
     '완료 · /app/patient/rebooking, migrations/20260406_education_visa_rebooking'),

    ('FR-19', '6개 언어 UI/UX', 'H',
     '한국어·영어·러시아어·카자흐어·중국어·일본어. Next.js App Router 다국어 라우팅 (/ru, /kz 등).',
     '완료 · 언어 접두어는 proxy.ts 가 처리한다(/ru/treatments → 내부 /treatments 로 넘기고 x-locale 머리값으로 언어 전달)'),

    ('FR-20', 'DB 콘텐츠 다국어 자동번역', 'H',
     '병원·치료 정보 i18n JSONB 컬럼 자동번역. Gemini 기반 배치 번역.',
     '완료 · /app/api/translate, src/lib/translate.ts(triggerMultiLangTranslation), scripts/batch-translate-all.mjs, /app/coordinator/content'),

    ('FR-21', '환자 대시보드', 'H',
     '환자 본인 상담 이력, 예약, 문서, 비용, 만족도 조회. 전체 여정 통합 뷰.',
     '완료 · /app/patient/*, PatientDashboardClient.jsx'),

    ('FR-22', '코디네이터 포털', 'H',
     '담당 환자 관리, 상담 처리, 실적 현황, 알림 수신.',
     '완료 · /app/coordinator/*'),

    ('FR-23', '국내 의료기관 포털', 'H',
     '협진 의뢰(리드) 수신·상태 갱신(치료 확정 → 유치 자동 반영)·견적 회신·CSV. 병원 정보·시술 카탈로그는 플래그 비활성(공개 연동 준비 중), 진료 결과 입력·가용 일정 화면은 없다. 의료진 등록은 /admin/doctors.',
     '부분구현 · /app/hospital/{,leads} (구 파트너 경로에서 개명), app/hospital/_components/featureFlags.js'),

    ('FR-24', '관리자(Admin) 포털', 'H',
     '사용자·병원·KPI·AI 성능 관리. 성과보고 자료 출력. 감사 로그.',
     '완료 · /app/admin/*, migrations/20260129_add_admin_audit_logs'),

    ('FR-25', '이메일 알림·리마인더', 'M',
     '접수 확인, 상담 초대(.ics), 리마인더, 설문, 교육 콘텐츠, 관리자 알림. Resend(폴백 AWS SES) + TS 템플릿 6개 언어. 발신 도메인 healwith.co.kr.',
     '완료 · src/lib/email/sendEmail.ts, src/lib/email/templates/*, /app/api/cron/dispatch-* (경로 /app/api/email 은 없다 — 옛 표기)'),

    ('FR-26', '실시간 In-app 알림', 'M',
     '신규 문의·환자 새 글·소견 도착·증상 경보·식은 문의 인앱 알림(벨, 30초 폴링) + 스토어 앱 FCM 푸시.',
     '완료 · src/lib/notifications/{inApp,pushBridge,pushPolicy}.ts, src/lib/push/*, src/components/NotificationBell.jsx (Realtime 은 상담방 메시지에만)'),

    ('FR-27', '환자 PII 암호화', 'H',
     '환자 성명·연락처·의료정보 AES-256-GCM 암호화 저장 (암호화 칸 20개(*_encrypted, voice_notes 3칸은 2026-09-04 추가) + inquiries 의 first_name·last_name·email·phone). 값은 전건 암호문이나 칸 이름은 평문 시절 그대로.',
     '완료 · src/lib/security/encryptionV2.ts (평문 칸 삭제 마이그레이션 20260420_drop_* 은 실행 금지 — 돌리면 26~29건 소실)'),

    ('FR-28', 'API Rate Limiting', 'H',
     '공개 POST 엔드포인트 IP 기반 요청 제한(기본 분당 20회, 관문별 5~30). 인스턴스별 메모리가 아니라 공용 저장소(rate_limit_buckets)로 센다.',
     '완료 · src/lib/rateLimit.ts, check:ratelimit-scope'),

    # ── 2026-09-06 추가: 착수 뒤 자체 도출·구현된 요구 (사업계획서에 없던 것) ──
    ('FR-29', '해외 파트너 포털 2계층', 'H',
     '해외 에이전시(/agency)·해외 의료기관(/clinic) 전용 포털. 의료기관만 검사결과·영상·소견을 올린다(사후관리 ICT ④). 케이스 메신저·단계 진행 확인.',
     '완료 · /app/agency, /app/clinic, /app/api/agency/*'),
    ('FR-30', '전문의 소견 수집·전달', 'H',
     '계정 없는 협진 전문의가 링크로 소견을 작성·제출 → 환자 언어 번역 초안 → 코디 교정 → 공식 문서로 환자 전달. 전달 시점이 K-02 판정 기준.',
     '완료 · /app/opinion/[token], /app/api/opinions/*, src/lib/opinions'),
    ('FR-31', '환자 진행상황 화면(가입 불필요)', 'H',
     '접수한 사람 누구나 링크로 6단계 진행·소식·우리가 준 문서·추가 자료 올리기. 환자 글은 코디에게 배지·알림.',
     '완료 · /app/claim/[token], /app/api/inquiries/claim/*'),
    ('FR-32', '만족도 설문 자동 발송·집계', 'H',
     '상담 완료 +24h 에 5문항 설문 발송, /survey/<토큰> 응답, 코디·관리자 화면 집계, 시험 데이터 제외, 표본부족 가드. K-03 의 유일한 원천.',
     '완료 · /app/survey/[token], /app/api/cron/dispatch-surveys, src/lib/surveys (실환자 응답 0)'),
    ('FR-33', '스토어 앱(iOS/Android)', 'M',
     'Capacitor 라이브로드 앱 — 푸시·딥링크·네이티브 로그인·계정 삭제·구판 감지. 스토어 요구사항(광고 ID 없음·AI 고지) 반영.',
     '완료 · capacitor.config.ts, android/, ios/, codemagic.yaml, /app/app (구글 검토·애플 심사 대기)'),
    ('FR-34', '음성 메모 판독', 'M',
     '왓츠앱·텔레그램 음성 메모를 AI 가 받아써 코디가 듣지 않고 읽는다. 문의 생성 전에도 가능. 연락처·국적 자동 채움.',
     '완료 · /app/coordinator/voice, /app/api/coordinator/voice-notes, voice_notes(암호화)'),
    ('FR-35', '콘텐츠 편집기·번역 품질 가드', 'H',
     '코디가 화면 문구·치료법·암종·FAQ·병원 소개 등 249칸을 6개 언어로 직접 수정. CI 가 ru·kz 누락·사실 유실·용어 흔들림을 차단. 검수 전 「제안」 표시.',
     '완료 · /app/coordinator/content, src/lib/i18n/glossary.js, scripts/check-i18n-*.mjs'),
    ('FR-36', '유치 운영 자동화', 'H',
     '환자 새 글 알림·배지, 7일 무동작 「식은 문의」 감지, 결과(유치 확정/종료/이탈) 기록, 구글 캘린더 자동 등록, 시험 문의 자동 정리.',
     '완료 · /app/api/cron/{detect-cold-leads,kpi-snapshot}, /app/coordinator/inbox, src/lib/khidi/inquiryOutcome.ts'),
    ('FR-37', '검색엔진 색인 자동 제출', 'M',
     'IndexNow(빙·얀덱스·네이버) 매일 자동 제출, 사이트맵 170쪽·hreflang·언어 접두어 링크 강제. 구글은 별도.',
     '완료 · /app/api/cron/indexnow, src/lib/seo/indexNow.ts (효과는 2주 뒤 재측정)'),
]

for row in fr_data:
    add_data_row(fr_tbl, list(row), bold_first=True)

doc.add_paragraph()
add_heading(doc, '3-1. 추적표: 요구 → 화면 → 기능 명세', 2)
add_para(doc, '향후 개발자가 요구 하나를 잡으면 어느 화면(01-1)과 어느 기능 명세(02)를 봐야 하는지 한 줄로 찾게 한다.')
trace_tbl = doc.add_table(rows=0, cols=4)
trace_tbl.style = 'Table Grid'
add_header_row(trace_tbl, ['FR', '화면(01-1)', '기능 명세(02)', '6대 ICT'])
for r in [
    ('FR-01~03', 'P-01, P-04(초대)', 'FN-AUTH-01~03', '기반'),
    ('FR-04~06', 'P-02, P-07', 'FN-INTAKE-01~02, FN-DOC-01', '②③'),
    ('FR-07~10', 'P-01, P-03', 'FN-MATCH-01~02, FN-CHAT-01, FN-RAG-01', '①②'),
    ('FR-11', 'P-06, P-07', 'FN-CHAT-02', '②'),
    ('FR-12~13', 'P-04', 'FN-VIDEO-01~02, FN-TRANS-01~02', '②④'),
    ('FR-14~15', 'P-07, P-08', 'FN-SCHED-01~02', '③'),
    ('FR-16~18', 'P-08', 'FN-POST-01~04', '④⑤⑥'),
    ('FR-19~20', '전 화면', 'FN-I18N-01~02, FN-CONTENT-01', '기반'),
    ('FR-21', 'P-08, P-11', 'FN-CLAIM-01', '②③'),
    ('FR-22', 'P-06, P-07', 'FN-COORD-01~02', '운영'),
    ('FR-23', '—', 'FN-PARTNER-01', '②③'),
    ('FR-24', 'P-09', 'FN-ADMIN-01', '운영'),
    ('FR-25~26', 'P-06, P-08', 'FN-NOTIF-01~02', '운영'),
    ('FR-27~28', '—', 'FN-SEC-01~03', '기반'),
    ('FR-29', '—', 'FN-PARTNER-02~03', '②④'),
    ('FR-30', 'P-05, P-07', 'FN-OPINION-01', '②'),
    ('FR-31', 'P-11', 'FN-CLAIM-01', '②③'),
    ('FR-32', 'P-08', 'FN-SURVEY-01', '⑤'),
    ('FR-33~34', 'P-08, P-06', 'FN-APP-01, FN-VOICE-01', '④⑤ · ②'),
    ('FR-35', '—', 'FN-CONTENT-01', '①'),
    ('FR-36', 'P-06, P-09', 'FN-COORD-01, FN-OPS-01', '②③'),
    ('FR-37', 'P-01', 'FN-SEO-01', '①'),
]:
    add_data_row(trace_tbl, list(r), bold_first=True)

doc.add_page_break()

# ================================================================
# 4. 비기능 요구사항
# ================================================================
add_heading(doc, '4. 비기능 요구사항', 1)

# 실측 — 어떻게 쟀는지까지 적는다. 못 잰 것은 왜 못 쟀는지 적는다.
# 2026-08-20~21 실측 + 2026-09-06 재실측(라우트 수·i18n·자가시험·감시·접근성 범위). 값마다 날짜를 단다.
NFR_VERDICT = {
 'NFR-01': '충족(실사용 관측) — 실제 브라우저 측정 LCP 0.42초, 데스크톱 5회 중앙값 1.12초(2026. 8. 21.). 채점 도구가 저속 4G 회선과 4배 느린 기기를 가정해 계산한 시뮬레이션 값은 5.31초(이후 실측 5.9~6.0초)로 목표를 넘고, 모바일 점수 편차(74~85)의 원인은 총 차단 시간(TBT 80~3,850ms)이다. 두 값을 갈라 적는다. 시뮬레이션 기준 개선은 남은 일',
 'NFR-02': '충족(경계) — ①실서비스 스트리밍 10건(2026. 8. 21.): 첫 글자 중앙값 2.30초 · 상위 5% 3.84초 ②같은 경로 자가시험 50건(2026. 8. 21.): P50 3.90초 · P95 5.33초 · 5초 초과 5건(10%). 표본이 큰 ②를 기준으로 하면 목표를 넘는 회차가 있다. 8/24~9/03 실행은 옛 경로라 첫 글자를 못 쟀고 9/05 복원(#1617) 뒤 9/07 실행부터 다시 잡힌다',
 'NFR-03': '충족(망 구간 실측) — 카자흐스탄 현지 측정망으로 2026. 8. 21. 측정. 알마티·악타우·파블로다르 등 7개 지점에서 화상 서버까지 왕복 중앙값 306ms, 한국 5개 지점에서 같은 서버까지 33ms. 한쪽 방향 망 지연 170ms 이며, 여기에 영상 부호화·흔들림 완충(통상 40~80ms)을 더해도 210~250ms 로 목표 300ms 이내다. 망 구간은 실측이고 단말 처리 구간은 표준값 추정이다',
 'NFR-04': '충족(부분) — 실서비스 상태 확인 주소로 동시 100건 요청: 100/100 성공, 응답 중앙값 417ms · 상위 5% 1,650ms. 다만 이는 읽기 전용 가벼운 요청이며, 화면 렌더·DB 쓰기를 포함한 본격 부하 시험은 실사용자에게 영향이 가므로 유치 확대 시점에 별도 진행',
 'NFR-05': '충족 — 운영DB 조회: 성명(first_name·last_name)·이메일·전화 값이 있는 행 전부 암호문, 평문 0건',
 'NFR-06': '충족 — 자동 검사 check:err-exposure 통과(매 변경 실행)',
 'NFR-07': '충족 — API 라우트 231개(2026. 9. 6.) 전수 확인. 관문 흔적이 없는 것은 공개 조회·상태 확인·서명 검증·크론(CRON_SECRET)·폐쇄된 라우트뿐. 2026-09-05 anon 이 부를 수 있던 RPC 4종의 EXECUTE 회수(#1622)',
 'NFR-08': '충족 — 권한 판정에 user_metadata 를 쓰는 곳 0건(전수 검색). app_metadata.role 기준',
 'NFR-09': '충족 — 자동 검사 check:ratelimit-scope 통과',
 'NFR-10': '충족 — 실서비스 머리값 확인: Strict-Transport-Security(2년·하위도메인·preload) · CSP · X-Content-Type-Options · X-Frame-Options',
 'NFR-11': '충족 — 외부 감시 UptimeRobot(5분, 2026-09-04 개설) + 깃허브 예약 감시(선언은 10분, 실측 하루 7회 — 이 저장소의 예약 큐 특성). 깃허브 감시 최근 40회 전건 성공(2026-09-06). 7/24 DB 장애 55분은 감시 도입 전 사고',
 'NFR-12': '충족 — Supabase Pro 플랜 자동 백업(7일 보관). 플랜 상태는 콘솔에서 확인',
 'NFR-13': '충족 — sentry.client·server·edge.config.js 세 개 모두 존재',
 'NFR-14': '충족 — 화면이 쓰는 문구 2,052개(2026. 9. 6.)를 6개 언어에 전건 채움. 러시아어·카자흐어는 자동 검사가 매 변경마다 100% 확인. 콘텐츠 칸(치료법·암종·병원 소개) 249개는 2026-09-06 AI 번역으로 채워 원어민 검수 대기',
 'NFR-15': '충족(공개 화면) — 라이트하우스 접근성 100점(2026. 8. 21.), axe-core WCAG 2.1 AA 주 1회 자동. 로그인 뒤 화면 22개는 8/17~8/31 3주 연속 측정 실패(networkidle 함정)였고 2026-09-06 #1648 로 고쳐 9/07 실행이 첫 재측정 — 그 범위는 «미측정»으로 적는다',
 'NFR-16': '충족 — 375px 폭 기준 글자 잘림·겹침 자동 검사 상시 실행(content-clip-sweep · header-no-overlap)',
 'NFR-17': "충족 — tsconfig paths: {'@/*': ['src/*']}",
 'NFR-18': '충족 — strict:false 유지, Zod 런타임 검증 사용',
 'NFR-19': '충족 — package.json build 가 next build --webpack',
 'NFR-20': '충족 — /api/rag/ingest 로 코드 수정 없이 올린다. 현재 문서 63건·조각 105개(2026. 9. 6.). 자가시험이 RAG 조각이 프롬프트에 «실제로» 들어가는지 매회 기록',
 'NFR-21': '충족 — 웹 접수 6건 전부 동의 기록과 동의 판번호 보존. 나머지 2건은 플랫폼 밖에서 받아 소급 등록한 건이라 화면 동의 기록이 없다(별도 서면 보관)',
 'NFR-22': '유지 — 문구·화면 모두 정보 제공·연결 서비스로 표기. 법률 판단은 외부 검토 대상',
 'NFR-23': '진행 중 — 유치 0/12 · 사전상담+사후관리 7/120(영상 1 + 전달된 소견 6) · 만족도 실환자 표본 0건(2026. 9. 6. 운영DB). 잔여 기간 유입에 달려 있다',
}

nfr_cats = [
    ('성능 (Performance)', [
        ('NFR-01', '페이지 초기 로드 ≤ 3초 (LCP 기준, 4G 환경)'),
        ('NFR-02', 'AI 챗봇 첫 토큰 응답 ≤ 5초'),
        ('NFR-03', 'WebRTC 화상상담 지연 ≤ 300ms (카자흐스탄↔한국)'),
        ('NFR-04', '동시 접속 100명+ 처리 (Vercel Edge + Supabase 자동 확장)'),
    ]),
    ('보안 (Security)', [
        ('NFR-05', '환자 PII AES-256-GCM 암호화 (개인정보보호법 준수)'),
        ('NFR-06', 'API 응답에 error.message 직접 노출 금지 → 코드형 오류만 반환'),
        ('NFR-07', '모든 API 엔드포인트 인증 헬퍼 적용 (requireAdminAuth 등)'),
        ('NFR-08', 'RBAC: app_metadata.role 기준 (user_metadata 금지)'),
        ('NFR-09', '공개 POST API Rate Limiting (IP 기반, rateLimit.ts)'),
        ('NFR-10', 'HTTPS 전용, HSTS 헤더 적용'),
    ]),
    ('가용성 (Availability)', [
        ('NFR-11', '월 가용성 99.5%+ (Vercel SLA 기준)'),
        ('NFR-12', 'DB 자동 백업 일일 (Supabase 기본 제공)'),
        ('NFR-13', '장애 감지 및 알림 (@sentry/nextjs 연동)'),
    ]),
    ('다국어·접근성 (i18n/a11y)', [
        ('NFR-14', '6개 언어 UI 지원 (ko/en/ru/kz/zh/ja)'),
        ('NFR-15', 'WCAG 2.1 AA 기준 접근성 기본 준수'),
        ('NFR-16', '모바일 반응형 (375px 이상 전 화면 지원)'),
    ]),
    ('유지보수·확장성', [
        ('NFR-17', 'Path alias @/* → src/* 일관 적용'),
        ('NFR-18', 'TypeScript strict:false 유지 (점진적 전환), Zod 런타임 검증'),
        ('NFR-19', '빌드: npx next build --webpack 전용 (Turbopack 금지)'),
        ('NFR-20', 'RAG 벡터DB에 신규 의료 지식 비코드 업로드 가능 구조'),
    ]),
    ('법적 준수 (Compliance)', [
        ('NFR-21', '개인정보보호법(PIPA) 준수: 명시적 동의, 최소 수집, 보관 기간 준수'),
        ('NFR-22', '의료법 준수: 원격 사전상담은 정보 제공·연결 서비스로 포지셔닝'),
        ('NFR-23', 'KHIDI KPI 달성: 유치 12건, 상담 120건, 만족도 90점 (수행계획서 확정본 목표. 공고문 최소치는 10·80·80)'),
    ]),
]

for cat_name, items in nfr_cats:
    add_heading(doc, cat_name, 2)
    ntbl = doc.add_table(rows=0, cols=3)
    ntbl.style = 'Table Grid'
    add_header_row(ntbl, ['NFR-ID', '요구사항 내용', '검증 결과 (2026. 8. 20.~21. 실측 · 2026. 9. 6. 재실측 — 값마다 날짜 표기)'])
    for nfr_id, nfr_desc in items:
        add_data_row(ntbl, [nfr_id, nfr_desc, NFR_VERDICT.get(nfr_id, '미검증')], bold_first=True)
    doc.add_paragraph()

doc.add_page_break()

# ================================================================
# 5. 외부 시스템 연계
# ================================================================
add_heading(doc, '5. 외부 시스템 연계', 1)

ext_tbl = doc.add_table(rows=0, cols=4)
ext_tbl.style = 'Table Grid'
add_header_row(ext_tbl, ['시스템', '용도', '연동 방식', '코드 위치'])

ext_data = [
    ('Supabase (PostgreSQL 17.6)',
     'DB, Auth, Storage, Realtime, RLS\n환자 정보·상담·병원 데이터 저장',
     '@supabase/ssr: SSR 쿠키 기반 세션\nservice_role: server-only 모듈',
     'src/lib/supabase/*'),
    ('Google Gemini Flash (최신 별칭)',
     'AI 챗봇, RAG 응답 생성\n다국어 번역, 문서 분석',
     'Vercel AI SDK (@ai-sdk/google)\nstreamText / generateText',
     'src/lib/chat/generateReply.ts'),
    ('LiveKit Cloud',
     'WebRTC 화상상담 서버\n오디오/비디오/화면공유',
     'livekit-server-sdk (토큰 발급)\n@livekit/components-react (UI)',
     '/app/api/khidi/consultation/token\n/app/api/livekit/webhook\n/app/consultation/[id]'),
    ('Resend (폴백 AWS SES)',
     '이메일 발송\n접수 확인, 상담 초대, 리마인더, 설문, 교육, 관리자 알림',
     'Resend SDK → 실패 시 @aws-sdk/client-ses\n템플릿은 TS 함수(react-email 아님), 발신 healwith.co.kr',
     'src/lib/email/*'),
    ('텔레그램 · 왓츠앱 (Cloud API)',
     '환자 상담 채널(봇) + 코디 양방향 릴레이',
     '웹훅 서명 검증 · AI 자동응답 · 상담기록 축적',
     '/app/api/webhooks/{telegram,whatsapp}, src/lib/messaging/*'),
    ('Firebase Cloud Messaging',
     '스토어 앱 푸시',
     '기기 등록(device_tokens) · 옵트인 · 정책',
     'src/lib/push/*, /app/api/push/*'),
    ('Google Calendar API',
     '화상상담 예약 자동 등록',
     '서비스 계정이 소유한 캘린더를 공유',
     'src/lib/calendar/googleCalendar.ts'),
    ('UptimeRobot',
     '실서비스 생사 감시(5분)',
     'HEAD /api/health, 2xx·3xx 만 정상',
     '외부 대시보드(PO 계정)'),
    ('IndexNow (빙·얀덱스·네이버)',
     '검색엔진 색인 자동 제출',
     '공개 키 파일 + 매일 크론',
     'src/lib/seo/indexNow.ts'),
    ('Google Maps API',
     '병원 위치 지도 표시\n병원 주변 정보',
     '@react-google-maps/api\nClient-side 렌더링',
     '/app/hospitals/*'),
    ('Sentry',
     '오류 추적·성능 모니터링\n배포 후 장애 감지',
     '@sentry/nextjs + withSentryConfig',
     'sentry.client·server·edge.config.js'),
    ('Capacitor (iOS/Android)',
     '웹→모바일 앱 래핑\n푸시 알림, 딥링크',
     '@capacitor/core\n@capacitor/push-notifications',
     'capacitor.config.ts'),
    ('HIRA 공공 API · 네이버 지역검색',
     'RAG 2계층: 답할 때 실시간 보조검색\n(크롤 파이프라인은 별도, 지금은 안 돈다)',
     'HTTP REST 병렬 호출, 짧은 타임아웃',
     'src/lib/chat/externalSearch.ts\nsrc/lib/crawl/sources/*'),
    ('AES-256-GCM (자체 구현)',
     '환자 PII 암호화·복호화\n암호화 칸 데이터 보호(*_encrypted 20개 + inquiries 의 성명·이메일·전화)',
     'Web Crypto API 기반 자체 구현',
     'src/lib/security/encryptionV2.ts'),
]
for row in ext_data:
    add_data_row(ext_tbl, list(row), bold_first=True)

doc.add_page_break()

# ================================================================
# 6. 제약사항
# ================================================================
add_heading(doc, '6. 제약사항', 1)

constraints_data = [
    ('법적 제약', [
        '개인정보보호법(PIPA): 환자 개인정보 수집·이용 시 명시적 동의 필수. 수집 최소화 원칙. 보관 기간 이후 파기 의무.',
        '의료법: 국내 원격 진료 규제를 고려하여 사전상담·사후상담은 「의료 정보 제공 및 연결 서비스」로 포지셔닝. 실제 진료·처방은 면허 의사와 협약 기반 별도 수행.',
        '카자흐스탄 개인정보보호법(2013): 카자흐스탄 국민 개인정보 처리 시 현지 법령 준거.',
        '의료기기법: AI 진단 기능이 아닌 정보 제공·매칭·상담 플랫폼으로 설계 (의료기기 해당 없음).',
    ]),
    ('KHIDI 사업 제약 (공고문 p.2~3)', [
        '성과지표 의무: 외국인 환자 유치 10건+ (목표 12건), 사전상담·사후관리 80건+ (목표 120건), 만족도 80점+ (목표 90점).',
        '원격협진 서비스: 4개월 이상 실시 필수.',
        '국고보조금 SW 개발비 필수 편성 (SW 개발·고도화 인건비 포함).',
        '사업 산출물 소유권: 보건복지부·KHIDI 소유 (공고문 p.6).',
        '중복 사업 금지: 타 공공기관 지원 사업과 동일·중복 불가.',
    ]),
    ('기술적 제약', [
        'Turbopack 빌드 금지: npx next build --webpack 전용.',
        'TypeScript strict:false 유지 (점진적 전환 중).',
        'LiveKit Cloud 의존: 화상상담은 LiveKit SaaS 의존. 자체 서버 운영 미계획.',
        '카자흐스탄 네트워크: 지방 WebRTC 품질은 현지 5G 인프라가 2025년 말 구축 완료되어 개선 추세이나, 지방은 여전히 회선 편차가 있음. 2026-08-21 현지 측정망 실측 왕복 306ms(7개 지점 중앙값).',
        '구글 검색 그라운딩 미작동: 설치된 @ai-sdk/google 에 없는 옵션이라 조용히 무시된다. 살리려면 정식 googleSearch 도구로.',
    ]),
    ('운영적 제약', [
        '사업 기간: 협약체결일 ~ 2026.11.20 (단년도 사업).',
        '총 인력: 12명 (기존 8명 + 신규 4명). 신규 참여연구원 채용 별지 서식 제출 필요.',
        '자기부담금: 총사업비의 20% 이상 매칭 (현금 5%+, 현물 15% 이하). 총사업비 1억원 기준 자기부담 2천만원.',
    ]),
]

for cat, items in constraints_data:
    add_heading(doc, cat, 2)
    for item in items:
        add_para(doc, f'• {item}', indent=0.5)
    doc.add_paragraph()

doc.add_page_break()

# ================================================================
# 7. 용어 정의
# ================================================================
add_heading(doc, '7. 용어 정의', 1)

term_tbl = doc.add_table(rows=0, cols=2)
term_tbl.style = 'Table Grid'
add_header_row(term_tbl, ['용어 / 약어', '정의'])

terms = [
    ('healwith', '본로이가 개발·운영하는 AI 기반 외국인환자 상담·매칭·사후관리 통합 플랫폼명(healwith.co.kr). 옛 이름은 2026-08 에 버렸고 병기하지 않는다'),
    ('ICT', 'Information and Communication Technology. 정보통신기술'),
    ('RAG', 'Retrieval-Augmented Generation. 검색 기반 AI 응답 생성. healwith 는 자체 DB → HIRA·네이버 실시간 보조검색 2계층이 가동 중이고 구글 검색 그라운딩(3계층)은 미작동'),
    ('AI Agent', 'AI 기반 자동 상담 봇. Gemini Flash 기반. 복잡 케이스 시 Human Agent로 이관'),
    ('Human Agent', '코디네이터. AI가 처리하지 못한 상담을 직접 수행'),
    ('WebRTC', 'Web Real-Time Communication. 브라우저 기반 실시간 화상·음성 통신 기술'),
    ('LiveKit', 'WebRTC 기반 화상상담 SaaS 플랫폼. 서버/클라이언트 SDK 제공'),
    ('RLS', 'Row Level Security. Supabase PostgreSQL 기반 행 단위 접근 제어'),
    ('RBAC', 'Role-Based Access Control. 역할 기반 접근 제어'),
    ('PII', 'Personally Identifiable Information. 개인식별정보 (성명, 연락처, 의료정보 등)'),
    ('AES-256-GCM', '대칭키 암호화 방식. 환자 PII 암호화에 사용'),
    ('Supabase', 'PostgreSQL 기반 BaaS. DB, Auth, Storage, Realtime 통합 제공'),
    ('Resend', '개발자용 이메일 발송 API 서비스'),
    ('Capacitor', 'Ionic 기반 웹→iOS/Android 네이티브 앱 변환 프레임워크'),
    ('KHIDI', '한국보건산업진흥원. Korea Health Industry Development Institute'),
    ('PIPA', '개인정보보호법. Personal Information Protection Act (대한민국)'),
    ('인테이크(Intake)', '환자가 플랫폼에 최초 정보를 입력하는 단계. 통합 문의 퍼널(/inquiry) 내에서 수행'),
    ('사전상담', '환자 내원 전 병원 정보 제공, 진료의뢰, 예약 안내 포함 ICT 서비스 (공고문 p.8)'),
    ('사후관리', '환자 귀국 후 경과 추적, 모니터링, 교육, 재방문 연계 ICT 서비스 (공고문 p.8)'),
    ('협진', '면력한방병원(면역치료)과 이대서울·고대구로 등 상급종합병원 간 연계 진료'),
    ('STT', 'Speech-To-Text. 음성→텍스트 변환. 상담방 실시간 자막과 코디 음성 메모 판독에 사용'),
    ('MVP', 'Minimum Viable Product. 핵심 기능 구현 초기 제품. 2026-06 배포·실서비스 운영 중'),
    ('PoC', 'Proof of Concept. 개념 검증 단계. 사업계획서 p.4 기준 완료 상태'),
    ('CIS', '독립국가연합. 카자흐스탄·러시아·우즈베키스탄 등 구소련 국가'),
    ('SaaS', 'Software as a Service. 클라우드 기반 소프트웨어 서비스 제공 방식'),
    ('pgvector', 'PostgreSQL 벡터 검색 확장. RAG 임베딩 벡터 저장·검색에 활용'),
    ('i18n', 'Internationalization. 다국어 지원 설계. HEALO는 6개 언어 지원'),
    ('Gemini Flash', 'Google DeepMind 의 멀티모달 AI 모델. healwith 는 최신 별칭(gemini-flash-latest)으로 고정 사용'),
]
for term, defn in terms:
    add_data_row(term_tbl, [term, defn], bold_first=True)

# 저장
import os as _os
# 저장 위치는 «이 스크립트가 있는 폴더» 기준으로 잡는다.
# (전에는 특정 PC 의 절대경로가 박혀 있어 그 PC 밖에서는 재생성이 아예 불가능했다.)
out_path = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), '01_요구사항정의서.docx')
doc.save(out_path)
print(f'저장 완료: {out_path}')
print(f'총 단락 수: {len(doc.paragraphs)}')
